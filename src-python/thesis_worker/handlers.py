"""
@file: handlers.py
@description: NDJSON 请求路由。收到 {id, cmd, ...} → 返回 {id, ok, result?, error?, code?}
              错误处理原则（spec Section 7）：不降级，异常直接抛出成 response
              - 规则异常 → 整批中止，不跳过其它规则
              - 文件错误 → 明确 code (ENOENT / EPERM / PARSE_ERROR)
@author: Atlas.oi
@date: 2026-04-17
"""
import traceback
from pathlib import Path
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

def handle(req: dict) -> dict:
    """顶层路由。任何未捕获异常都转成 code=INTERNAL 的 response（sidecar 进程不崩）"""
    req_id = req.get('id')
    cmd = req.get('cmd')

    try:
        if cmd == 'ping':
            return {'id': req_id, 'ok': True, 'result': 'pong'}

        if cmd == 'detect':
            return _handle_detect(req_id, req)

        if cmd == 'fix':
            return _handle_fix(req_id, req)

        if cmd == 'fix_preview':
            return _handle_fix_preview(req_id, req)

        if cmd == 'cancel':
            # P3 单线程串行，cancel 只做 ack；真实中断留 P4 实现
            return {'id': req_id, 'ok': True, 'result': {'cancelled': True}}

        if cmd == 'extract_all':
            return _handle_extract_all(req_id, req)

        if cmd == 'extract_from_selection':
            return _handle_extract_from_selection(req_id, req)

        if cmd == 'list_fields':
            return _handle_list_fields(req_id)

        return {
            'id': req_id, 'ok': False,
            'error': f'unknown cmd: {cmd}',
            'code': 'UNKNOWN_CMD',
        }
    except Exception as e:
        return {
            'id': req_id, 'ok': False,
            'error': f'{type(e).__name__}: {e}\n{traceback.format_exc()}',
            'code': 'INTERNAL',
        }


def _handle_detect(req_id: str, req: dict) -> dict:
    from .engine_v2.detector import detect_v2
    file = req['file']
    template = req['template']

    # ENOENT 检查必须在 Document() 前，否则 python-docx 的报错信息不统一
    if not Path(file).exists():
        return {'id': req_id, 'ok': False, 'error': f'file not found: {file}', 'code': 'ENOENT'}

    # 打开一次确认文件格式有效（detect_v2 内部也会打开，此处保证 PARSE_ERROR / EPERM 语义对齐）
    try:
        Document(file)
    except PackageNotFoundError:
        return {'id': req_id, 'ok': False, 'error': f'docx malformed: {file}', 'code': 'PARSE_ERROR'}
    except PermissionError as e:
        return {'id': req_id, 'ok': False, 'error': str(e), 'code': 'EPERM'}

    try:
        issues = detect_v2(file, template)
    except Exception as e:
        return {
            'id': req_id, 'ok': False,
            'error': f'detect_v2 raised: {type(e).__name__}: {e}\n{traceback.format_exc()}',
            'code': 'RULE_ERROR',
        }

    # 给每个 issue 分配稳定 id（rule_id-全局偏移，跨字段唯一）
    for idx, issue in enumerate(issues):
        issue['issue_id'] = f"{issue['rule_id']}-{idx}"

    return {'id': req_id, 'ok': True, 'result': {'issues': issues}}


def _handle_fix(req_id: str, req: dict) -> dict:
    """v2 fix：按 issue.attr + value 修改 docx 段落属性，成功后蓝色标记

    业务逻辑：
    1. ENOENT / PARSE_ERROR / EPERM 前置检查
    2. 调用 fix_v2 执行修改并写回文件
    3. 返回 {diff, applied, xml_changed}
    """
    from .engine_v2.fixer import fix_v2
    file = req['file']
    issue = req['issue']
    value = req['value']

    if not Path(file).exists():
        return {'id': req_id, 'ok': False, 'error': f'file not found: {file}', 'code': 'ENOENT'}
    try:
        Document(file)
    except PackageNotFoundError:
        return {'id': req_id, 'ok': False, 'error': f'docx malformed: {file}', 'code': 'PARSE_ERROR'}
    except PermissionError as e:
        return {'id': req_id, 'ok': False, 'error': str(e), 'code': 'EPERM'}

    try:
        result = fix_v2(file, issue, value)
    except Exception as e:
        return {
            'id': req_id, 'ok': False,
            'error': f'fix_v2 raised: {type(e).__name__}: {e}\n{traceback.format_exc()}',
            'code': 'RULE_ERROR',
        }

    return {'id': req_id, 'ok': True, 'result': result}


def _handle_fix_preview(req_id: str, req: dict) -> dict:
    """v2 fix_preview：在临时副本上执行修复，返回 diff，原文件不动

    业务逻辑：
    1. ENOENT / PARSE_ERROR / EPERM 前置检查
    2. 复制文件到临时路径，在副本上调用 fix_v2
    3. 取出 diff 后删除临时文件，applied 标记置 False
    """
    from .engine_v2.fixer import fix_v2
    import shutil
    import os
    import tempfile
    file = req['file']
    issue = req['issue']
    value = req['value']

    if not Path(file).exists():
        return {'id': req_id, 'ok': False, 'error': f'file not found: {file}', 'code': 'ENOENT'}
    try:
        Document(file)
    except PackageNotFoundError:
        return {'id': req_id, 'ok': False, 'error': f'docx malformed: {file}', 'code': 'PARSE_ERROR'}
    except PermissionError as e:
        return {'id': req_id, 'ok': False, 'error': str(e), 'code': 'EPERM'}

    # 复制到临时文件上执行 fix，diff 回读后删除，原文件不动
    tmp_fd, tmp_path = tempfile.mkstemp(suffix='.docx')
    os.close(tmp_fd)
    try:
        shutil.copy(file, tmp_path)
        try:
            result = fix_v2(tmp_path, issue, value)
        except Exception as e:
            return {
                'id': req_id, 'ok': False,
                'error': f'fix_v2 raised: {type(e).__name__}: {e}\n{traceback.format_exc()}',
                'code': 'RULE_ERROR',
            }
        # 预览模式：applied 回写为 False，diff 保留
        return {
            'id': req_id,
            'ok': True,
            'result': {
                'diff': result['diff'],
                'applied': False,
                'xml_changed': result['xml_changed'],
            },
        }
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


def _handle_extract_all(req_id: str, req: dict) -> dict:
    """
    P4 语义字段抽取：扫描整个 docx，返回所有识别到的语义字段规则和证据。

    业务逻辑：
    1. ENOENT 检查先于 pipeline，避免 docx 解析阶段抛非预期异常
    2. 调用 pipeline.extract_all(file)，返回 {rules, evidence}
    3. PackageNotFoundError → PARSE_ERROR，与其它命令保持一致
    """
    file = req['file']
    if not Path(file).exists():
        return {'id': req_id, 'ok': False, 'error': f'file not found: {file}', 'code': 'ENOENT'}
    try:
        from .extractor.pipeline import extract_all
        result = extract_all(file)
        return {'id': req_id, 'ok': True, 'result': result}
    except PackageNotFoundError:
        return {'id': req_id, 'ok': False, 'error': f'docx malformed: {file}', 'code': 'PARSE_ERROR'}


def _handle_extract_from_selection(req_id: str, req: dict) -> dict:
    """
    P4 基于用户选中段落的精确字段抽取。

    业务逻辑：
    1. 接受 para_indices（用户在预览中点选的段落索引列表）和 field_id（目标字段）
    2. ENOENT 检查先于 pipeline
    3. 调用 pipeline.extract_from_selection(file, para_indices, field_id)
    """
    file = req['file']
    para_indices = req['para_indices']
    field_id = req['field_id']
    # Task 4 新增：前端按 Shift 选句时携带 selected_text，此处透传给 pipeline
    # 未传则 None，pipeline 内部回退为全段提取
    selected_text: str | None = req.get('selected_text')
    if not Path(file).exists():
        return {'id': req_id, 'ok': False, 'error': f'file not found: {file}', 'code': 'ENOENT'}
    try:
        from .extractor.pipeline import extract_from_selection
        result = extract_from_selection(file, para_indices, field_id, selected_text)
        return {'id': req_id, 'ok': True, 'result': result}
    except PackageNotFoundError:
        return {'id': req_id, 'ok': False, 'error': f'docx malformed: {file}', 'code': 'PARSE_ERROR'}


def _handle_list_fields(req_id: str) -> dict:
    from .engine_v2.field_defs import FIELD_DEFS
    return {'id': req_id, 'ok': True, 'result': {'fields': FIELD_DEFS}}
