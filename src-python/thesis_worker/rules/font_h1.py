"""
@file: font_h1.py
@description: 检测和修复一级标题段落（Heading 1）的字体名称、字号和加粗
              config 示例：{"family": "黑体", "size_pt": 16, "bold": true}
              bold 字段可选，None 表示不检查加粗；Spec 内置模板 bold=true
@author: Atlas.oi
@date: 2026-04-18
"""
from typing import Any, Optional
from docx.document import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

from ..models import Issue, Location, FixResult


# 蓝色标记：Office 标准"蓝色, 个性色 1" = #0070C0
_MARK_COLOR = RGBColor(0x00, 0x70, 0xC0)

# context 预览长度：段落前 N 字符，供用户 Ctrl-F 定位段落
_CONTEXT_MAX = 30

# snippet 截断长度：取 run 文本前 N 字符展示给用户
_SNIPPET_MAX = 20

# 一级标题 style 名称集合：兼容两种 python-docx 命名约定
_H1_STYLES = {'Heading 1', 'Heading1'}


def _make_context(text: str) -> str:
    """段落预览：首尾 strip 后取前 N 字符"""
    stripped = text.strip()
    if len(stripped) <= _CONTEXT_MAX:
        return stripped
    return stripped[:_CONTEXT_MAX] + '…'


def _is_h1_paragraph(para) -> bool:
    """只匹配一级标题段落，兼容 'Heading 1' 和 'Heading1' 两种命名"""
    style_name = para.style.name if para.style else ''
    return style_name in _H1_STYLES


def _read_actual_font_name(run) -> Optional[str]:
    """读取 run 实际字体名称
    run.font.name 为 None 时表示继承 style，
    需要通过 XML 层 rFonts 元素读取（eastAsia 优先，次取 ascii）"""
    if run.font.name:
        return run.font.name
    rpr = run._element.rPr
    if rpr is not None and rpr.rFonts is not None:
        # 中文字符集优先读 eastAsia
        ea = rpr.rFonts.get(qn('w:eastAsia'))
        if ea:
            return ea
        # 回退读 ascii（英文字符集）
        ascii_name = rpr.rFonts.get(qn('w:ascii'))
        if ascii_name:
            return ascii_name
    # None = 真正继承 doc/style 默认字体
    return None


def _read_actual_size_pt(run) -> Optional[int]:
    """读取 run 实际字号（磅）；为 None 表示继承 style 默认值"""
    if run.font.size is None:
        return None
    return int(run.font.size.pt)


class FontH1Rule:
    id = 'font.h1'
    category = 'format'
    severity = 'warning'
    fix_available = True

    @staticmethod
    def detect(doc: Document, value: Any) -> list[Issue]:
        """
        遍历所有一级标题段落（Heading 1），逐 run 检查字体名、字号和加粗。

        业务逻辑：
        1. 从 value 读取期望字体/字号/加粗设置
        2. value 不完整（缺 family 或 size_pt）时静默跳过，不误报
        3. 遍历段落，只处理 Heading 1 段落
        4. 遍历 run，跳过纯空白 run
        5. 对每个 run 比较字体/字号/(可选)加粗，不符则生成 Issue
        """
        expected_family: Optional[str] = value.get('family')
        expected_size: Optional[int] = value.get('size_pt')
        expected_bold: Optional[bool] = value.get('bold')  # None 表示不检查；Spec 模板通常为 True

        # value 不完整时静默跳过，避免模板缺字段时产生大量误报
        if not expected_family or not expected_size:
            return []

        issues: list[Issue] = []
        for p_idx, para in enumerate(doc.paragraphs):
            if not _is_h1_paragraph(para):
                continue
            context = _make_context(para.text)
            for r_idx, run in enumerate(para.runs):
                if not run.text.strip():
                    # 空白 run 不检查——它对用户不可见，字体无意义
                    continue
                actual_family = _read_actual_font_name(run)
                actual_size = _read_actual_size_pt(run)

                problems: list[str] = []
                if actual_family != expected_family:
                    problems.append(f'family={actual_family!r}→{expected_family!r}')
                if actual_size != expected_size:
                    problems.append(f'size_pt={actual_size}→{expected_size}')
                if expected_bold is not None:
                    actual_bold = bool(run.font.bold)
                    if actual_bold != expected_bold:
                        problems.append(f'bold={actual_bold}→{expected_bold}')

                if not problems:
                    continue

                snippet = run.text[:_SNIPPET_MAX]
                if len(run.text) > _SNIPPET_MAX:
                    snippet += '…'

                issues.append(Issue(
                    rule_id='font.h1',
                    loc=Location(para=p_idx, run=r_idx),
                    message='一级标题字体/字号不符：' + '; '.join(problems),
                    current={'family': actual_family, 'size_pt': actual_size},
                    expected={'family': expected_family, 'size_pt': expected_size},
                    fix_available=True,
                    snippet=snippet,
                    context=context,
                ))
        return issues

    @staticmethod
    def fix(doc: Document, issue: Issue, value: Any) -> FixResult:
        """
        修复指定 run 的字体/字号/(可选)加粗，并用蓝色标记提示用户。

        业务逻辑：
        1. 按 loc.para / loc.run 定位目标 run
        2. 记录修复前的状态供 diff 展示
        3. 写入期望字体名、字号、(可选)加粗
        4. 将文字颜色设为蓝色 #0070C0，提示人工复核
        """
        para = doc.paragraphs[issue.loc.para]
        run = para.runs[issue.loc.run]

        # 记录修复前摘要，供 diff 展示
        before_family = run.font.name or '(继承)'
        before_size = f'{int(run.font.size.pt)}pt' if run.font.size else '(继承)'
        before_summary = f'{before_family} {before_size}'

        # 应用字体/字号修复
        run.font.name = value['family']
        run.font.size = Pt(value['size_pt'])

        # 可选：同步修复加粗（Spec 模板 bold=true，标题应加粗）
        if value.get('bold') is not None:
            run.font.bold = value['bold']

        # 蓝色标记，提示用户已修改（与 font.body 保持一致的标记风格）
        run.font.color.rgb = _MARK_COLOR

        after_summary = f'{value["family"]} {value["size_pt"]}pt'
        return FixResult(
            diff=f'- {before_summary}\n+ {after_summary}',
            applied=True,
            xml_changed=[f'w:p[{issue.loc.para}]/w:r[{issue.loc.run}]'],
        )

    @staticmethod
    def extract(doc: Document) -> dict | None:
        """
        从 docx 反推一级标题的字体、字号和加粗设置。

        策略：扫所有 Heading 1 段落的所有 run，
        统计 majority 字体名、字号和加粗，返回置信度 = 字体多数票占比。
        """
        from collections import Counter
        fc: Counter = Counter()
        sc: Counter = Counter()
        bc: Counter = Counter()
        for para in doc.paragraphs:
            if not _is_h1_paragraph(para):
                continue
            for run in para.runs:
                if not run.text.strip():
                    continue
                fname = _read_actual_font_name(run)
                fsize = _read_actual_size_pt(run)
                if fname:
                    fc[fname] += 1
                if fsize:
                    sc[fsize] += 1
                # bold 为 None 表示继承，只统计明确设置的值
                if run.font.bold is not None:
                    bc[bool(run.font.bold)] += 1
        if not fc or not sc:
            # 文档无一级标题内容，无法提取
            return None
        family = fc.most_common(1)[0][0]
        size = sc.most_common(1)[0][0]
        bold = bc.most_common(1)[0][0] if bc else False
        total = sum(fc.values())
        return {
            'value': {'family': family, 'size_pt': size, 'bold': bold},
            'source_xml': f'h1: <w:rFonts w:eastAsia="{family}"/>',
            'confidence': round(fc[family] / total, 2) if total else 0.0,
        }
