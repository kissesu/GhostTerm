"""
@file: pipeline.py
@description: extract_all / extract_from_selection 主流程
              读段落文本+样式 → 按 field_matcher 关联 → 用 patterns+gazetteer 抽属性
@author: Atlas.oi
@date: 2026-04-18
"""
import os
import re  # 用于空格占位 fallback 的正则匹配
import logging
from typing import Any, Optional
from docx import Document

from .gazetteer import find_font, find_align, is_bold_keyword
from .patterns import extract_size_name, extract_size_pt_raw
from .field_matcher import match_all_fields
from ..utils.size import name_to_pt

# 模块级 logger，遵循 Python 最佳实践：用包层级名称，不新建 handler
logger = logging.getLogger(__name__)

# 开发者诊断日志中 context_snippet 的截断长度（字符数）。
# 30 字足以定位段落来源，同时避免日志行过长影响可读性。
_LOG_SNIPPET_LEN = 30

# 所有字段 applicable_attributes 的并集，导入时一次性构建并冻结为 frozenset。
# 任何不在此集合内的 attr key 即为 unsupported，需写开发者诊断日志。
# field_defs 不反向依赖 pipeline，无循环导入风险；此处直接调用无副作用。
def _build_known_attr_keys() -> frozenset[str]:
    """从 field_defs 汇总所有已知 attr key 并集。
    此函数仅在模块首次被导入时调用一次，结果缓存在模块级常量中。
    """
    from ..engine_v2.field_defs import FIELD_DEFS
    keys: set[str] = set()
    for field in FIELD_DEFS:
        keys.update(field.get('applicable_attributes', []))
    return frozenset(keys)

# 模块级常量：已知 attr key 全集（导入时计算一次，后续 O(1) 查询）
_KNOWN_ATTR_KEYS: frozenset[str] = _build_known_attr_keys()

# OOXML namespace，用于直接读 run 的 rFonts XML 属性
_W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
_W_RFONTS = f'{{{_W_NS}}}rFonts'
_W_EAST_ASIA = f'{{{_W_NS}}}eastAsia'
_W_ASCII = f'{{{_W_NS}}}ascii'
# run 级字间距（rPr/w:spacing），单位 twips
_W_SPACING = f'{{{_W_NS}}}spacing'
# w:spacing/@w:val 属性名
_W_VAL = f'{{{_W_NS}}}val'


def _extract_attributes_from_text(text: str) -> dict[str, Any]:
    """从单段文本里抽取所有可识别的格式属性

    业务逻辑：
    1. 字号：先尝试 pt 数字（如 12pt/15磅），再尝试字号名（如"小三号"）
    2. 字体：从词典中匹配 CJK 或 ASCII 字体名
    3. 加粗：检测"加粗"/"粗体"关键词
    4. 对齐：检测"居中"/"左对齐"等对齐词
    """
    attrs: dict[str, Any] = {}

    # 字号（字号名优先，fallback pt 数字）
    # 调换原优先级原因：T3.5 集成段前/段后/字符间距等新抽取后，单段可能同时出现
    # "小三号" + "段前 6 磅"——"6 磅"会被 extract_size_pt_raw 误识别为字号 6pt。
    # 中文字号名（小三/小四/六号）语义无歧义，应优先；pt 数字仅作纯英文规范的兜底。
    size_name = extract_size_name(text)
    pt: Optional[float] = None
    if size_name is not None:
        pt = name_to_pt(size_name)
    if pt is None:
        pt = extract_size_pt_raw(text)
    if pt is not None:
        attrs['font.size_pt'] = pt

    # 字体（cjk / ascii 分类）
    font_info = find_font(text)
    if font_info is not None:
        kind, name = font_info
        if kind == 'cjk':
            attrs['font.cjk'] = name
        else:
            attrs['font.ascii'] = name

    # 加粗
    if is_bold_keyword(text):
        attrs['font.bold'] = True

    # 对齐
    align = find_align(text)
    if align is not None:
        attrs['para.align'] = align

    # ============================================
    # T3.5: 集成 5 个新自然语言抽取函数
    # 段前/段后 / 缩进 / 行距 / 字符间距 / 表线
    # 顺序无依赖；后写入会覆盖同 key（当前各函数 attr key 不冲突）
    # ============================================
    from .patterns import (
        extract_indent,
        extract_line_spacing,
        extract_letter_spacing,
        extract_table_borders_text,
        _RE_PARA_SPACING,
    )
    from .units import length_to_pt

    # 段前/段后：同段可能两个 prefix 同时出现（"段前 6 磅，段后 3 磅"），
    # 用 finditer 扫描全部匹配，避免单匹配版本丢失第二个值
    for m in _RE_PARA_SPACING.finditer(text):
        prefix = m.group(1)
        val = float(m.group(2))
        unit = m.group(3)
        suffix = 'before' if prefix == '段前' else 'after'
        if unit == '行':
            attrs[f'para.space_{suffix}_lines'] = val
        else:
            pt = length_to_pt(val, unit)
            if pt is not None:
                attrs[f'para.space_{suffix}_pt'] = round(pt, 2)

    # 缩进（首行 / 悬挂，单匹配即可：同段同时出现首行+悬挂概率极低）
    if (ind := extract_indent(text)) is not None:
        sink_key, val = ind
        attrs[sink_key] = val

    # 行距类型 + 值（dict 形式，含两个绑定 key）
    if (ls := extract_line_spacing(text)) is not None:
        attrs.update(ls)

    # 字符间距（单匹配，同段同时出现两种单位概率极低）
    if (letter := extract_letter_spacing(text)) is not None:
        sink_key, val = letter
        attrs[sink_key] = val

    # 表线规范关键词（三线表 / 上下表线 / 表头下线）
    table_attrs = extract_table_borders_text(text)
    attrs.update(table_attrs)

    return attrs


def _read_paragraph_style_attrs(para) -> dict[str, Any]:
    """从段落 XML 样式里抽取属性

    业务逻辑：
    1. 遍历 runs，跳过空白 run，取第一个有实际文字的 run 的字体/字号/加粗
    2. 通过 OOXML rFonts 读 eastAsia（CJK 字体）属性
    3. 读段落对齐属性
    4. 读首行缩进（换算为字数，用于判断正文段落标准缩进）
    """
    attrs: dict[str, Any] = {}

    # 从第一个非空 run 读字体/字号/加粗
    for run in para.runs:
        if not run.text.strip():
            # 跳过空白 run，避免把段尾换行占位 run 的属性当有效属性
            continue
        if run.font.size is not None:
            attrs['font.size_pt'] = float(run.font.size.pt)
        if run.font.bold is True:
            attrs['font.bold'] = True
        # 直接读 OOXML XML 层的 rFonts 属性（python-docx 高级 API 不暴露 eastAsia）
        rpr = run._element.rPr
        if rpr is not None:
            rfonts = rpr.find(_W_RFONTS)
            if rfonts is not None:
                ea = rfonts.get(_W_EAST_ASIA)
                if ea:
                    attrs['font.cjk'] = ea
                asc = rfonts.get(_W_ASCII)
                if asc and 'font.cjk' not in attrs:
                    # 只有在没读到 CJK 字体时才保存 ascii 字体
                    attrs['font.ascii'] = asc
            # 字间距（OOXML rPr/w:spacing，单位 twips）
            # 240 twips = 1 字宽（@ 12pt 正文基准），换算为字数方便规范校验
            spacing_el = rpr.find(_W_SPACING)
            if spacing_el is not None:
                val = spacing_el.get(_W_VAL)
                if val:
                    try:
                        attrs['para.letter_spacing_chars'] = round(int(val) / 240, 1)
                    except ValueError:
                        pass  # 非法 val 值跳过（不影响其他属性）
        # 只读第一个非空 run 的属性：规范模板同段 run 格式通常一致，取首 run 已足够。
        # 若未来发现模板同段多 run 格式差异大，再改为 setdefault 遍历策略。
        break

    # 段落对齐
    if para.paragraph_format.alignment is not None:
        # WD_ALIGN_PARAGRAPH 枚举值：LEFT=0, CENTER=1, RIGHT=2, JUSTIFY=3
        align_map = {0: 'left', 1: 'center', 2: 'right', 3: 'justify'}
        val = para.paragraph_format.alignment
        if val in align_map:
            attrs['para.align'] = align_map[val]

    # 首行缩进（粗略换算为字数，以 12pt 正文为基准）
    fli = para.paragraph_format.first_line_indent
    if fli is not None:
        attrs['para.first_line_indent_chars'] = round(fli.pt / 12)

    # 行距：paragraph_format.line_spacing 返回 float（倍数）或 Emu（固定值）
    # 此处仅处理倍数模式（如 1.5 倍），固定值用 Emu 表示，暂不换算
    ls = para.paragraph_format.line_spacing
    if ls is not None and isinstance(ls, (int, float)):
        attrs['para.line_spacing'] = round(float(ls), 2)

    # 段前行数：space_before 以 Emu 返回，除以 12pt 换算为行数
    # 12pt = 1 行基准（与首行缩进换算统一基准）
    sb = para.paragraph_format.space_before
    if sb is not None:
        attrs['para.space_before_lines'] = round(sb.pt / 12, 1)
        # T3.1: 同时写入 pt 版本（直接保留磅值，不除以行基准）。
        # 规范文本用"磅"描述间距时用 _pt key，用"行"描述时用 _lines key，
        # 前端按字段 applicable_attributes 决定显示哪个。
        attrs['para.space_before_pt'] = round(sb.pt, 1)

    # 段后行数
    sa = para.paragraph_format.space_after
    if sa is not None:
        attrs['para.space_after_lines'] = round(sa.pt / 12, 1)
        # T3.1: 同时写入 pt 版本（理由同 space_before_pt）
        attrs['para.space_after_pt'] = round(sa.pt, 1)

    # 若 XML 未设 w:spacing，尝试识别空格占位字间距风格（如"摘  要"、"目　录"）
    # 同时接受半角空格（\s）和全角空格（U+3000），真实 Word 模板常用全角空格做字间距占位
    # 仅匹配 {单个非空字符}+{连续空格}+{单个非空字符} 的短标题模式，避免误匹配正文句子
    if 'para.letter_spacing_chars' not in attrs:
        stripped = para.text.strip()
        m = re.match(r'^(\S)([\s\u3000]+)(\S)$', stripped)
        if m:
            attrs['para.letter_spacing_chars'] = len(m.group(2))

    return attrs


def _read_run_list_style_attrs(run_list: list) -> dict[str, Any]:
    """从指定 run 列表中读取字体/加粗/字间距属性（不含段落级 align/indent/spacing）

    业务逻辑：
    - 与 _read_paragraph_style_attrs 的 run 部分逻辑一致，但接收外部 run 子集
    - 段落级属性（对齐、行距、首行缩进、段前/段后）属于整段，不跟 run 走，
      因此本函数只读 run 级属性，调用方负责补全段落级属性
    - 取 run_list 中第一个非空文字 run 的属性
    """
    attrs: dict[str, Any] = {}
    for run in run_list:
        if not run.text.strip():
            # 跳过空白 run，避免取到占位 run 的默认属性
            continue
        if run.font.size is not None:
            attrs['font.size_pt'] = float(run.font.size.pt)
        if run.font.bold is True:
            attrs['font.bold'] = True
        rpr = run._element.rPr
        if rpr is not None:
            rfonts = rpr.find(_W_RFONTS)
            if rfonts is not None:
                ea = rfonts.get(_W_EAST_ASIA)
                if ea:
                    attrs['font.cjk'] = ea
                asc = rfonts.get(_W_ASCII)
                if asc and 'font.cjk' not in attrs:
                    attrs['font.ascii'] = asc
            spacing_el = rpr.find(_W_SPACING)
            if spacing_el is not None:
                val = spacing_el.get(_W_VAL)
                if val:
                    try:
                        attrs['para.letter_spacing_chars'] = round(int(val) / 240, 1)
                    except ValueError:
                        pass
        # 只取第一个非空 run（与 _read_paragraph_style_attrs 策略一致）
        break
    return attrs


def _find_runs_for_text(para, selected_text: str) -> list:
    """在段落的 runs 中定位覆盖 selected_text 的 run 子集

    业务逻辑：
    1. 遍历 para.runs，逐步累加字符偏移，找到 selected_text 在 para.text 中的
       字符范围 [start, start+len)
    2. 返回与该范围有重叠的 run 列表（部分重叠也算覆盖）
    3. 若 selected_text 在 para.text 中找不到，返回空列表（调用方会回退到全段提取）

    @param para - python-docx Paragraph 对象
    @param selected_text - 用户选中的文本字符串
    @returns 覆盖选中文本的 run 列表，可能为空（未找到时回退）
    """
    full_text = para.text
    start = full_text.find(selected_text)
    if start == -1:
        # 找不到目标文本：无法定位 run，返回空列表让上层回退到全段
        return []

    end = start + len(selected_text)
    matched_runs: list = []
    cursor = 0

    for run in para.runs:
        run_len = len(run.text)
        run_start = cursor
        run_end = cursor + run_len
        cursor = run_end

        # run 与 [start, end) 有重叠则纳入子集
        if run_start < end and run_end > start:
            matched_runs.append(run)

    return matched_runs


def _merge_attrs(from_text: dict[str, Any], from_style: dict[str, Any]) -> dict[str, Any]:
    """合并文本抽取和样式抽取的属性

    文本抽取（规范说明括号内容）优先于 XML 样式读取。
    两者取并集，文本结果覆盖样式结果中的同名键。
    """
    merged = {**from_style, **from_text}
    return merged


def _log_and_filter_unsupported(
    attrs: dict[str, Any],
    spec_file: str,
    field_id: str,
    context_snippet: str,
) -> dict[str, Any]:
    """检测并过滤 attrs 中不属于当前字段白名单的 attr key。

    W6 修复：改为字段级白名单过滤（替代原全局 _KNOWN_ATTR_KEYS 并集过滤）。
    原因：section/table/numbering 全文注入时，结果可能带有该字段未声明的 attr key
    （例如 page_margin 字段注入了 table.* attr），违反 P4 白名单合约。

    业务逻辑：
    1. 从 field_defs 取该字段的 applicable_attributes 白名单
    2. 对比每个 key：
       a. 在白名单内 → 保留（写入 clean）
       b. 不在白名单但在 _KNOWN_ATTR_KEYS（已知但越界）→ 写 attr_field_mismatch 诊断日志
       c. 不在全局已知集（未知 attr）→ 写 unsupported_attr 诊断日志（原有逻辑）
    3. 两种日志均为开发者诊断，不暴露给用户面板

    注意：field_id 未知（如 '' 或 None）时 applicable_attrs 返回空列表，
    所有 key 会走 b/c 分支被过滤——这是正确行为（未知字段不应有任何属性）。

    @param attrs          - 合并后的属性字典（可能含越界或未知 key）
    @param spec_file      - 来源规范文件名（用于日志定位）
    @param field_id       - 当前字段 id（用于字段级白名单查询和日志定位）
    @param context_snippet - 来源段落文本前 30 字（用于日志定位）
    @returns 过滤掉白名单外 key 后的干净属性字典
    """
    from ..engine_v2.field_defs import applicable_attrs as _field_whitelist
    # 取该字段声明的 attr key 白名单（未知字段返回空列表）
    field_whitelist: frozenset[str] = frozenset(_field_whitelist(field_id))
    clean: dict[str, Any] = {}
    for key, val in attrs.items():
        if key in field_whitelist:
            # 在字段白名单内：保留
            clean[key] = val
        elif key in _KNOWN_ATTR_KEYS:
            # 全局已知 attr，但不属于该字段白名单：attr 越界，写专项诊断日志
            # 开发者可见，用于发现 section/table/numbering 全文注入越界问题
            logger.info(
                'attr_field_mismatch: key=%r is known but not in field=%r whitelist, will be dropped',
                key,
                field_id,
                extra={
                    'attr_key': key,
                    'spec_file': spec_file,
                    'context_snippet': context_snippet[:_LOG_SNIPPET_LEN],
                    'suspected_field_id': field_id,
                },
            )
        else:
            # 完全未知的 attr key（不在任何字段白名单中）：写原有 unsupported_attr 诊断日志
            logger.info(
                'unsupported_attr: key=%r not in schema, will be dropped',
                key,
                extra={
                    'attr_key': key,
                    'spec_file': spec_file,
                    'context_snippet': context_snippet[:_LOG_SNIPPET_LEN],
                    'suspected_field_id': field_id,
                },
            )
    return clean


def _calculate_confidence(attrs: dict[str, Any], text_len: int) -> float:
    """根据属性数量估算置信度（启发式）。
    阈值 0.0 / 0.5 / 0.7 / 0.9 对应抽到 0 / 1 / 2 / ≥3 个属性。
    数值为经验启发，非统计模型；Phase B 引入校验集后再替换为回归式估算。"""
    if len(attrs) == 0:
        return 0.0
    if len(attrs) >= 3:
        return 0.9
    if len(attrs) == 2:
        return 0.7
    # len == 1
    return 0.5


def _read_section_attrs(doc) -> dict[str, Any]:
    """读取 Section 级页面布局属性（装订线/页眉脚距/打印模式）

    业务逻辑：
    1. 从 doc.sections[0] 读取装订线宽、页眉距、页脚距（均取 cm 并保留 2 位小数）
    2. 通过 w:settings 根元素是否含 w:evenAndOddHeaders 判断打印模式
    3. 返回 4 个属性的字典，供 extract_all 合并到 page_margin 字段

    注意：若 sections 为空（异常文档），直接返回空字典，不抛异常。
    """
    from docx.oxml.ns import qn
    attrs: dict[str, Any] = {}
    if not doc.sections:
        return attrs
    section = doc.sections[0]

    # 装订线宽度（cm）
    # python-docx 对未显式设置的 section 属性返回 None 而非 0，
    # 真实 docx（尤其 textutil 从 .doc 转换）大量触发此情况，必须守卫
    if section.gutter is not None:
        attrs['page.margin_gutter_cm'] = round(section.gutter.cm, 2)

    # 页眉距页面顶端距离（cm）
    if section.header_distance is not None:
        attrs['page.header_offset_cm'] = round(section.header_distance.cm, 2)

    # 页脚距页面底端距离（cm）
    if section.footer_distance is not None:
        attrs['page.footer_offset_cm'] = round(section.footer_distance.cm, 2)

    # 打印模式：检测 w:evenAndOddHeaders 元素是否存在
    # 存在 → 奇偶页眉分设（双面打印） → 'double'；不存在 → 'single'
    settings_el = doc.settings.element
    even_odd = settings_el.find(qn('w:evenAndOddHeaders'))
    attrs['page.print_mode'] = 'double' if even_odd is not None else 'single'

    return attrs


def _read_first_row_bottom_border(tbl) -> 'float | None':
    """读第一行所有 tc 的 tcBorders/bottom 线宽，取众数（eighth-points → pt）。

    W5 修复：Word 三线表惯例是表头下边框写在第一行单元格的 tcBorders/bottom，
    而不是 tblBorders/insideH。本函数优先读取 tcBorders/bottom，
    作为 table.header_border_pt 的第一优先来源。

    @param tbl - python-docx Table._element（OOXML 元素）
    @returns 第一行 tc bottom border 众数（pt），无则返回 None（触发 insideH fallback）
    """
    from docx.oxml.ns import qn
    from collections import Counter
    rows = tbl.findall(qn('w:tr'))
    if not rows:
        return None
    first_row = rows[0]
    cells = first_row.findall(qn('w:tc'))
    if not cells:
        return None
    sizes: list[float] = []
    for tc in cells:
        tc_pr = tc.find(qn('w:tcPr'))
        if tc_pr is None:
            continue
        tc_borders = tc_pr.find(qn('w:tcBorders'))
        if tc_borders is None:
            continue
        bottom = tc_borders.find(qn('w:bottom'))
        if bottom is None:
            continue
        sz_val = bottom.get(qn('w:sz'))
        if sz_val:
            try:
                # eighth-points → pt
                sizes.append(int(sz_val) / 8.0)
            except ValueError:
                continue
    if not sizes:
        return None
    # 取众数：三线表同行所有单元格通常一致，极端情况取多数值
    return Counter(sizes).most_common(1)[0][0]


def _read_table_attrs(doc) -> dict[str, Any]:
    """读取文档第一个表格的 OOXML tblBorders 线宽属性，返回 table.* attr 字典。

    业务逻辑：
    1. 若 doc.tables 为空，直接返回空字典（不报错，文档无表格是合理情况）
    2. 取第一个表格，查找 tblPr/tblBorders 子元素
    3. 读 top/bottom/insideH/insideV 各方向的 w:sz（eighth-points → pt 除以 8）
    4. W4 修复：三线表判定加 left/right == 0（排除带左右外框的误判）
    5. W5 修复：table.header_border_pt 优先读第一行 tcBorders/bottom，fallback insideH
    6. 返回 4 个 attr key；table.is_three_line 固定写入（即使为 False 也写，便于 UI 展示）

    注意：OOXML w:sz 单位为 eighth-points（1/8 pt），除以 8 得磅值。
    """
    from docx.oxml.ns import qn
    attrs: dict[str, Any] = {}
    if not doc.tables:
        return attrs
    tbl = doc.tables[0]
    tbl_pr = tbl._element.find(qn('w:tblPr'))
    if tbl_pr is None:
        return attrs
    tbl_borders = tbl_pr.find(qn('w:tblBorders'))
    if tbl_borders is None:
        return attrs

    # 从 tblBorders 子元素中提取各方向线宽（eighth-points → pt）
    border_pt: dict[str, float] = {}
    for child in tbl_borders:
        local = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        sz_val = child.get(qn('w:sz'))
        if sz_val is not None:
            try:
                border_pt[local] = int(sz_val) / 8.0
            except ValueError:
                pass  # sz 非法值跳过

    # tblBorders 存在时，无条件写入 4 个 attr 全集（缺方向填 0.0）。
    # 保持与 checker 侧 borders.get('top', 0.0) 的行为对等，
    # 确保 UI 始终展示完整的 4 个 attr 行，不出现数量不固定的问题。
    top_pt = border_pt.get('top', 0.0)
    bottom_pt = border_pt.get('bottom', 0.0)
    inside_v_pt = border_pt.get('insideV', 0.0)
    # W4 修复：三线表额外检查左/右外侧竖线（left/right 也必须为 0）
    # 带左右边框但无 insideV 的表应判 False，否则会误判为三线表
    left_pt = border_pt.get('left', 0.0)
    right_pt = border_pt.get('right', 0.0)
    attrs['table.is_three_line'] = (
        top_pt > 0 and bottom_pt > 0
        and inside_v_pt == 0.0 and left_pt == 0.0 and right_pt == 0.0
    )
    # 上下边框线宽（缺方向时填 0.0，与 checker 侧 fallback 对齐）
    attrs['table.border_top_pt'] = round(top_pt, 4)
    attrs['table.border_bottom_pt'] = round(bottom_pt, 4)
    # W5 修复：表头下边框优先读第一行 tcBorders/bottom，无则 fallback 到 insideH
    # 原因：Word 三线表惯例是表头下线写在 tcBorders/bottom，而非 tblBorders/insideH
    tc_bottom = _read_first_row_bottom_border(tbl._element)
    header_pt = tc_bottom if tc_bottom is not None else border_pt.get('insideH', 0.0)
    attrs['table.header_border_pt'] = round(header_pt, 4)

    return attrs


# ───────────────────────────────────────────────────────────────────
# T3.3: 编号风格推断正则（模块级常量，避免每次调用重复编译）
# 全部含 [\s　] 以兼容全角空格（U+3000），参 feedback_regex_fullwidth_space.md
# ───────────────────────────────────────────────────────────────────

# 连续图编号：图1 / 图 2 / 图3，后面不跟 . 或 -（排除章节式前缀）
_RE_FIG_CONTINUOUS = re.compile(r'图[\s　]*(\d+)(?![\.\-\d])', re.UNICODE)
# 章节式图编号：图1-1 / 图1.2 / 图 2-3
_RE_FIG_CHAPTER = re.compile(r'图[\s　]*(\d+)[\.\-](\d+)', re.UNICODE)

# W2 修复：图题开头锚定正则，只匹配段落行首的图题（排除"如图1所示"/"见图2"等正文引用）
# 必须以 "图N" 模式开头才算图题计票，不匹配正文中间的 "图N" 引用
_RE_FIG_CAPTION_START = re.compile(r'^图[\s　]*\d+', re.UNICODE)

# 分图字母标记：图1(a) / 图2（b）/ 图1.2(c)
_RE_SUBFIG_LETTER = re.compile(
    r'图[\s　]*\d+[\.\-]?\d*[\s　]*[\(（][a-zA-Z][\)）]', re.UNICODE
)
# W3 修复：分图数字标记：图1(1) / 图2（2）/ 图1.2(3)
_RE_SUBFIG_NUMBER = re.compile(
    r'图[\s　]*\d+[\.\-]?\d*[\s　]*[\(（]\d+[\)）]', re.UNICODE
)
# 连续公式编号：(1) / （2）
_RE_FORMULA_CONTINUOUS = re.compile(r'[\(（][\s　]*(\d+)[\s　]*[\)）]')
# 章节式公式编号：(1-1) / (2.3) / （1-2）
_RE_FORMULA_CHAPTER = re.compile(r'[\(（][\s　]*(\d+)[\.\-](\d+)[\s　]*[\)）]')

# W1 修复：含数学符号的正则，用于判断段落是否疑似公式
# 含等号或常见数学运算符，表明该段落更可能是公式，而非引用/参考文献中的括号
_RE_MATH_SYMBOLS = re.compile(r'[+\-×÷±∑∫√≤≥≠]', re.UNICODE)
# 短行以编号结束（如 "... (1-1)" "... (2)"），判别力强于仅匹配括号
_RE_FORMULA_TAIL = re.compile(r'[\(（](?:\d+[\.\-]?\d*)[\)）]\s*$', re.UNICODE)


def _looks_like_formula_paragraph(text: str) -> bool:
    """判断段落是否疑似公式段落，用于 W1 过滤非公式括号计票。

    只有疑似公式段落才参与 formula_continuous/chapter 计票，
    避免 APA 年份引用 (2020)/(2021) 等误触发公式编号推断。

    判定条件（任一满足即为公式段落）：
    1. 含等号 '=' 或常见数学运算符（+/×/∑/∫ 等）
    2. 短行（<=30 字符）且以编号括号结束（如 "某公式 (1)" "E=mc² (1-1)"）

    空段落或超长段落（>200 字符）直接排除，避免正文句子误匹配。

    @param text - 段落文本（已 strip 之前的原始文本）
    @returns True 表示疑似公式段落，False 表示跳过不计票
    """
    stripped = text.strip()
    # 空段落或超长段落（>200字）排除：公式段落通常极短
    if not stripped or len(stripped) > 200:
        return False
    # 条件1：含等号（公式核心特征）或常见数学运算符
    if '=' in stripped or _RE_MATH_SYMBOLS.search(stripped):
        return True
    # 条件2：短行且以括号编号结束（编号标注在行末的纯编号行，如 "(1-1)"）
    if len(stripped) <= 30 and _RE_FORMULA_TAIL.search(stripped):
        return True
    return False


def _read_numbering_styles(doc) -> 'dict[str, str]':
    """从文档图题和公式段落推断编号风格（启发式多数票算法）。

    业务逻辑：
    1. 遍历所有段落，用 _RE_FIG_* 正则匹配图题文本
       - W2 修复：仅限段落开头以"图N"开始的行（排除"如图1所示"等正文引用）
       - 章节式（图1-1/图2.3）vs 连续式（图1/图2）各自计票
       - >=2 个样本时按多数票决定 figure_style；冲突时偏保守选 chapter_based
    2. 仅在 figure_style 已确定时推断 subfigure_style：
       - W3 修复：同时识别字母子图 (a)/(b) 和数字子图 (1)/(2)，多数票决定类型
    3. 用 _RE_FORMULA_* 正则匹配含编号括号的段落（可能在行末）
       - W1 修复：仅对疑似公式段落（含数学符号/等号/短行编号尾）计票
         排除 APA 年份引用 (2020)/(2021) 等误触发

    @param doc - python-docx Document 对象
    @returns  { 'numbering.figure_style': str, 'numbering.subfigure_style': str,
                'numbering.formula_style': str }，仅写入有足够样本的 key
    """
    # 图编号计票
    fig_continuous = 0
    fig_chapter = 0
    # W3 修复：分图标记分别计票（字母 vs 数字）
    subfig_letter_count = 0
    subfig_number_count = 0

    # 公式编号计票
    formula_continuous = 0
    formula_chapter = 0

    for para in doc.paragraphs:
        text = para.text
        if not text.strip():
            continue

        # ── 图题检测 ────────────────────────────────────────────
        # W2 修复：只对段落开头以"图N"开始的行计票（排除正文"如图1所示"等引用）
        stripped = text.strip()
        if _RE_FIG_CAPTION_START.match(stripped):
            # 章节式优先判断（否则"图1-1"也会被连续式 regex 匹配前缀"图1"）
            if _RE_FIG_CHAPTER.search(text):
                fig_chapter += 1
                # 同段检测子图标记（W3：字母/数字分别计票）
                if _RE_SUBFIG_LETTER.search(text):
                    subfig_letter_count += 1
                if _RE_SUBFIG_NUMBER.search(text):
                    subfig_number_count += 1
            elif _RE_FIG_CONTINUOUS.search(text):
                fig_continuous += 1
                if _RE_SUBFIG_LETTER.search(text):
                    subfig_letter_count += 1
                if _RE_SUBFIG_NUMBER.search(text):
                    subfig_number_count += 1

        # ── 公式编号检测 ─────────────────────────────────────────
        # W1 修复：只对疑似公式段落（含数学符号/等号/短行编号尾）计票
        # 排除 APA 年份引用 (2020)/(2021) 等正文括号被误算为公式编号
        if _looks_like_formula_paragraph(text):
            # 章节式先判（避免"(1-1)"被连续式错误匹配括号内数字）
            if _RE_FORMULA_CHAPTER.search(text):
                formula_chapter += 1
            elif _RE_FORMULA_CONTINUOUS.search(text):
                formula_continuous += 1

    result: dict[str, str] = {}

    # figure_style 多数票：需要至少 2 个样本才做判断
    total_fig = fig_continuous + fig_chapter
    if total_fig >= 2:
        if fig_chapter > fig_continuous:
            # 章节式占多数，或相等时保守选章节式
            result['numbering.figure_style'] = 'chapter_based'
        elif fig_continuous > fig_chapter:
            result['numbering.figure_style'] = 'continuous'
        else:
            # 票数相等：保守选 chapter_based
            result['numbering.figure_style'] = 'chapter_based'

        # W3 修复：仅在 figure_style 已确定时推断 subfigure_style，多数票决定类型
        total_subfig = subfig_letter_count + subfig_number_count
        if total_subfig >= 2:
            if subfig_letter_count > subfig_number_count:
                result['numbering.subfigure_style'] = 'a_b_c'
            elif subfig_number_count > subfig_letter_count:
                result['numbering.subfigure_style'] = '1_2_3'
            else:
                # 票数相等：偏字母（a_b_c 更通用）
                result['numbering.subfigure_style'] = 'a_b_c'

    # formula_style 多数票：至少 2 个样本
    total_formula = formula_continuous + formula_chapter
    if total_formula >= 2:
        if formula_chapter > formula_continuous:
            result['numbering.formula_style'] = 'chapter_based'
        elif formula_continuous > formula_chapter:
            result['numbering.formula_style'] = 'continuous'
        else:
            # 票数相等：保守选 chapter_based
            result['numbering.formula_style'] = 'chapter_based'

    return result


def _detect_punct_space_after(doc) -> 'bool | None':
    """检测英文标点后是否规范地空一字符。

    业务逻辑：
    1. 拼接全文所有 run.text，得到整个文档的纯文本
    2. 统计 ASCII 标点 [.,;:!?] 后紧跟空白字符的次数（space_after_count）
    3. 统计 ASCII 标点后紧跟非空白字符的次数（no_space_count）
    4. 若 space_after_count 明显占优（>= 2x 无空格），判定为 True；
       若无空格明显占优（no_space >= 2x space_after），判定为 False；
       样本过少（两者加起来 < 3）时返回 None 表示"无法判定"
    阈值 2x 的依据：标点后空格有时夹杂引号/括号等上下文紧接情况，
    允许少量例外，但超过 2:1 才认定主导规范成立。
    """
    # 收集全文文本（逐 run 而非逐段，可避免段落拼接引入误差）
    fulltext_parts = []
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text:
                fulltext_parts.append(run.text)
    fulltext = ''.join(fulltext_parts)

    space_after = len(re.findall(r'[.,;:!?](?=\s)', fulltext))
    no_space = len(re.findall(r'[.,;:!?](?=\S)', fulltext))
    total = space_after + no_space
    if total < 3:
        # 样本不足，无法做可靠的统计推断
        return None
    return space_after >= 2 * no_space


def extract_all(file: str) -> dict[str, Any]:
    """全文自动抽取字段规则

    业务逻辑：
    1. 读取 docx 所有段落文本
    2. 调用 field_matcher 关联字段 id
    3. 对每个命中段落，合并文本抽取 + 样式抽取的属性
    4. 同一字段取首次命中（跳过重复）
    5. 返回 rules 字典 + evidence 列表 + unmatched_paragraphs 列表

    @param file - docx 文件路径
    @returns {rules: dict[field_id, {enabled, value}], evidence: list, unmatched_paragraphs: list}
    """
    doc = Document(file)
    paragraphs_text = [p.text for p in doc.paragraphs]

    field_matches = match_all_fields(paragraphs_text)

    rules: dict[str, Any] = {}
    evidence: list[dict[str, Any]] = []
    unmatched: list[dict[str, Any]] = []

    for para_idx, field_id, confidence in field_matches:
        if field_id is None:
            text = paragraphs_text[para_idx].strip()
            if text:
                unmatched.append({
                    'idx': para_idx,
                    'text': text[:60],
                    'reason': 'no_field_keyword',
                })
            continue

        # 同一字段只取首次命中（模板文档中同字段极少出现两次，后出现的通常是示例）
        if field_id in rules:
            continue

        para = doc.paragraphs[para_idx]
        text_attrs = _extract_attributes_from_text(para.text)
        style_attrs = _read_paragraph_style_attrs(para)
        value = _merge_attrs(text_attrs, style_attrs)

        # 过滤掉 schema 未声明的 attr key，同时写开发者诊断日志
        value = _log_and_filter_unsupported(
            value,
            spec_file=os.path.basename(file),
            field_id=field_id,
            context_snippet=para.text,
        )

        final_conf = _calculate_confidence(value, len(para.text))

        rules[field_id] = {
            'enabled': True,
            'value': value,
        }
        evidence.append({
            'field_id': field_id,
            'source_para_idx': para_idx,
            'source_text': para.text[:100],
            'confidence': final_conf,
        })

    # ============================================
    # 全文扫描：英文标点后空格规范推断
    # 无法通过段落级 field_matcher 命中，需要独立扫全文一次
    # ============================================
    punct_result = _detect_punct_space_after(doc)
    if punct_result is not None:
        # setdefault 避免覆盖已由段落级流程填入的 mixed_script_global 条目
        rules.setdefault('mixed_script_global', {'enabled': True, 'value': {}})
        # 写入 value 子字典；该字段全文启发式推断，置信度约 0.7
        # （非段落直接匹配，0.7 与 _calculate_confidence 中"2个属性"档位对齐）
        rules['mixed_script_global']['value']['mixed_script.punct_space_after'] = punct_result

    # ============================================
    # T3.1: Section 级页面属性抽取（装订线/页眉脚距/打印模式）
    # 这四项属性无法通过段落扫描命中，需独立读取文档 Section
    # 结果合并到 page_margin 字段（按 plan "或并入 page_margin" 选并入）
    # ============================================
    section_attrs = _read_section_attrs(doc)
    if section_attrs:
        # 过滤掉 schema 未声明的 attr key（与段落流程保持一致）
        section_attrs = _log_and_filter_unsupported(
            section_attrs,
            spec_file=os.path.basename(file),
            field_id='page_margin',
            context_snippet='[section-level attrs]',
        )
        if section_attrs:
            rules.setdefault('page_margin', {'enabled': True, 'value': {}})
            rules['page_margin']['value'].update(section_attrs)

    # ============================================
    # T3.2: 表格级属性抽取（三线表判定 + 边框线宽）
    # 表格结构无法通过段落扫描命中，需独立读取文档 tables
    # 结果合并到 table_header 字段（线宽约束挂在表头字段，
    # 因为三线表中表头行承载了最关键的边框约束）
    # ============================================
    table_attrs = _read_table_attrs(doc)
    if table_attrs:
        # 过滤掉 schema 未声明的 attr key（与 section_attrs 流程保持一致）
        table_attrs = _log_and_filter_unsupported(
            table_attrs,
            spec_file=os.path.basename(file),
            field_id='table_header',
            context_snippet='[table-level attrs]',
        )
        if table_attrs:
            rules.setdefault('table_header', {'enabled': True, 'value': {}})
            rules['table_header']['value'].update(table_attrs)

    # ============================================
    # T3.3: 编号风格抽取（图/分图/公式，启发式全文扫描）
    # 编号风格无法通过段落级 field_matcher 命中，需独立扫全文推断
    # figure_style + subfigure_style → figure_caption 字段
    # formula_style → formula_block 字段
    # ============================================
    numbering_styles = _read_numbering_styles(doc)
    if numbering_styles:
        # figure_caption 相关属性（figure_style + subfigure_style）
        fig_attrs: dict[str, Any] = {}
        for key in ('numbering.figure_style', 'numbering.subfigure_style'):
            if key in numbering_styles:
                fig_attrs[key] = numbering_styles[key]
        if fig_attrs:
            fig_attrs = _log_and_filter_unsupported(
                fig_attrs,
                spec_file=os.path.basename(file),
                field_id='figure_caption',
                context_snippet='[numbering-figure attrs]',
            )
            if fig_attrs:
                rules.setdefault('figure_caption', {'enabled': True, 'value': {}})
                rules['figure_caption']['value'].update(fig_attrs)

        # formula_block 相关属性（formula_style）
        if 'numbering.formula_style' in numbering_styles:
            formula_attrs: dict[str, Any] = {
                'numbering.formula_style': numbering_styles['numbering.formula_style'],
            }
            formula_attrs = _log_and_filter_unsupported(
                formula_attrs,
                spec_file=os.path.basename(file),
                field_id='formula_block',
                context_snippet='[numbering-formula attrs]',
            )
            if formula_attrs:
                rules.setdefault('formula_block', {'enabled': True, 'value': {}})
                rules['formula_block']['value'].update(formula_attrs)

    return {
        'rules': rules,
        'evidence': evidence,
        'unmatched_paragraphs': unmatched,
    }


def extract_from_selection(
    file: str,
    para_indices: list[int],
    field_id: str,
    selected_text: str | None = None,
) -> dict[str, Any]:
    """从用户选定的段落抽取属性，赋给指定字段

    业务逻辑：
    1. 读取指定索引段落的文本和样式
    2. 若传入 selected_text：
       a. 在 para_indices[0] 的段落中定位 selected_text 覆盖的 run 子集
       b. 用 _read_run_list_style_attrs(matched_runs) 读 run 级属性
       c. 补全段落级属性（行距、对齐、缩进等属于整段，不跟 run 走）
       d. 若 selected_text 在段落中找不到，回退为全段提取（不报错，安全降级）
    3. 若未传入 selected_text：合并多段属性（后段覆盖前段同名键）
    4. 返回字段 id + 合并属性 + 置信度 + 证据信息

    @param file - docx 文件路径
    @param para_indices - 用户选定的段落索引列表（可多段）
    @param field_id - 用户指定的字段 id
    @param selected_text - 用户按句选取的文本字符串（可选）
    @returns {field_id, value, confidence, evidence}
    """
    doc = Document(file)
    all_paras = list(doc.paragraphs)

    # ============================================
    # selected_text 路径：缩小到单段特定 run 子集
    # ============================================
    if selected_text is not None and para_indices:
        first_idx = para_indices[0]
        if 0 <= first_idx < len(all_paras):
            para = all_paras[first_idx]
            matched_runs = _find_runs_for_text(para, selected_text)

            if matched_runs:
                # 仅对覆盖选中文本的 run 读字体/加粗等 run 级属性
                run_attrs = _read_run_list_style_attrs(matched_runs)
                # 段落级属性（行距/对齐/首行缩进）属于整段，无论选哪句都应读取
                para_level_attrs = _extract_para_level_attrs(para)
                style_attrs = {**para_level_attrs, **run_attrs}
                text_attrs = _extract_attributes_from_text(selected_text)
                value = _merge_attrs(text_attrs, style_attrs)
                # 过滤 schema 未声明的 attr key，同时写开发者诊断日志
                value = _log_and_filter_unsupported(
                    value,
                    spec_file=os.path.basename(file),
                    field_id=field_id,
                    context_snippet=selected_text,
                )
                confidence = _calculate_confidence(value, len(selected_text))
                return {
                    'field_id': field_id,
                    'value': value,
                    'confidence': confidence,
                    'evidence': {
                        'source_text': selected_text[:200],
                        'matched_patterns': list(value.keys()),
                    },
                }
            # 找不到 selected_text：回退到全段提取（不报错，走下方通用路径）

    # ============================================
    # 通用路径：多段合并提取（selected_text 未传 or 定位失败时）
    # ============================================
    text_parts: list[str] = []
    combined_style_attrs: dict[str, Any] = {}

    for idx in para_indices:
        if idx < 0 or idx >= len(all_paras):
            continue
        para = all_paras[idx]
        text_parts.append(para.text)
        style_attrs = _read_paragraph_style_attrs(para)
        # 后段属性覆盖前段，取到的属性集合更丰富
        combined_style_attrs.update(style_attrs)

    combined_text = '\n'.join(text_parts)
    text_attrs = _extract_attributes_from_text(combined_text)
    value = _merge_attrs(text_attrs, combined_style_attrs)
    # 过滤 schema 未声明的 attr key，同时写开发者诊断日志
    value = _log_and_filter_unsupported(
        value,
        spec_file=os.path.basename(file),
        field_id=field_id,
        context_snippet=combined_text,
    )

    confidence = _calculate_confidence(value, len(combined_text))

    return {
        'field_id': field_id,
        'value': value,
        'confidence': confidence,
        'evidence': {
            'source_text': combined_text[:200],
            'matched_patterns': list(value.keys()),
        },
    }


def _extract_para_level_attrs(para) -> dict[str, Any]:
    """从段落读取段落级格式属性（不含 run 级字体/字号/加粗）

    业务逻辑：
    - 段落级属性（对齐、行距、首行缩进、段前/段后）属于整段语义，
      即使用户只选了一句话，这些属性仍应从整段读取
    - 与 _read_paragraph_style_attrs 的后半段逻辑一致，抽为独立函数供
      selected_text 路径使用
    """
    attrs: dict[str, Any] = {}

    # 段落对齐
    if para.paragraph_format.alignment is not None:
        align_map = {0: 'left', 1: 'center', 2: 'right', 3: 'justify'}
        val = para.paragraph_format.alignment
        if val in align_map:
            attrs['para.align'] = align_map[val]

    # 首行缩进
    fli = para.paragraph_format.first_line_indent
    if fli is not None:
        attrs['para.first_line_indent_chars'] = round(fli.pt / 12)

    # 行距
    ls = para.paragraph_format.line_spacing
    if ls is not None and isinstance(ls, (int, float)):
        attrs['para.line_spacing'] = round(float(ls), 2)

    # 段前行数（T3.1: 同时写入 pt 版本，与 _read_paragraph_style_attrs 保持一致）
    sb = para.paragraph_format.space_before
    if sb is not None:
        attrs['para.space_before_lines'] = round(sb.pt / 12, 1)
        attrs['para.space_before_pt'] = round(sb.pt, 1)

    # 段后行数（T3.1: 同时写入 pt 版本）
    sa = para.paragraph_format.space_after
    if sa is not None:
        attrs['para.space_after_lines'] = round(sa.pt / 12, 1)
        attrs['para.space_after_pt'] = round(sa.pt, 1)

    # 空格占位字间距（字间距 fallback，与 _read_paragraph_style_attrs 一致）
    stripped = para.text.strip()
    m = re.match(r'^(\S)([\s\u3000]+)(\S)$', stripped)
    if m:
        attrs['para.letter_spacing_chars'] = len(m.group(2))

    return attrs
