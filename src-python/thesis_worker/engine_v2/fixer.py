"""
@file: fixer.py
@description: v2 fix 实现：按 issue.attr 修改段落属性，成功后蓝色标记，
              标记作用是提示用户该处已被自动修复，方便人工复核
@author: Atlas.oi
@date: 2026-04-18
"""
from typing import Any
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

# 修复标记蓝色：Microsoft Word 修订蓝，视觉辨识度高且不干扰文字可读性
_MARK_COLOR = RGBColor(0x00, 0x70, 0xC0)

# 可修复的属性集合，detector 用此过滤 fix_available 标志
# 注意：此处列出的 attr_key 必须与 checkers.py 中的键名一致
FIXABLE_ATTRS: frozenset[str] = frozenset({
    'font.size_pt',
    'font.cjk',
    'font.bold',
    'para.align',
    'para.first_line_indent_chars',
})


def fix_v2(file: str, issue: dict, value: dict[str, Any]) -> dict:
    """按 issue.attr + value[attr] 修改段落，并对修改的 run 标蓝

    业务逻辑：
    1. 从 issue 中取 para_idx 和 attr，从 value 中取期望值
    2. 按 attr 分发到对应的属性写入逻辑
    3. 写入成功后对第一个 run 标蓝（用户复核标记）
    4. 保存文件，返回 diff 摘要

    @param file   - docx 文件绝对路径（直接修改原文件）
    @param issue  - 检测器产出的 issue 字典，必须含 para_idx 和 attr
    @param value  - 期望值字典，key 与 attr 对应，如 {'font.size_pt': 12}
    @returns      - {diff, applied, xml_changed}；applied=False 表示无法处理
    """
    doc = Document(file)
    para_idx: int = issue['para_idx']
    attr: str = issue['attr']
    expected = value.get(attr)

    # 期望值缺失：无法修复，返回 applied=False
    if expected is None:
        return {'diff': '', 'applied': False, 'xml_changed': []}

    # 文档级属性（para_idx=-1）暂不支持 v2 fixer，留后续扩展
    if para_idx < 0 or para_idx >= len(doc.paragraphs):
        return {'diff': '', 'applied': False, 'xml_changed': []}

    para = doc.paragraphs[para_idx]
    if not para.runs:
        return {'diff': '', 'applied': False, 'xml_changed': []}

    run = para.runs[0]
    before_summary = f'{attr}: ?'

    # ── 按 attr 分发修改逻辑 ──
    if attr == 'font.size_pt':
        before_summary = f'font.size_pt: {run.font.size.pt if run.font.size else "?"}'
        run.font.size = Pt(expected)

    elif attr == 'font.cjk':
        before_summary = f'font.cjk: ?'
        # 中文字体写在 w:rFonts/@w:eastAsia，python-docx 无直接属性，需操作 XML
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            from lxml import etree
            rfonts = etree.SubElement(rpr, qn('w:rFonts'))
        rfonts.set(qn('w:eastAsia'), expected)

    elif attr == 'font.bold':
        before_summary = f'font.bold: {run.font.bold}'
        run.font.bold = expected

    elif attr == 'para.align':
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        # 字符串对齐名 → WD_ALIGN_PARAGRAPH 枚举值
        align_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        before_summary = f'para.align: {para.paragraph_format.alignment}'
        para.paragraph_format.alignment = align_map.get(expected)

    elif attr == 'para.first_line_indent_chars':
        before_summary = f'para.first_line_indent_chars: ?'
        # 按正文 12pt 字号换算字符数到 EMU（1pt = 12700 EMU）
        # 2个字符首行缩进 = 2 × 12pt = 24pt
        para.paragraph_format.first_line_indent = Pt(expected * 12)

    else:
        # 未知 attr，不支持修复
        return {'diff': '', 'applied': False, 'xml_changed': []}

    # ── 蓝色标记：提示用户此处已被自动修复 ──
    run.font.color.rgb = _MARK_COLOR

    doc.save(file)

    return {
        'diff': f'- {before_summary}\n+ {attr}: {expected}',
        'applied': True,
        'xml_changed': [f'w:p[{para_idx}]'],
    }
