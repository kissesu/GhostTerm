"""
@file: detector.py
@description: v2 规则引擎检测器：按字段遍历 template.rules，
              对每个字段定位论文对应段落，用 checkers 检查每条属性约束。
              支持段落级属性和文档级属性的分流处理。
@author: Atlas.oi
@date: 2026-04-18
"""
from typing import Any
from docx import Document

from .field_defs import get_field
from .checkers import CHECKER_MAP, DOC_LEVEL_KEYS


def _find_paragraphs_for_field(doc, field_id: str) -> list[int]:
    """定位字段对应段落 index 列表

    路由规则：
    - chapter_title：所有 style 含 Heading 1 的段，加上 FIELD_KEYWORDS 关键词命中的段
    - body_para：所有非 heading 非空段落
    - 其他：按 FIELD_KEYWORDS 关键词匹配正文文本
    """
    from ..extractor.field_matcher import FIELD_KEYWORDS
    keywords = FIELD_KEYWORDS.get(field_id, [])
    indices: list[int] = []

    # 先按关键词匹配
    for idx, p in enumerate(doc.paragraphs):
        if any(kw in p.text for kw in keywords):
            indices.append(idx)

    if field_id == 'chapter_title':
        # 补充所有 Heading 1 style 段（中英文 style 名都覆盖）
        for idx, p in enumerate(doc.paragraphs):
            style_name = p.style.name if p.style else ''
            if style_name in ('Heading 1', 'Heading1', '一级标题'):
                if idx not in indices:
                    indices.append(idx)

    if field_id == 'body_para':
        # 所有非空非 heading 段均视为正文段
        for idx, p in enumerate(doc.paragraphs):
            style_name = p.style.name if p.style else ''
            if 'Heading' in style_name or '标题' in style_name:
                continue
            if not p.text.strip():
                continue
            if idx not in indices:
                indices.append(idx)

    return indices


def detect_v2(file: str, template: dict[str, Any]) -> list[dict]:
    """执行 v2 规则检测，返回 Issue 字典列表

    每条 Issue 包含：
      - rule_id: 字段 ID（如 'chapter_title'）
      - attr: 违规属性 key（如 'font.size_pt'）
      - actual: 文档中检测到的实际值
      - expected: 规则期望值
      - para_idx: 段落索引（-1 表示文档级属性）
      - message: 人可读的错误描述
      - loc: {para, run} 用于 fix 定位
      - current: 同 actual（兼容 P3 fix handler 签名）
      - fix_available: 是否可自动修复
      - snippet: 段落前 30 字（方便 UI 展示上下文）
      - context: 段落前 60 字（方便 UI 展示更完整上下文）

    业务逻辑：
    1. 遍历 template.rules，跳过 enabled=False 的字段
    2. 按 DOC_LEVEL_KEYS 将属性分为文档级和段落级两组
    3. 文档级属性直接对 doc 对象调用 checker
    4. 段落级属性先定位目标段落，再逐段逐属性调用 checker
    5. checker 返回 None 表示符合规范，返回 dict 表示违规
    """
    doc = Document(file)
    rules = template.get('rules', {})
    issues: list[dict] = []

    for field_id, field_cfg in rules.items():
        # 跳过禁用的字段规则
        if not field_cfg.get('enabled', False):
            continue
        value = field_cfg.get('value', {})
        if not value:
            continue

        # 将属性分流为文档级和段落级两组
        doc_level_attrs = {k: v for k, v in value.items() if k in DOC_LEVEL_KEYS}
        para_level_attrs = {k: v for k, v in value.items() if k not in DOC_LEVEL_KEYS}

        # ── 文档级属性检查（不需要定位段落）──
        for attr_key, expected in doc_level_attrs.items():
            checker = CHECKER_MAP.get(attr_key)
            if checker is None:
                continue
            result = checker(doc, expected)
            if result is None:
                continue
            issues.append({
                'rule_id': field_id,
                'attr': attr_key,
                'actual': result['actual'],
                'expected': expected,
                'para_idx': -1,
                'message': f'{field_id}.{attr_key}: actual={result["actual"]} expected={expected}',
                'loc': {'para': -1, 'run': 0},
                'current': result['actual'],
                'fix_available': True,
                'snippet': '',
                'context': f'文档级属性 {attr_key}',
            })

        if not para_level_attrs:
            continue

        # ── 段落级属性检查（先定位目标段落）──
        para_indices = _find_paragraphs_for_field(doc, field_id)
        if not para_indices:
            continue

        for para_idx in para_indices:
            para = doc.paragraphs[para_idx]
            for attr_key, expected in para_level_attrs.items():
                checker = CHECKER_MAP.get(attr_key)
                if checker is None:
                    continue
                result = checker(para, expected)
                if result is None:
                    continue
                # snippet/context 截取段落文本供 UI 展示
                snippet = para.text[:30]
                context = para.text[:60]
                issues.append({
                    'rule_id': field_id,
                    'attr': attr_key,
                    'actual': result['actual'],
                    'expected': expected,
                    'para_idx': para_idx,
                    'message': f'{field_id}.{attr_key}: actual={result["actual"]} expected={expected}',
                    'loc': {'para': para_idx, 'run': 0},
                    'current': result['actual'],
                    'fix_available': True,
                    'snippet': snippet,
                    'context': context,
                })

    return issues
