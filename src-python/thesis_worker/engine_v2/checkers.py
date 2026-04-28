"""
@file: checkers.py
@description: 属性 key → checker 函数映射
              每个 checker 接收段落对象（或 Document 对象）和期望值，
              返回 None（符合规范）或 dict（违规描述）。
              分三类：
                A — 段落级 (para, expected) -> Optional[dict]
                B — 文档级 (doc, expected) -> Optional[dict]
                C — 延后存根，暂不支持，固定返回 None
@author: Atlas.oi
@date: 2026-04-18
"""
from typing import Any, Optional

from docx.document import Document
from docx.text.paragraph import Paragraph

# 复用 extractor 中已实现的全文启发式扫描，避免重复维护两份正则+阈值
# extractor 包不依赖 engine_v2，反向 import 无循环
from ..extractor.pipeline import _detect_punct_space_after, _read_numbering_styles

# ───────────────────────────────────────────────
# 容差常量（用于浮点属性的近似匹配）
# ───────────────────────────────────────────────
_TOL_FONT_PT = 0.10       # 字号容差（pt）
_TOL_MARGIN_CM = 0.05     # 页边距容差（cm）
_TOL_SIZE_INCH = 0.05     # 页面尺寸容差（英寸）
_TOL_LINE_SPACING = 0.05  # 行距倍数容差
_TOL_SPACE_PT = 0.5       # 段前/段后间距容差（pt，对应 Word UI 显示精度）
_TOL_OFFSET_CM = 0.05     # 页眉/页脚距边界容差（语义独立于 _TOL_MARGIN_CM，数值相同）

# ───────────────────────────────────────────────
# 内部工具函数
# ───────────────────────────────────────────────

# OOXML 命名空间 — Word 主命名空间 URI
_W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def _w(tag: str) -> str:
    """生成带命名空间的 OOXML 标签，避免反复拼接字符串"""
    return f'{{{_W_NS}}}{tag}'


def _first_nonempty_run(para: Paragraph):
    """返回段落第一个有可见文字的 run；没有则返回 None"""
    for run in para.runs:
        if run.text.strip():
            return run
    return None


def _read_cjk_font(para: Paragraph) -> Optional[str]:
    """读段落第一个非空 run 的 eastAsia 字体名（中文字体）"""
    run = _first_nonempty_run(para)
    if run is None:
        return None
    rpr = run._element.rPr
    if rpr is None:
        return None
    rfonts = rpr.find(_w('rFonts'))
    if rfonts is None:
        return None
    return rfonts.get(_w('eastAsia'))


def _read_ascii_font(para: Paragraph) -> Optional[str]:
    """读段落第一个非空 run 的 ASCII/西文字体名"""
    run = _first_nonempty_run(para)
    if run is None:
        return None
    # python-docx 的 run.font.name 读的是 w:ascii 属性
    name = run.font.name
    if name:
        return name
    # 兜底：直接读 XML
    rpr = run._element.rPr
    if rpr is None:
        return None
    rfonts = rpr.find(_w('rFonts'))
    if rfonts is None:
        return None
    return rfonts.get(_w('ascii'))


# ───────────────────────────────────────────────
# 类别 A：段落级 checker
# ───────────────────────────────────────────────

def check_font_cjk(para: Paragraph, expected: str) -> Optional[dict]:
    """检查段落中文字体（eastAsia）是否符合预期"""
    actual = _read_cjk_font(para)
    if actual == expected:
        return None
    return {'attr': 'font.cjk', 'actual': actual, 'expected': expected}


def check_font_ascii(para: Paragraph, expected: str) -> Optional[dict]:
    """检查段落西文/ASCII 字体是否符合预期"""
    actual = _read_ascii_font(para)
    if actual == expected:
        return None
    return {'attr': 'font.ascii', 'actual': actual, 'expected': expected}


def check_font_size_pt(para: Paragraph, expected: float) -> Optional[dict]:
    """检查段落字号（pt），容差 0.1pt，取第一个非空 run"""
    run = _first_nonempty_run(para)
    if run is None:
        return None
    if run.font.size is None:
        return {'attr': 'font.size_pt', 'actual': None, 'expected': expected}
    actual = float(run.font.size.pt)
    if abs(actual - expected) < _TOL_FONT_PT:
        return None
    return {'attr': 'font.size_pt', 'actual': actual, 'expected': expected}


def check_font_bold(para: Paragraph, expected: bool) -> Optional[dict]:
    """检查段落第一个非空 run 是否加粗；None 视为 False"""
    run = _first_nonempty_run(para)
    if run is None:
        return None
    actual = bool(run.font.bold)
    if actual == expected:
        return None
    return {'attr': 'font.bold', 'actual': actual, 'expected': expected}


def check_font_italic(para: Paragraph, expected: bool) -> Optional[dict]:
    """检查段落第一个非空 run 是否斜体；None 视为 False"""
    run = _first_nonempty_run(para)
    if run is None:
        return None
    # run.font.italic 三值：None / True / False，None 视为 False
    actual = bool(run.font.italic)
    if actual == expected:
        return None
    return {'attr': 'font.italic', 'actual': actual, 'expected': expected}


def check_para_align(para: Paragraph, expected: str) -> Optional[dict]:
    """检查段落对齐方式；expected: 'left'/'center'/'right'/'justify'"""
    # python-docx WD_ALIGN_PARAGRAPH 枚举值：LEFT=0,CENTER=1,RIGHT=2,JUSTIFY=3
    align_map = {0: 'left', 1: 'center', 2: 'right', 3: 'justify'}
    val = para.paragraph_format.alignment
    actual = align_map.get(int(val), 'left') if val is not None else 'left'
    if actual == expected:
        return None
    return {'attr': 'para.align', 'actual': actual, 'expected': expected}


def check_para_first_line_indent_chars(
    para: Paragraph,
    expected: int,
    body_size_pt: float = 12,
) -> Optional[dict]:
    """
    检查首行缩进字符数；body_size_pt 用于将 pt 换算为字符数（默认 12pt）。
    first_line_indent > 0 为首行缩进；< 0 为悬挂缩进（不属于本检查器）。
    """
    fli = para.paragraph_format.first_line_indent
    if fli is None:
        actual_chars = 0
    else:
        # 四舍五入：1 个字宽 = body_size_pt pt
        actual_chars = round(fli.pt / body_size_pt)
    if actual_chars == expected:
        return None
    return {'attr': 'para.first_line_indent_chars', 'actual': actual_chars, 'expected': expected}


def check_para_hanging_indent_chars(
    para: Paragraph,
    expected: int,
    body_size_pt: float = 12,
) -> Optional[dict]:
    """
    检查悬挂缩进字符数（参考文献条目）。
    python-docx 用 first_line_indent < 0 表示悬挂缩进，abs 后换算为字符数。
    """
    fli = para.paragraph_format.first_line_indent
    if fli is None:
        actual_chars = 0
    elif fli.pt >= 0:
        # 首行缩进而非悬挂缩进，字符数记为 0
        actual_chars = 0
    else:
        actual_chars = round(abs(fli.pt) / body_size_pt)
    if actual_chars == expected:
        return None
    return {'attr': 'para.hanging_indent_chars', 'actual': actual_chars, 'expected': expected}


def check_para_line_spacing(para: Paragraph, expected: float) -> Optional[dict]:
    """
    检查行距倍数（如 1.5、2.0）。
    line_spacing 为 None 时视为 1.0（单倍）；容差 0.05。
    """
    ls = para.paragraph_format.line_spacing
    if ls is None:
        actual = 1.0
    else:
        # line_spacing 可能是 float（倍数）也可能是 Length（固定值，单位 emu）
        # 若为 Length 对象则取 pt 值再除以字号 12 近似换算
        try:
            actual = float(ls)
        except (TypeError, AttributeError):
            actual = 1.0
    if abs(actual - expected) < _TOL_LINE_SPACING:
        return None
    return {'attr': 'para.line_spacing', 'actual': actual, 'expected': expected}


def check_para_space_before_lines(
    para: Paragraph,
    expected: int,
    body_size_pt: float = 12,
) -> Optional[dict]:
    """检查段前间距（行数）；space_before 为 None 时视为 0"""
    sb = para.paragraph_format.space_before
    if sb is None:
        actual = 0
    else:
        actual = round(sb.pt / body_size_pt)
    if actual == expected:
        return None
    return {'attr': 'para.space_before_lines', 'actual': actual, 'expected': expected}


def check_para_space_after_lines(
    para: Paragraph,
    expected: int,
    body_size_pt: float = 12,
) -> Optional[dict]:
    """检查段后间距（行数）；space_after 为 None 时视为 0"""
    sa = para.paragraph_format.space_after
    if sa is None:
        actual = 0
    else:
        actual = round(sa.pt / body_size_pt)
    if actual == expected:
        return None
    return {'attr': 'para.space_after_lines', 'actual': actual, 'expected': expected}


def check_para_space_before_pt(para: Paragraph, expected: float) -> Optional[dict]:
    """检查段前间距（磅值）；space_before 为 None 时视为 0；容差 0.5pt

    与 check_para_space_before_lines 使用同一数据源（para.paragraph_format.space_before），
    但不除以行基准值，直接保留 pt 数值。规范文本用"磅"描述时用本检查器，
    用"行"描述时用 _lines 版本，两者互不冲突。
    """
    sb = para.paragraph_format.space_before
    actual = round(sb.pt, 1) if sb is not None else 0.0
    if abs(actual - expected) < _TOL_SPACE_PT:
        return None
    return {'attr': 'para.space_before_pt', 'actual': actual, 'expected': expected}


def check_para_space_after_pt(para: Paragraph, expected: float) -> Optional[dict]:
    """检查段后间距（磅值）；space_after 为 None 时视为 0；容差 0.5pt"""
    sa = para.paragraph_format.space_after
    actual = round(sa.pt, 1) if sa is not None else 0.0
    if abs(actual - expected) < _TOL_SPACE_PT:
        return None
    return {'attr': 'para.space_after_pt', 'actual': actual, 'expected': expected}


def check_para_letter_spacing_chars(
    para: Paragraph,
    expected: int,
    body_size_pt: float = 12,
) -> Optional[dict]:
    """
    检查字符间距（字符数）。
    OOXML w:spacing 在 run 的 rPr 中，单位为"二十分之一磅"（1/20 pt）。
    换算：actual_chars = round(spacing_pt / body_size_pt)。
    """
    run = _first_nonempty_run(para)
    if run is None:
        return None
    rpr = run._element.rPr
    if rpr is None:
        spacing_pt = 0.0
    else:
        spacing_el = rpr.find(_w('spacing'))
        if spacing_el is None:
            spacing_pt = 0.0
        else:
            val = spacing_el.get(_w('val'))
            # val 单位：twip (1/20 pt)
            spacing_pt = int(val) / 20.0 if val is not None else 0.0
    actual = round(spacing_pt / body_size_pt)
    if actual == expected:
        return None
    return {'attr': 'para.letter_spacing_chars', 'actual': actual, 'expected': expected}


def check_page_new_page_before(para: Paragraph, expected: bool) -> Optional[dict]:
    """检查段前分页（一级标题/参考文献/致谢等需要另起一页）"""
    val = para.paragraph_format.page_break_before
    # python-docx 返回 None/True/False，None 视为 False
    actual = bool(val)
    if actual == expected:
        return None
    return {'attr': 'page.new_page_before', 'actual': actual, 'expected': expected}


# page_break_after 不在 python-docx ParagraphFormat API 中，
# 需要直接读 XML pPr/pageBreakAfter 元素（deferred to v3，实际不在 spec）


def check_content_char_count_min(para: Paragraph, expected: int) -> Optional[dict]:
    """检查段落字符数不少于下限（适用于摘要正文最少字数）"""
    actual = len(para.text)
    if actual >= expected:
        return None
    return {'attr': 'content.char_count_min', 'actual': actual, 'expected': expected}


def check_content_char_count_max(para: Paragraph, expected: int) -> Optional[dict]:
    """检查段落字符数不超过上限（适用于摘要正文最多字数）"""
    actual = len(para.text)
    if actual <= expected:
        return None
    return {'attr': 'content.char_count_max', 'actual': actual, 'expected': expected}


def check_content_item_count_min(
    para: Paragraph,
    expected: int,
    separator: str = '；',
) -> Optional[dict]:
    """
    检查段落以分隔符划分的条目数不少于下限（如关键词至少 3 个）。
    默认分隔符为中文分号。
    """
    actual = len(para.text.split(separator))
    if actual >= expected:
        return None
    return {'attr': 'content.item_count_min', 'actual': actual, 'expected': expected}


def check_content_item_count_max(
    para: Paragraph,
    expected: int,
    separator: str = '；',
) -> Optional[dict]:
    """检查段落条目数不超过上限（如关键词最多 8 个）"""
    actual = len(para.text.split(separator))
    if actual <= expected:
        return None
    return {'attr': 'content.item_count_max', 'actual': actual, 'expected': expected}


def check_content_item_separator(para: Paragraph, expected: str) -> Optional[dict]:
    """检查段落内是否使用了规定的分隔符（如中文分号）"""
    if expected in para.text:
        return None
    return {'attr': 'content.item_separator', 'actual': None, 'expected': expected}


def check_content_specific_text(para: Paragraph, expected: str) -> Optional[dict]:
    """检查段落是否包含特定文本（如「摘要」、「关键词：」等固定标题词）"""
    if expected in para.text:
        return None
    return {'attr': 'content.specific_text', 'actual': para.text, 'expected': expected}


# ───────────────────────────────────────────────
# 类别 B：文档级 checker（接收 docx.Document）
# ───────────────────────────────────────────────

# A4 页面尺寸（英寸），容差 0.05 英寸
_PAGE_SIZES = {
    'A4':     (8.27, 11.69),
    'Letter': (8.50, 11.00),
}


def check_page_size(doc: Document, expected: str) -> Optional[dict]:
    """
    检查页面大小是否符合预期（'A4' 或 'Letter'）。
    通过 sections[0] 的宽高（英寸）与已知规格比对，容差 0.05 英寸。
    """
    section = doc.sections[0]
    w_in = section.page_width.inches
    h_in = section.page_height.inches
    for name, (sw, sh) in _PAGE_SIZES.items():
        if abs(w_in - sw) < _TOL_SIZE_INCH and abs(h_in - sh) < _TOL_SIZE_INCH:
            actual = name
            break
    else:
        actual = f'{w_in:.2f}x{h_in:.2f}in'
    if actual == expected:
        return None
    return {'attr': 'page.size', 'actual': actual, 'expected': expected}


def check_page_margin_top_cm(doc: Document, expected: float) -> Optional[dict]:
    """检查页面上边距（cm），容差 0.05cm"""
    actual = doc.sections[0].top_margin.cm
    if abs(actual - expected) < _TOL_MARGIN_CM:
        return None
    return {'attr': 'page.margin_top_cm', 'actual': round(actual, 2), 'expected': expected}


def check_page_margin_bottom_cm(doc: Document, expected: float) -> Optional[dict]:
    """检查页面下边距（cm），容差 0.05cm"""
    actual = doc.sections[0].bottom_margin.cm
    if abs(actual - expected) < _TOL_MARGIN_CM:
        return None
    return {'attr': 'page.margin_bottom_cm', 'actual': round(actual, 2), 'expected': expected}


def check_page_margin_left_cm(doc: Document, expected: float) -> Optional[dict]:
    """检查页面左边距（cm），容差 0.05cm"""
    actual = doc.sections[0].left_margin.cm
    if abs(actual - expected) < _TOL_MARGIN_CM:
        return None
    return {'attr': 'page.margin_left_cm', 'actual': round(actual, 2), 'expected': expected}


def check_page_margin_right_cm(doc: Document, expected: float) -> Optional[dict]:
    """检查页面右边距（cm），容差 0.05cm"""
    actual = doc.sections[0].right_margin.cm
    if abs(actual - expected) < _TOL_MARGIN_CM:
        return None
    return {'attr': 'page.margin_right_cm', 'actual': round(actual, 2), 'expected': expected}


def check_page_margin_gutter_cm(doc: Document, expected: float) -> Optional[dict]:
    """检查装订线宽度（cm）；容差 0.05cm

    gutter 为 0 也是有效值（表示无装订线），不做 None 特殊处理。
    python-docx sections[0].gutter.cm 直接读取 OOXML w:pgMar/@w:gutter 属性。
    """
    actual = doc.sections[0].gutter.cm
    if abs(actual - expected) < _TOL_OFFSET_CM:
        return None
    return {'attr': 'page.margin_gutter_cm', 'actual': round(actual, 2), 'expected': expected}


def check_page_header_offset_cm(doc: Document, expected: float) -> Optional[dict]:
    """检查页眉距页面边界的距离（cm）；容差 0.05cm

    对应 OOXML w:pgMar/@w:header 属性，Word 中称"页眉边距"。
    """
    actual = doc.sections[0].header_distance.cm
    if abs(actual - expected) < _TOL_OFFSET_CM:
        return None
    return {'attr': 'page.header_offset_cm', 'actual': round(actual, 2), 'expected': expected}


def check_page_footer_offset_cm(doc: Document, expected: float) -> Optional[dict]:
    """检查页脚距页面边界的距离（cm）；容差 0.05cm

    对应 OOXML w:pgMar/@w:footer 属性，Word 中称"页脚边距"。
    """
    actual = doc.sections[0].footer_distance.cm
    if abs(actual - expected) < _TOL_OFFSET_CM:
        return None
    return {'attr': 'page.footer_offset_cm', 'actual': round(actual, 2), 'expected': expected}


def check_page_print_mode(doc: Document, expected: str) -> Optional[dict]:
    """检查打印模式（'single' 单面 / 'double' 双面）；严格相等，无容差

    判别方式：检测 w:settings 根元素是否含 w:evenAndOddHeaders 子元素。
    存在该元素 → 文档启用奇偶页眉分设（双面打印模式）→ 'double'；
    不存在 → 'single'。这是 Word 保存双面打印设置时的标准 OOXML 语义。
    """
    from docx.oxml.ns import qn
    settings_el = doc.settings.element
    even_odd = settings_el.find(qn('w:evenAndOddHeaders'))
    actual = 'double' if even_odd is not None else 'single'
    if actual == expected:
        return None
    return {'attr': 'page.print_mode', 'actual': actual, 'expected': expected}


def check_mixed_script_ascii_is_tnr(doc: Document, expected: bool) -> Optional[dict]:
    """
    检查文档西文/数字字体是否全局使用 Times New Roman（TNR）。
    遍历正文段落 run，统计 ASCII 字体名，多数为 TNR 则视为满足。
    [延后 v3：当前仅做简单多数决；未来接入精确 field 过滤]
    """
    tnr_count = 0
    total_count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if not run.text.strip():
                continue
            name = run.font.name
            if name is None:
                continue
            total_count += 1
            if 'Times New Roman' in name or name.lower() == 'tnr':
                tnr_count += 1
    if total_count == 0:
        # 文档无有效 run，无法判断，视为符合
        return None
    actual = (tnr_count / total_count) >= 0.5
    if actual == expected:
        return None
    return {'attr': 'mixed_script.ascii_is_tnr', 'actual': actual, 'expected': expected}


def check_mixed_script_punct_space_after(doc: Document, expected: bool) -> Optional[dict]:
    """
    检查文档英文标点后是否规范空一字符。
    扫描逻辑复用 pipeline._detect_punct_space_after，保持 extractor 与 checker 口径一致。
    detector 返回 None 表示样本不足 — 此时视为"无法判定"，不报告 issue。
    """
    actual = _detect_punct_space_after(doc)
    if actual is None:
        return None
    if actual == expected:
        return None
    return {'attr': 'mixed_script.punct_space_after', 'actual': actual, 'expected': expected}


# ───────────────────────────────────────────────
# 类别 B（续）：T3.2 table.* namespace（表格线宽/三线表）
# ───────────────────────────────────────────────

def _read_tbl_borders(doc: Document) -> dict[str, float]:
    """从文档第一个表格读取 OOXML tblBorders 各方向线宽（单位 pt）。

    业务逻辑：
    1. 检查 doc.tables 是否为空，空则返回空字典（不报错）
    2. 取第一个表格，找 tblPr/tblBorders 元素
    3. 遍历子元素，读 w:sz 属性（eighth-points 单位，除以 8 得 pt）
    4. 返回 {tag: pt} 字典，只含能读到 sz 的方向

    注意：OOXML w:sz 单位为 eighth-points（1/8 pt），不是 twips（1/20 pt）。
    """
    from docx.oxml.ns import qn
    borders: dict[str, float] = {}
    if not doc.tables:
        return borders
    tbl = doc.tables[0]
    tbl_pr = tbl._element.find(qn('w:tblPr'))
    if tbl_pr is None:
        return borders
    tbl_borders = tbl_pr.find(qn('w:tblBorders'))
    if tbl_borders is None:
        return borders
    for child in tbl_borders:
        # 去掉命名空间前缀，取 localname（如 'top'/'bottom'/'insideH' 等）
        local = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        sz_val = child.get(qn('w:sz'))
        if sz_val is not None:
            try:
                # eighth-points → pt：除以 8
                borders[local] = int(sz_val) / 8.0
            except ValueError:
                pass  # sz 值非法，跳过
    return borders


def check_table_is_three_line(doc: Document, expected: bool) -> 'Optional[dict]':
    """检查文档第一个表格是否符合三线表规范。

    三线表判定逻辑（简化版）：
    - top > 0（有上边框）
    - bottom > 0（有下边框）
    - insideV 不存在或 insideV.sz == 0（无竖向内线，即无纵格线）
    若 doc 无表格，视为"不是三线表"（返回 False 与 expected 比对）。
    """
    borders = _read_tbl_borders(doc)
    # 统一走数值判定，不区分"无 tblBorders"与"sz 全缺"两种情况。
    # 两者结果均为 top=0/bottom=0 → actual=False，业务语义上都不能确认为三线表。
    # 这消除了"空 dict 被误判为无 tblBorders"的混淆，与 check_table_border_* 行为对齐。
    top_pt = borders.get('top', 0.0)
    bottom_pt = borders.get('bottom', 0.0)
    inside_v_pt = borders.get('insideV', 0.0)
    # 三线表：有顶线/底线，无竖向内线
    actual = top_pt > 0 and bottom_pt > 0 and inside_v_pt == 0.0
    if actual == expected:
        return None
    return {'attr': 'table.is_three_line', 'actual': actual, 'expected': expected}


def check_table_border_top_pt(doc: Document, expected: float) -> 'Optional[dict]':
    """检查文档第一个表格上边框线宽（pt），容差 0.1pt。

    若 doc 无表格或无 tblBorders，actual 视为 0.0。
    """
    borders = _read_tbl_borders(doc)
    actual = borders.get('top', 0.0)
    if abs(actual - expected) < _TOL_FONT_PT:
        return None
    return {'attr': 'table.border_top_pt', 'actual': actual, 'expected': expected}


def check_table_border_bottom_pt(doc: Document, expected: float) -> 'Optional[dict]':
    """检查文档第一个表格下边框线宽（pt），容差 0.1pt。

    若 doc 无表格或无 tblBorders，actual 视为 0.0。
    """
    borders = _read_tbl_borders(doc)
    actual = borders.get('bottom', 0.0)
    if abs(actual - expected) < _TOL_FONT_PT:
        return None
    return {'attr': 'table.border_bottom_pt', 'actual': actual, 'expected': expected}


def check_table_header_border_pt(doc: Document, expected: float) -> 'Optional[dict]':
    """检查文档第一个表格表头下边框线宽（pt），容差 0.1pt。

    使用 insideH（水平内线宽度）作为表头下边框的简化代理值。
    三线表中表头下线即内部水平线，此值与表头下边框线宽在绝大多数规范模板中一致。
    若 doc 无表格或无 tblBorders，actual 视为 0.0。
    """
    borders = _read_tbl_borders(doc)
    actual = borders.get('insideH', 0.0)
    if abs(actual - expected) < _TOL_FONT_PT:
        return None
    return {'attr': 'table.header_border_pt', 'actual': actual, 'expected': expected}


# ───────────────────────────────────────────────
# 类别 B（续）：T3.3 numbering.* namespace（编号风格）
# ───────────────────────────────────────────────

def check_numbering_figure_style(doc: Document, expected: str) -> Optional[dict]:
    """检查图编号风格是否符合预期（'continuous' 连续 / 'chapter_based' 章节式）。

    复用 pipeline._read_numbering_styles 的启发式多数票算法，
    返回 None 表示：符合规范 OR 样本不足无法判定（两者均不报 issue）。
    样本不足（< 2 个图题）时视为无法判定，保守处理不报违规。
    """
    styles = _read_numbering_styles(doc)
    actual = styles.get('numbering.figure_style')
    if actual is None:
        # 样本不足，无法判定，不报 issue
        return None
    if actual == expected:
        return None
    return {'attr': 'numbering.figure_style', 'actual': actual, 'expected': expected}


def check_numbering_subfigure_style(doc: Document, expected: str) -> Optional[dict]:
    """检查分图编号风格是否符合预期（'a_b_c' 字母 / '1_2_3' 数字点号）。

    subfigure_style 仅在 figure_style 已确定时由 pipeline 推断；
    若返回 None 表示文档无子图标记，视为无法判定，不报 issue。
    """
    styles = _read_numbering_styles(doc)
    actual = styles.get('numbering.subfigure_style')
    if actual is None:
        return None
    if actual == expected:
        return None
    return {'attr': 'numbering.subfigure_style', 'actual': actual, 'expected': expected}


def check_numbering_formula_style(doc: Document, expected: str) -> Optional[dict]:
    """检查公式编号风格是否符合预期（'continuous' 连续 / 'chapter_based' 章节式）。

    复用 pipeline._read_numbering_styles 多数票推断；
    样本不足（< 2 个公式编号）时返回 None，不报 issue。
    """
    styles = _read_numbering_styles(doc)
    actual = styles.get('numbering.formula_style')
    if actual is None:
        return None
    if actual == expected:
        return None
    return {'attr': 'numbering.formula_style', 'actual': actual, 'expected': expected}


# ───────────────────────────────────────────────
# 类别 C：T4.1 实现 deferred checker
# ───────────────────────────────────────────────

import re as _re

# 用于判断段落是否含图（w:drawing）或表（w:tbl）的 OOXML 标签名
_DRAWING_TAG = _w('drawing')
_TBL_TAG = _w('tbl')


def _para_el_contains_figure_or_table(el) -> bool:
    """判断给定 XML 元素（段落 _element 或其 r 元素）是否包含图片（w:drawing）或表格（w:tbl）。

    直接递归遍历子树，确保嵌套情况也能覆盖。
    """
    for child in el:
        if child.tag in (_DRAWING_TAG, _TBL_TAG):
            return True
        if _para_el_contains_figure_or_table(child):
            return True
    return False


def check_layout_position(para: Paragraph, expected: str) -> Optional[dict]:
    """T4.1 实现：检查图题/表题的位置（'above'/'below'）。

    业务逻辑：
    1. 通过 para._element 找到它在文档 body 中的索引
    2. 检查前驱元素（previous sibling）是否含 w:drawing 或 w:tbl
       - 前驱含图/表 → caption 在图/表"下方"（below）
    3. 检查后继元素（next sibling）是否含 w:drawing 或 w:tbl
       - 后继含图/表 → caption 在图/表"上方"（above）
    4. 前后都没有图/表邻接 → 无法判定，返回 None

    注意：只检查直接相邻的段落级元素（一步之遥），
    不做跨段落的深度搜索，避免误判。
    """
    el = para._element
    parent = el.getparent()
    if parent is None:
        # 段落无父节点（异常文档结构），无法判定
        return None

    children = list(parent)
    try:
        idx = children.index(el)
    except ValueError:
        return None

    # 检查前驱段落（index - 1）
    if idx > 0 and _para_el_contains_figure_or_table(children[idx - 1]):
        # 前驱是图/表 → caption 在图/表下方
        actual = 'below'
        if actual == expected:
            return None
        return {'attr': 'layout.position', 'actual': actual, 'expected': expected}

    # 检查后继段落（index + 1）
    if idx < len(children) - 1 and _para_el_contains_figure_or_table(children[idx + 1]):
        # 后继是图/表 → caption 在图/表上方
        actual = 'above'
        if actual == expected:
            return None
        return {'attr': 'layout.position', 'actual': actual, 'expected': expected}

    # 前后都不是图/表（独立 caption 段落），无法判定
    return None


# 引用风格判定正则
# GB/T 7714 特征：序号 [N] 开头
_RE_GBT7714 = _re.compile(r'^\[\d+\]')
# APA 特征：含括号包围年份 (YYYY)
_RE_APA = _re.compile(r'\(\d{4}\)')


def check_citation_style(para: Paragraph, expected: str) -> Optional[dict]:
    """T4.1 实现：检查参考文献条目的引用风格（'gbt7714' / 'apa' / 'mla'）。

    启发式判定规则（保守）：
    - 文本以 [N] 开头（regex ^\\[\\d+\\]）→ actual='gbt7714'
    - 文本含 (YYYY) 括号年份（regex \\(\\d{4}\\)）→ actual='apa'
    - 两条都不符合 → actual=None，视为无法判定，不报 issue

    MLA 格式无简单特征正则可判别，保守处理为 None（不报违规）。
    """
    text = para.text.strip()
    if not text:
        # 空段落，无法判定
        return None

    if _RE_GBT7714.match(text):
        actual: Optional[str] = 'gbt7714'
    elif _RE_APA.search(text):
        actual = 'apa'
    else:
        # 无法推断风格（含 MLA），保守处理
        actual = None

    if actual is None:
        return None
    if actual == expected:
        return None
    return {'attr': 'citation.style', 'actual': actual, 'expected': expected}


def _read_pgnumtype_fmt(section_pr) -> Optional[str]:
    """从 sectPr XML 元素读取 w:pgNumType/@w:fmt 属性值。

    返回 OOXML 原始 fmt 字符串（如 'lowerRoman'/'upperRoman'/'decimal'），
    若不存在则返回 None（表示使用默认值 decimal/阿拉伯数字）。
    """
    from docx.oxml.ns import qn
    pg_num_type = section_pr.find(qn('w:pgNumType'))
    if pg_num_type is None:
        return None
    return pg_num_type.get(qn('w:fmt'))


def _ooxml_fmt_to_style(fmt: Optional[str]) -> str:
    """将 OOXML w:pgNumType fmt 值映射为业务风格字符串。

    映射表：
    - 'lowerRoman' / 'upperRoman' → 'roman'
    - 'decimal' / None（默认）     → 'arabic'
    """
    if fmt in ('lowerRoman', 'upperRoman'):
        return 'roman'
    # decimal 或不存在（Word 默认为阿拉伯数字）
    return 'arabic'


def check_pagination_front_style(doc: Document, expected: str) -> Optional[dict]:
    """T4.1 实现：检查前置页码样式（'roman' / 'arabic' / 'none'）。

    从第一个 section 的 sectPr 读 w:pgNumType/@w:fmt：
    - lowerRoman / upperRoman → 'roman'
    - decimal 或不存在       → 'arabic'
    """
    section_pr = doc.sections[0]._sectPr
    fmt = _read_pgnumtype_fmt(section_pr)
    actual = _ooxml_fmt_to_style(fmt)
    if actual == expected:
        return None
    return {'attr': 'pagination.front_style', 'actual': actual, 'expected': expected}


def check_pagination_body_style(doc: Document, expected: str) -> Optional[dict]:
    """T4.1 实现：检查正文页码样式（'arabic' / 'roman'）。

    从最后一个 section 的 sectPr 读 w:pgNumType/@w:fmt。
    若文档只有 1 个 section，front 和 body 样式相同。
    """
    section_pr = doc.sections[-1]._sectPr
    fmt = _read_pgnumtype_fmt(section_pr)
    actual = _ooxml_fmt_to_style(fmt)
    if actual == expected:
        return None
    return {'attr': 'pagination.body_style', 'actual': actual, 'expected': expected}


# ───────────────────────────────────────────────
# CHECKER_MAP：属性 key → checker 函数
# ───────────────────────────────────────────────
# 段落级函数签名：(para: Paragraph, expected) -> Optional[dict]
# 文档级函数签名：(doc: Document, expected) -> Optional[dict]
# detector (Task 10) 需要区分两类，约定：
#   - DOC_LEVEL_KEYS 中列出的 key 用文档级调用
#   - 其余均为段落级调用

CHECKER_MAP: dict[str, Any] = {
    # font 字体
    'font.cjk':                       check_font_cjk,
    'font.ascii':                      check_font_ascii,
    'font.size_pt':                    check_font_size_pt,
    'font.bold':                       check_font_bold,
    'font.italic':                     check_font_italic,
    # para 段落格式
    'para.align':                      check_para_align,
    'para.first_line_indent_chars':    check_para_first_line_indent_chars,
    'para.hanging_indent_chars':       check_para_hanging_indent_chars,
    'para.line_spacing':               check_para_line_spacing,
    'para.space_before_lines':         check_para_space_before_lines,
    'para.space_after_lines':          check_para_space_after_lines,
    # T3.1: 段前/段后磅值版本（与 _lines 版本共存）
    'para.space_before_pt':            check_para_space_before_pt,
    'para.space_after_pt':             check_para_space_after_pt,
    'para.letter_spacing_chars':       check_para_letter_spacing_chars,
    # page 分页/页面
    'page.new_page_before':            check_page_new_page_before,
    'page.size':                       check_page_size,
    'page.margin_top_cm':              check_page_margin_top_cm,
    'page.margin_bottom_cm':           check_page_margin_bottom_cm,
    'page.margin_left_cm':             check_page_margin_left_cm,
    'page.margin_right_cm':            check_page_margin_right_cm,
    # T3.1: 装订线/页眉脚距/打印模式（文档级）
    'page.margin_gutter_cm':           check_page_margin_gutter_cm,
    'page.header_offset_cm':           check_page_header_offset_cm,
    'page.footer_offset_cm':           check_page_footer_offset_cm,
    'page.print_mode':                 check_page_print_mode,
    # content 内容
    'content.char_count_min':          check_content_char_count_min,
    'content.char_count_max':          check_content_char_count_max,
    'content.item_count_min':          check_content_item_count_min,
    'content.item_count_max':          check_content_item_count_max,
    'content.item_separator':          check_content_item_separator,
    'content.specific_text':           check_content_specific_text,
    # mixed_script 中西混排
    'mixed_script.ascii_is_tnr':       check_mixed_script_ascii_is_tnr,
    'mixed_script.punct_space_after':  check_mixed_script_punct_space_after,
    # T3.2: table namespace（三线表判定 + 三条边框线宽）
    'table.is_three_line':             check_table_is_three_line,
    'table.border_top_pt':             check_table_border_top_pt,
    'table.border_bottom_pt':          check_table_border_bottom_pt,
    'table.header_border_pt':          check_table_header_border_pt,
    # T3.3: numbering namespace（图/分图/公式编号风格，文档级启发式推断）
    'numbering.figure_style':          check_numbering_figure_style,
    'numbering.subfigure_style':       check_numbering_subfigure_style,
    'numbering.formula_style':         check_numbering_formula_style,
    # layout / citation / pagination（延后存根）
    'layout.position':                 check_layout_position,
    'citation.style':                  check_citation_style,
    'pagination.front_style':          check_pagination_front_style,
    'pagination.body_style':           check_pagination_body_style,
}

# Task 10 detector 使用：需要传入 Document 而非 Paragraph 的属性 key 集合
DOC_LEVEL_KEYS: frozenset[str] = frozenset({
    'page.size',
    'page.margin_top_cm',
    'page.margin_bottom_cm',
    'page.margin_left_cm',
    'page.margin_right_cm',
    # T3.1: 新增 4 个文档级 page key
    'page.margin_gutter_cm',
    'page.header_offset_cm',
    'page.footer_offset_cm',
    'page.print_mode',
    'mixed_script.ascii_is_tnr',
    'mixed_script.punct_space_after',
    'pagination.front_style',
    'pagination.body_style',
    # T3.2: table namespace 4 个文档级 key（表格结构属于文档级，不依赖段落）
    'table.is_three_line',
    'table.border_top_pt',
    'table.border_bottom_pt',
    'table.header_border_pt',
    # T3.3: numbering namespace 3 个文档级 key（全文启发式扫描，不依赖单段落）
    'numbering.figure_style',
    'numbering.subfigure_style',
    'numbering.formula_style',
})
