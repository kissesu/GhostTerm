"""
@file: make_fixtures.py
@description: 一次性脚本，生成测试 fixture docx。跑完可删。
@author: Atlas.oi
@date: 2026-04-17
"""
from docx import Document
from pathlib import Path

out = Path(__file__).parent / 'fixtures'
out.mkdir(parents=True, exist_ok=True)

# bad: 中英间有空格
bad = Document()
bad.add_paragraph('这是 AI 工具。')
bad.add_paragraph('版本 2.1.1 已发布。')
bad.add_paragraph('无违规段落。')
bad.save(out / 'cjk_space_bad.docx')

# clean: 中英紧贴，无空格
clean = Document()
clean.add_paragraph('这是AI工具。')
clean.add_paragraph('版本2.1.1已发布。')
clean.save(out / 'cjk_space_clean.docx')

print('fixtures generated')
