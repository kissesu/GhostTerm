"""
@file: test_font_body.py
@description: font.body 规则的单元测试
              fixture 全部通过 tmp_path 在运行时生成（不依赖静态 .docx 文件）
@author: Atlas.oi
@date: 2026-04-18
"""
import pytest
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor

from thesis_worker.rules.font_body import FontBodyRule
from thesis_worker.models import Issue, Location


# ===========================================================================
# fixture 工厂函数（tmp_path 策略，CI 自包含）
# ===========================================================================

def _make_bad_font_body_doc(tmp_path: Path) -> Path:
    """创建包含字体/字号不符正文的测试文档"""
    doc = Document()

    # 段落1：字号错误（14pt），字体为继承默认（None → Calibri），不符合期望 宋体 12pt
    p1 = doc.add_paragraph()
    r1 = p1.add_run('正文段落字体错误示例文本。')
    r1.font.size = Pt(14)  # 错字号；family 为 None（继承）

    # 段落2：字体名明确写为黑体，不符合期望
    p2 = doc.add_paragraph()
    r2 = p2.add_run('另一段显式设为黑体的正文内容。')
    r2.font.name = '黑体'
    r2.font.size = Pt(12)  # 字号正确，但字体错误

    # 段落3（标题）：标题段落不应被 font.body 检测到
    h = doc.add_paragraph('一级标题内容', style='Heading 1')
    h.runs[0].font.name = '黑体'
    h.runs[0].font.size = Pt(16)

    path = tmp_path / 'bad_font_body.docx'
    doc.save(path)
    return path


def _make_clean_font_body_doc(tmp_path: Path) -> Path:
    """创建完全合规的测试文档：宋体 12pt 正文"""
    doc = Document()

    p1 = doc.add_paragraph()
    r1 = p1.add_run('完全合规的正文内容，字体字号均正确。')
    r1.font.name = '宋体'
    r1.font.size = Pt(12)

    p2 = doc.add_paragraph()
    r2 = p2.add_run('第二段也是合规的正文。')
    r2.font.name = '宋体'
    r2.font.size = Pt(12)

    path = tmp_path / 'clean_font_body.docx'
    doc.save(path)
    return path


# ===========================================================================
# detect 测试
# ===========================================================================

class TestDetect:
    def test_detect_finds_wrong_font_in_body(self, tmp_path: Path):
        """期望宋体 12pt，文档包含 Calibri(None)/14pt 和 黑体/12pt → 检出两条 issue"""
        path = _make_bad_font_body_doc(tmp_path)
        doc = Document(path)
        value = {'family': '宋体', 'size_pt': 12}
        issues = FontBodyRule.detect(doc, value)
        # 段落0 (r0) 和 段落1 (r0) 都有问题；标题段落不计
        assert len(issues) == 2
        assert all(i.rule_id == 'font.body' for i in issues)

    def test_detect_skips_heading_paragraph(self, tmp_path: Path):
        """标题段落（Heading 1）必须被跳过，不产生 issue"""
        doc = Document()
        # 只有一个 Heading 1 段落
        h = doc.add_paragraph('仅有标题', style='Heading 1')
        h.runs[0].font.name = '黑体'
        h.runs[0].font.size = Pt(16)
        path = tmp_path / 'only_heading.docx'
        doc.save(path)
        doc2 = Document(path)
        issues = FontBodyRule.detect(doc2, {'family': '宋体', 'size_pt': 12})
        assert issues == []

    def test_detect_returns_empty_for_clean(self, tmp_path: Path):
        """完全合规文档 → 空列表"""
        path = _make_clean_font_body_doc(tmp_path)
        doc = Document(path)
        issues = FontBodyRule.detect(doc, {'family': '宋体', 'size_pt': 12})
        assert issues == []

    def test_detect_skips_when_family_missing(self, tmp_path: Path):
        """value 缺少 family → 静默返回 []（不完整 spec 不误报）"""
        path = _make_bad_font_body_doc(tmp_path)
        doc = Document(path)
        issues = FontBodyRule.detect(doc, {'size_pt': 12})  # 无 family
        assert issues == []

    def test_detect_skips_when_size_pt_missing(self, tmp_path: Path):
        """value 缺少 size_pt → 静默返回 []"""
        path = _make_bad_font_body_doc(tmp_path)
        doc = Document(path)
        issues = FontBodyRule.detect(doc, {'family': '宋体'})  # 无 size_pt
        assert issues == []

    def test_detect_snippet_length(self, tmp_path: Path):
        """snippet 不超过 21 字符（20 + 可能有省略号），context 不超过 31 字符"""
        # 创建一个有超长文本的段落确保截断逻辑被触发
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run('这是一段非常非常非常非常非常非常非常非常非常长的正文内容，字体为黑体。')
        r.font.name = '黑体'
        r.font.size = Pt(14)
        path = tmp_path / 'long_text.docx'
        doc.save(path)

        doc2 = Document(path)
        issues = FontBodyRule.detect(doc2, {'family': '宋体', 'size_pt': 12})
        assert len(issues) == 1
        issue = issues[0]
        # snippet 最多 20 字符正文 + 1 个省略号
        assert len(issue.snippet) <= 21
        # context 最多 30 字符正文 + 1 个省略号
        assert len(issue.context) <= 31

    def test_detect_snippet_no_ellipsis_for_short_text(self, tmp_path: Path):
        """短文本（<=20 字符）的 snippet 不加省略号"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run('短文本')  # 3 字符
        r.font.name = '黑体'
        r.font.size = Pt(12)
        path = tmp_path / 'short.docx'
        doc.save(path)

        doc2 = Document(path)
        issues = FontBodyRule.detect(doc2, {'family': '宋体', 'size_pt': 12})
        assert len(issues) == 1
        assert '…' not in issues[0].snippet
        assert issues[0].snippet == '短文本'

    def test_detect_bold_check_optional(self, tmp_path: Path):
        """value 含 bold=false，run 实际 bold=True → 额外报告 bold 问题"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run('加粗文本')
        r.font.name = '宋体'
        r.font.size = Pt(12)
        r.font.bold = True  # 不符合 bold=False 要求
        path = tmp_path / 'bold.docx'
        doc.save(path)

        doc2 = Document(path)
        issues = FontBodyRule.detect(doc2, {'family': '宋体', 'size_pt': 12, 'bold': False})
        assert len(issues) == 1
        assert 'bold' in issues[0].message


# ===========================================================================
# fix 测试
# ===========================================================================

class TestFix:
    def test_fix_and_reopen_no_issue(self, tmp_path: Path):
        """修复后保存并重新打开文档，detect 必须返回空列表（硬验证）"""
        path = _make_bad_font_body_doc(tmp_path)
        doc = Document(path)
        value = {'family': '宋体', 'size_pt': 12}
        issues = FontBodyRule.detect(doc, value)
        assert len(issues) > 0

        # 逐条修复
        for issue in issues:
            FontBodyRule.fix(doc, issue, value)

        fixed_path = tmp_path / 'fixed.docx'
        doc.save(fixed_path)

        # 重新打开验证
        doc2 = Document(fixed_path)
        remaining = FontBodyRule.detect(doc2, value)
        assert remaining == [], f'修复后仍有 {len(remaining)} 条 issue: {remaining}'

    def test_fix_marks_blue(self, tmp_path: Path):
        """修复后 run 的颜色必须变为蓝色 #0070C0"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run('需修复的正文文本')
        r.font.name = '黑体'
        r.font.size = Pt(14)
        path = tmp_path / 'for_fix.docx'
        doc.save(path)

        doc2 = Document(path)
        value = {'family': '宋体', 'size_pt': 12}
        issues = FontBodyRule.detect(doc2, value)
        assert len(issues) == 1

        FontBodyRule.fix(doc2, issues[0], value)

        # 验证颜色
        target_run = doc2.paragraphs[issues[0].loc.para].runs[issues[0].loc.run]
        color = target_run.font.color.rgb
        assert color == RGBColor(0x00, 0x70, 0xC0), f'颜色应为 #0070C0，实际为 {color}'

    def test_fix_handles_bold_optional(self, tmp_path: Path):
        """value 含 bold=False，修复后 run.font.bold 必须为 False"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run('加粗需修复文本')
        r.font.name = '黑体'
        r.font.size = Pt(14)
        r.font.bold = True
        path = tmp_path / 'bold_fix.docx'
        doc.save(path)

        doc2 = Document(path)
        value = {'family': '宋体', 'size_pt': 12, 'bold': False}
        issues = FontBodyRule.detect(doc2, value)
        assert len(issues) == 1

        result = FontBodyRule.fix(doc2, issues[0], value)
        assert result.applied is True

        target_run = doc2.paragraphs[issues[0].loc.para].runs[issues[0].loc.run]
        assert target_run.font.bold is False, 'fix 后 bold 应为 False'

    def test_fix_returns_diff_string(self, tmp_path: Path):
        """fix 返回的 FixResult.diff 包含修复前后字体信息"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run('测试 diff 输出')
        r.font.name = '黑体'
        r.font.size = Pt(14)
        path = tmp_path / 'diff_test.docx'
        doc.save(path)

        doc2 = Document(path)
        value = {'family': '宋体', 'size_pt': 12}
        issues = FontBodyRule.detect(doc2, value)
        assert len(issues) == 1

        result = FontBodyRule.fix(doc2, issues[0], value)
        assert result.applied is True
        assert '宋体' in result.diff
        assert '12pt' in result.diff
        assert result.xml_changed  # 非空列表
