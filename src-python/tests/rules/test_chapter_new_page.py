"""
@file: test_chapter_new_page.py
@description: chapter.new_page 规则的单元测试
              fixture 全部通过 tmp_path 在运行时生成（不依赖静态 .docx 文件）
@author: Atlas.oi
@date: 2026-04-18
"""
import pytest
from pathlib import Path
from docx import Document
from docx.shared import RGBColor

from thesis_worker.rules.chapter_new_page import ChapterNewPageRule
from thesis_worker.models import Issue, Location


# ===========================================================================
# fixture 工厂函数（tmp_path 策略，CI 自包含）
# ===========================================================================

def _make_multi_chapter_doc(tmp_path: Path) -> Path:
    """
    创建含多章的测试文档：
    - 第一章：文档开头，不需要分页（合规）
    - 第二章：缺少 page_break_before（违规）
    - 第三章：已设置 page_break_before（合规）
    """
    doc = Document()
    # 第一个 Heading 1，文档开头，无需前置分页
    doc.add_paragraph('第一章 引言', style='Heading 1')
    doc.add_paragraph('正文段落，第一章内容。')

    # 第二个 Heading 1，缺少分页符 → 违规
    doc.add_paragraph('第二章 方法', style='Heading 1')
    doc.add_paragraph('正文段落，第二章内容。')

    # 第三个 Heading 1，已设置分页符 → 合规
    ch3 = doc.add_paragraph('第三章 实验', style='Heading 1')
    ch3.paragraph_format.page_break_before = True

    path = tmp_path / 'multi_chapter.docx'
    doc.save(path)
    return path


# ===========================================================================
# detect 测试
# ===========================================================================

class TestDetect:
    def test_detect_finds_h1_without_page_break(self, tmp_path: Path):
        """第二章缺分页符 → 检出 1 条 issue，第一章和第三章不违规"""
        path = _make_multi_chapter_doc(tmp_path)
        doc = Document(path)
        issues = ChapterNewPageRule.detect(doc, True)

        # 只有第二章违规（缺 page_break_before）
        assert len(issues) == 1
        assert issues[0].rule_id == 'chapter.new_page'
        assert issues[0].current is False
        assert issues[0].expected is True

    def test_detect_skips_first_h1(self, tmp_path: Path):
        """文档第一个段落是 Heading 1 时，不应报告违规"""
        doc = Document()
        # 第一个段落是 Heading 1，索引 0，应跳过
        doc.add_paragraph('第一章 只有一章', style='Heading 1')
        doc.add_paragraph('正文内容。')
        path = tmp_path / 'single_chapter.docx'
        doc.save(path)

        doc2 = Document(path)
        issues = ChapterNewPageRule.detect(doc2, True)
        assert issues == []

    def test_detect_skips_non_h1(self, tmp_path: Path):
        """Heading 2/3 和正文段落不应被检测"""
        doc = Document()
        doc.add_paragraph('第一章', style='Heading 1')
        doc.add_paragraph('第一节', style='Heading 2')  # 不检测
        doc.add_paragraph('第一小节', style='Heading 3')  # 不检测
        doc.add_paragraph('正文内容。')
        path = tmp_path / 'h2_h3_only.docx'
        doc.save(path)

        doc2 = Document(path)
        issues = ChapterNewPageRule.detect(doc2, True)
        # Heading 2/3 不触发 chapter.new_page 规则
        assert issues == []

    def test_detect_returns_empty_when_value_false(self, tmp_path: Path):
        """value=False 时不要求章节分页，直接返回 []"""
        path = _make_multi_chapter_doc(tmp_path)
        doc = Document(path)
        issues = ChapterNewPageRule.detect(doc, False)
        assert issues == []

    def test_detect_returns_empty_when_all_h1_have_break(self, tmp_path: Path):
        """所有非首个 Heading 1 都设了 page_break_before → []"""
        doc = Document()
        doc.add_paragraph('第一章 引言', style='Heading 1')
        doc.add_paragraph('正文内容。')

        # 第二章设有分页符
        ch2 = doc.add_paragraph('第二章 方法', style='Heading 1')
        ch2.paragraph_format.page_break_before = True

        # 第三章设有分页符
        ch3 = doc.add_paragraph('第三章 实验', style='Heading 1')
        ch3.paragraph_format.page_break_before = True

        path = tmp_path / 'all_have_break.docx'
        doc.save(path)

        doc2 = Document(path)
        issues = ChapterNewPageRule.detect(doc2, True)
        assert issues == []


# ===========================================================================
# fix 测试
# ===========================================================================

class TestFix:
    def test_fix_and_reopen_no_issue(self, tmp_path: Path):
        """修复后保存并重新打开文档，detect 必须返回空列表（硬验证）"""
        path = _make_multi_chapter_doc(tmp_path)
        doc = Document(path)
        issues = ChapterNewPageRule.detect(doc, True)
        assert len(issues) > 0

        # 逐条修复
        for issue in issues:
            ChapterNewPageRule.fix(doc, issue, True)

        fixed_path = tmp_path / 'fixed_chapter.docx'
        doc.save(fixed_path)

        # 重新打开验证
        doc2 = Document(fixed_path)
        remaining = ChapterNewPageRule.detect(doc2, True)
        assert remaining == [], f'修复后仍有 {len(remaining)} 条 issue: {remaining}'

    def test_fix_marks_blue(self, tmp_path: Path):
        """修复后段落第一个 run 的颜色必须变为蓝色 #0070C0"""
        doc = Document()
        doc.add_paragraph('第一章 引言', style='Heading 1')
        # 第二章缺分页，且有 run 文本（需要蓝色标记）
        doc.add_paragraph('第二章 方法', style='Heading 1')
        path = tmp_path / 'for_fix_blue.docx'
        doc.save(path)

        doc2 = Document(path)
        issues = ChapterNewPageRule.detect(doc2, True)
        assert len(issues) == 1

        ChapterNewPageRule.fix(doc2, issues[0], True)

        target_para = doc2.paragraphs[issues[0].loc.para]
        assert target_para.runs, '目标段落应有 run 用于蓝色标记'
        color = target_para.runs[0].font.color.rgb
        assert color == RGBColor(0x00, 0x70, 0xC0), f'颜色应为 #0070C0，实际为 {color}'
