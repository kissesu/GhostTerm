"""
@file: test_figure_table_caption.py
@description: figure.caption_pos + table.caption_pos 规则检测测试
              测试题注位置的正确性判断（只检测，不 fix）
@author: Atlas.oi
@date: 2026-04-18
"""
import base64
import io

import pytest
from docx import Document
from docx.shared import Inches
from lxml import etree

from thesis_worker.rules.figure_table_caption import (
    FigureCaptionPosRule,
    TableCaptionPosRule,
)

# 1x1 透明 PNG，用于构造最小内联图片
_PNG_1X1 = base64.b64decode(
    'iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNkYPhfDwAChwGA60e6kgAAAABJRU5ErkJggg=='
)


def _add_figure_para(doc: Document) -> None:
    """向 doc 追加含内联图片的段落（代表图）"""
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_picture(io.BytesIO(_PNG_1X1), width=Inches(0.1))


def _add_caption_para(doc: Document, text: str = '图 1 测试图') -> None:
    """向 doc 追加 Caption 样式段落"""
    doc.add_paragraph(text, style='Caption')


def _add_table(doc: Document) -> None:
    """向 doc 追加单行单列表格（代表表）"""
    doc.add_table(rows=1, cols=1)


class TestFigureCaptionPos:
    def test_caption_below_figure_when_expected_above_detected(self, tmp_path):
        """
        图在上，题注在下（below），但期望 above → 应检出 issue。
        文档顺序：[图段落] → [Caption 段落]
        """
        doc = Document()
        _add_figure_para(doc)
        _add_caption_para(doc, '图 1 测试图')

        issues = FigureCaptionPosRule.detect(doc, 'above')

        assert len(issues) == 1
        assert issues[0].rule_id == 'figure.caption_pos'
        assert issues[0].current == 'below'
        assert issues[0].expected == 'above'
        assert issues[0].fix_available is False
        assert '图 1 测试图' in issues[0].snippet

    def test_caption_above_figure_when_expected_above_clean(self, tmp_path):
        """
        题注在上（above），图在下，期望 above → 合规，应返回空列表。
        文档顺序：[Caption 段落] → [图段落]
        """
        doc = Document()
        _add_caption_para(doc, '图 1 测试图')
        _add_figure_para(doc)

        issues = FigureCaptionPosRule.detect(doc, 'above')

        assert issues == []

    def test_caption_below_figure_when_expected_below_clean(self, tmp_path):
        """
        题注在下（below），期望 below → 合规，应返回空列表。
        文档顺序：[图段落] → [Caption 段落]
        """
        doc = Document()
        _add_figure_para(doc)
        _add_caption_para(doc, '图 2 测试图')

        issues = FigureCaptionPosRule.detect(doc, 'below')

        assert issues == []

    def test_skip_when_value_invalid(self, tmp_path):
        """value='middle' 不在允许值内 → 直接返回空列表"""
        doc = Document()
        _add_figure_para(doc)
        _add_caption_para(doc)

        issues = FigureCaptionPosRule.detect(doc, 'middle')

        assert issues == []


class TestTableCaptionPos:
    def test_caption_below_table_when_expected_above_detected(self, tmp_path):
        """
        表在上，题注在下（below），但期望 above → 应检出 issue。
        文档顺序：[表格] → [Caption 段落]
        """
        doc = Document()
        _add_table(doc)
        _add_caption_para(doc, '表 1 测试表')

        issues = TableCaptionPosRule.detect(doc, 'above')

        assert len(issues) == 1
        assert issues[0].rule_id == 'table.caption_pos'
        assert issues[0].current == 'below'
        assert issues[0].expected == 'above'
        assert issues[0].fix_available is False

    def test_caption_above_table_when_expected_above_clean(self, tmp_path):
        """
        题注在上（above），表在下，期望 above → 合规，应返回空列表。
        文档顺序：[Caption 段落] → [表格]
        """
        doc = Document()
        _add_caption_para(doc, '表 1 测试表')
        _add_table(doc)

        issues = TableCaptionPosRule.detect(doc, 'above')

        assert issues == []


class TestFixUnavailable:
    def test_fix_raises_not_implemented(self, tmp_path):
        """fix 必须 raise NotImplementedError，确认 fix_available=False 语义一致"""
        doc = Document()
        _add_figure_para(doc)
        _add_caption_para(doc, '图 1 测试图')

        issues = FigureCaptionPosRule.detect(doc, 'above')
        assert len(issues) == 1

        with pytest.raises(NotImplementedError):
            FigureCaptionPosRule.fix(doc, issues[0], 'above')

    def test_table_fix_raises_not_implemented(self, tmp_path):
        """TableCaptionPosRule.fix 同样 raise NotImplementedError"""
        doc = Document()
        _add_table(doc)
        _add_caption_para(doc, '表 1 测试表')

        issues = TableCaptionPosRule.detect(doc, 'above')
        assert len(issues) == 1

        with pytest.raises(NotImplementedError):
            TableCaptionPosRule.fix(doc, issues[0], 'above')
