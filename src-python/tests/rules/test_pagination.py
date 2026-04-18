"""
@file: test_pagination.py
@description: pagination 规则的单元测试。
              fixture 全部通过 tmp_path 在运行时生成（不依赖静态 .docx 文件）。
@author: Atlas.oi
@date: 2026-04-18
"""
import pytest
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

from thesis_worker.rules.pagination import PaginationRule, _footer_has_page_field
from thesis_worker.models import Issue


# ===========================================================================
# 辅助：向 footer 插入 PAGE field（fldSimple 写法）
# ===========================================================================

def _add_page_field_to_footer(section) -> None:
    """在 section 的 footer 第一段插入 w:fldSimple PAGE field"""
    footer = section.footer
    p = footer.paragraphs[0]
    fld = etree.SubElement(p._element, qn('w:fldSimple'))
    fld.set(qn('w:instr'), ' PAGE ')


# ===========================================================================
# fixture 工厂函数
# ===========================================================================

_VALUE = {'front_matter': 'roman', 'body': 'arabic'}


def _make_doc_no_page_field() -> Document:
    """创建单 section 的文档，footer 无 PAGE field（默认空 footer）"""
    doc = Document()
    doc.add_paragraph('正文内容')
    return doc


def _make_doc_with_page_field() -> Document:
    """创建单 section 的文档，footer 含 PAGE field"""
    doc = Document()
    doc.add_paragraph('正文内容')
    _add_page_field_to_footer(doc.sections[0])
    return doc


def _make_doc_multi_section_partial() -> Document:
    """
    创建两个 section 的文档：
    - section[0]：无 PAGE field（违规）
    - section[1]：有 PAGE field（合规）

    python-docx 新分节符的 footer 默认 is_linked_to_previous=True，
    必须先 unlink 才能让两个 section 拥有独立的 footer XML。
    """
    from docx.oxml import OxmlElement

    doc = Document()
    doc.add_paragraph('第一节正文')

    # 插入连续分节符，产生第二个 section
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    sectPr = OxmlElement('w:sectPr')
    pPr.append(sectPr)

    doc.add_paragraph('第二节正文')

    # unlink section[1] footer，使其拥有独立 XML，再插入 PAGE field
    doc.sections[1].footer.is_linked_to_previous = False
    _add_page_field_to_footer(doc.sections[1])

    return doc


# ===========================================================================
# 测试类
# ===========================================================================

class TestFooterHasPageField:
    """_footer_has_page_field 辅助函数的独立测试"""

    def test_returns_false_when_footer_empty(self):
        doc = _make_doc_no_page_field()
        assert _footer_has_page_field(doc.sections[0].footer) is False

    def test_returns_true_when_fld_simple_present(self):
        doc = _make_doc_with_page_field()
        assert _footer_has_page_field(doc.sections[0].footer) is True

    def test_returns_true_when_instr_text_present(self):
        """验证 w:instrText 复杂域写法也能被检测到"""
        from docx.oxml import OxmlElement

        doc = Document()
        section = doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0]

        # 构造 w:r > w:instrText 结构（复杂域写法）
        r = OxmlElement('w:r')
        instr = OxmlElement('w:instrText')
        instr.text = ' PAGE '
        r.append(instr)
        p._element.append(r)

        assert _footer_has_page_field(footer) is True


class TestDetect:
    def test_detect_finds_missing_page_field(self):
        """footer 无 PAGE field → 检出 1 条 issue"""
        doc = _make_doc_no_page_field()
        issues = PaginationRule.detect(doc, _VALUE)

        assert len(issues) == 1
        issue = issues[0]
        assert issue.rule_id == 'pagination'
        assert issue.fix_available is False
        assert 'section[0]' in issue.snippet
        assert 'no PAGE field' == issue.current

    def test_detect_returns_empty_when_page_field_present(self):
        """footer 有 PAGE field → 返回空列表"""
        doc = _make_doc_with_page_field()
        issues = PaginationRule.detect(doc, _VALUE)

        assert issues == []

    def test_detect_skips_when_value_is_none(self):
        """value=None → 跳过检测，返回空列表"""
        doc = _make_doc_no_page_field()
        assert PaginationRule.detect(doc, None) == []

    def test_detect_skips_when_value_missing_keys(self):
        """value 缺少 front_matter 或 body → 跳过检测"""
        doc = _make_doc_no_page_field()

        # 缺 body
        assert PaginationRule.detect(doc, {'front_matter': 'roman'}) == []
        # 缺 front_matter
        assert PaginationRule.detect(doc, {'body': 'arabic'}) == []
        # 空 dict
        assert PaginationRule.detect(doc, {}) == []

    def test_detect_one_issue_per_section_without_page(self):
        """多 section，只有部分缺 PAGE field → 检出对应数量 issue"""
        doc = _make_doc_multi_section_partial()
        issues = PaginationRule.detect(doc, _VALUE)

        # section[0] 无 PAGE field → 1 条 issue
        assert len(issues) == 1
        assert issues[0].loc.para == 0

    def test_detect_issue_message_contains_section_number(self):
        """issue.message 应包含人类可读的 section 序号（从 1 开始）"""
        doc = _make_doc_no_page_field()
        issues = PaginationRule.detect(doc, _VALUE)

        assert '1' in issues[0].message

    def test_detect_context_shows_total_sections(self):
        """issue.context 应包含文档 section 总数"""
        doc = _make_doc_no_page_field()
        issues = PaginationRule.detect(doc, _VALUE)

        assert 'sections' in issues[0].context


class TestFixUnavailable:
    def test_fix_raises_not_implemented(self):
        """pagination 不支持自动修复，fix 必须抛 NotImplementedError"""
        doc = _make_doc_no_page_field()
        issues = PaginationRule.detect(doc, _VALUE)
        assert len(issues) == 1

        with pytest.raises(NotImplementedError):
            PaginationRule.fix(doc, issues[0], _VALUE)


class TestRuleMetadata:
    def test_rule_attributes(self):
        assert PaginationRule.id == 'pagination'
        assert PaginationRule.category == 'structure'
        assert PaginationRule.severity == 'warning'
        assert PaginationRule.fix_available is False
