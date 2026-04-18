"""
@file: test_ai_pattern.py
@description: ai_pattern.check 规则 detect 覆盖测试（只检测，无 fix 用例）
@author: Atlas.oi
@date: 2026-04-18
"""
import pytest
from docx import Document

from thesis_worker.rules.ai_pattern import AiPatternCheckRule

# thesis-default ruleset 的标准 value
_VALUE = {'ruleset': 'thesis-default'}


def _make_doc(paragraphs: list[str]) -> Document:
    """用指定段落文本创建内存中的 Document（每段一个 run）"""
    doc = Document()
    # 移除默认空段落
    for para in doc.paragraphs:
        p = para._element
        p.getparent().remove(p)
    for text in paragraphs:
        doc.add_paragraph(text)
    return doc


class TestDetect:
    def test_detect_finds_all_markers(self, tmp_path):
        """段落同时含 '综上所述' 和 '值得注意的是' → 两个 issue"""
        doc = _make_doc(['综上所述，值得注意的是本研究存在局限性。'])
        issues = AiPatternCheckRule.detect(doc, _VALUE)
        rule_ids = {i.rule_id for i in issues}
        assert rule_ids == {'ai_pattern.check'}
        # 找到 '综上所述' 和 '值得注意的是' 各一个 issue
        markers_found = {i.current for i in issues}
        assert '综上所述' in markers_found
        assert '值得注意的是' in markers_found
        assert len(issues) >= 2

    def test_detect_returns_empty_for_clean_text(self, tmp_path):
        """不含 marker 的段落 → 空列表"""
        doc = _make_doc(['本章节介绍实验方法及数据分析过程。'])
        issues = AiPatternCheckRule.detect(doc, _VALUE)
        assert issues == []

    def test_detect_skips_when_ruleset_invalid(self, tmp_path):
        """ruleset 不是 thesis-default → 跳过，返回空列表"""
        doc = _make_doc(['综上所述，本研究完成了目标。'])
        issues = AiPatternCheckRule.detect(doc, {'ruleset': 'custom-ruleset'})
        assert issues == []

    def test_detect_skips_when_value_not_dict(self, tmp_path):
        """value 不是 dict（如 None）→ 跳过，返回空列表"""
        doc = _make_doc(['综上所述，实验证明了假设。'])
        issues = AiPatternCheckRule.detect(doc, None)
        assert issues == []

    def test_detect_handles_multiple_occurrences_same_marker(self, tmp_path):
        """同一 run 中 '值得注意的是' 出现两次 → 两个独立 issue"""
        doc = _make_doc(['值得注意的是前者；同样值得注意的是后者。'])
        issues = AiPatternCheckRule.detect(doc, _VALUE)
        target = [i for i in issues if i.current == '值得注意的是']
        assert len(target) == 2
        # 两个 issue 的 char 偏移不同
        chars = {i.loc.char for i in target}
        assert len(chars) == 2

    def test_detect_issue_fields(self, tmp_path):
        """验证 issue 字段：fix_available=False，snippet 非空，context 非空"""
        doc = _make_doc(['毋庸置疑，这是正确的结论。'])
        issues = AiPatternCheckRule.detect(doc, _VALUE)
        assert len(issues) >= 1
        issue = next(i for i in issues if i.current == '毋庸置疑')
        assert issue.fix_available is False
        assert '毋庸置疑' in issue.snippet
        assert issue.context != ''
        assert issue.message == 'AI 化 marker 疑似：毋庸置疑'


class TestFixUnavailable:
    def test_fix_raises_not_implemented(self):
        """fix 方法必须抛出 NotImplementedError"""
        doc = _make_doc(['综上所述。'])
        issues = AiPatternCheckRule.detect(doc, _VALUE)
        assert issues  # 至少有一个 issue
        with pytest.raises(NotImplementedError):
            AiPatternCheckRule.fix(doc, issues[0], _VALUE)
