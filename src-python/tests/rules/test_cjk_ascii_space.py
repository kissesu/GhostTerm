"""
@file: test_cjk_ascii_space.py
@description: cjk_ascii_space 规则 detect + fix + reopen 硬测试
@author: Atlas.oi
@date: 2026-04-17
"""
import tempfile
import shutil
from pathlib import Path
from docx import Document

from thesis_worker.rules.cjk_ascii_space import CjkAsciiSpaceRule

FIXTURES = Path(__file__).parent.parent / 'fixtures'
CONFIG_FORBID = {'allowed': False}  # 院校要求：不允许空格


class TestDetect:
    def test_detect_finds_violations_in_bad_doc(self):
        doc = Document(FIXTURES / 'cjk_space_bad.docx')
        issues = CjkAsciiSpaceRule.detect(doc, CONFIG_FORBID)
        # "这是 AI 工具" 有两处违规（中-空-英 + 英-空-中）
        # "版本 2.1.1 已发布" 有两处违规（中-空-数 + 数-空-中）
        # "无违规段落。" 0 处
        assert len(issues) == 4
        for i in issues:
            assert i.rule_id == 'cjk_ascii_space'
            assert i.fix_available is True

    def test_detect_returns_empty_on_clean_doc(self):
        doc = Document(FIXTURES / 'cjk_space_clean.docx')
        issues = CjkAsciiSpaceRule.detect(doc, CONFIG_FORBID)
        assert issues == []

    def test_detect_skip_when_allowed_true(self):
        """config 设为 allowed:true（允许空格）→ 不检测"""
        doc = Document(FIXTURES / 'cjk_space_bad.docx')
        issues = CjkAsciiSpaceRule.detect(doc, {'allowed': True})
        assert issues == []

    def test_detect_snippet_and_context_fields(self):
        """snippet 扩展 + context 段落预览，用户可直接在 WPS 里定位违规"""
        doc = Document(FIXTURES / 'cjk_space_bad.docx')
        issues = CjkAsciiSpaceRule.detect(doc, CONFIG_FORBID)
        assert len(issues) == 4

        # 段落 0: "这是 AI 工具。" → 两处违规，snippet 各扩到 "这是 AI" 和 "AI 工具"
        assert issues[0].snippet == '这是 AI'
        assert issues[1].snippet == 'AI 工具'
        assert issues[0].context == '这是 AI 工具。'
        assert issues[1].context == '这是 AI 工具。'

        # 段落 1: "版本 2.1.1 已发布。" → . 不是 ASCII 停止符，数字串整块保留
        assert issues[2].snippet == '版本 2.1.1'
        assert issues[3].snippet == '2.1.1 已发布'
        assert '版本 2.1.1' in issues[2].context


class TestSnippetExpansion:
    """_expand_snippet 的边界单元测试，不依赖 docx fixture"""

    def test_cjk_side_stops_at_non_cjk(self):
        from thesis_worker.rules.cjk_ascii_space import _expand_snippet
        text = '使用 AI 完成'
        # 违规 "用 A" 位置
        start = text.index('用 A')
        assert _expand_snippet(text, start, start + 3) == '使用 AI'

    def test_ascii_side_keeps_dots_in_version(self):
        from thesis_worker.rules.cjk_ascii_space import _expand_snippet
        text = '版本 2.1.1 已发布'
        start = text.index('本 2')
        assert _expand_snippet(text, start, start + 3) == '版本 2.1.1'

    def test_ascii_side_keeps_paren_inside_token(self):
        from thesis_worker.rules.cjk_ascii_space import _expand_snippet
        text = '使用 (AI) 工具'
        start = text.index('用 (')
        # 右侧吃 "(AI)" 直到遇到空格停，不把中文标点吸入
        assert _expand_snippet(text, start, start + 3) == '使用 (AI)'

    def test_cjk_punct_stops_ascii_expansion(self):
        from thesis_worker.rules.cjk_ascii_space import _expand_snippet
        # 中文句号应停止 ASCII 侧扩展
        text = '通过 submit。下一步'
        start = text.index('过 s')
        assert _expand_snippet(text, start, start + 3) == '通过 submit'


class TestFix:
    def test_fix_and_reopen_produces_no_issues(self, tmp_path):
        """关键硬测试：修复后重开跑 detect 必须返回空"""
        origin = FIXTURES / 'cjk_space_bad.docx'
        tmp = tmp_path / 'copy.docx'
        shutil.copy(origin, tmp)

        doc = Document(tmp)
        issues = CjkAsciiSpaceRule.detect(doc, CONFIG_FORBID)
        assert len(issues) == 4

        # 逐条修复
        for issue in issues:
            CjkAsciiSpaceRule.fix(doc, issue, CONFIG_FORBID)
        doc.save(tmp)

        # 重开验证
        reopened = Document(tmp)
        remaining = CjkAsciiSpaceRule.detect(reopened, CONFIG_FORBID)
        assert remaining == []

    def test_fix_marks_blue_color(self, tmp_path):
        from docx.shared import RGBColor
        origin = FIXTURES / 'cjk_space_bad.docx'
        tmp = tmp_path / 'copy.docx'
        shutil.copy(origin, tmp)

        doc = Document(tmp)
        issues = CjkAsciiSpaceRule.detect(doc, CONFIG_FORBID)
        assert issues  # 非空

        CjkAsciiSpaceRule.fix(doc, issues[0], CONFIG_FORBID)
        doc.save(tmp)

        reopened = Document(tmp)
        # 第 0 段第 0 run（修改发生处）应有蓝色标记 #0070C0
        first_run = reopened.paragraphs[0].runs[0]
        assert first_run.font.color.rgb == RGBColor(0x00, 0x70, 0xC0)
