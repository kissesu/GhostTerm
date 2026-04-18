"""
@file: test_paragraph_indent.py
@description: paragraph.indent 规则的单元测试
              fixture 全部通过 tmp_path 在运行时生成（不依赖静态 .docx 文件）
@author: Atlas.oi
@date: 2026-04-18
"""
import pytest
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor

from thesis_worker.rules.paragraph_indent import ParagraphIndentRule
from thesis_worker.models import Issue, Location


# ===========================================================================
# fixture 工厂函数（tmp_path 策略，CI 自包含）
# ===========================================================================

# 默认测试用 value 配置：2字缩进，12pt正文
_DEFAULT_VALUE = {'first_line_chars': 2, 'body_font_size_pt': 12}
# 期望缩进 = 2 * 12 = 24pt
_EXPECTED_PT = 24.0


def _save(doc: Document, tmp_path: Path, name: str) -> Path:
    """保存 docx 并返回路径"""
    path = tmp_path / name
    doc.save(path)
    return path


def _make_doc_no_indent(tmp_path: Path) -> Path:
    """无首行缩进的正文段落"""
    doc = Document()
    p = doc.add_paragraph()
    p.add_run('这是一段没有首行缩进的正文内容，应当被检测出来。')
    # first_line_indent 默认为 None（无缩进）
    return _save(doc, tmp_path, 'no_indent.docx')


def _make_doc_wrong_indent(tmp_path: Path) -> Path:
    """首行缩进 = 12pt（1字），期望 24pt，应被检测"""
    doc = Document()
    p = doc.add_paragraph()
    p.add_run('这是一段缩进只有一个字宽的正文，缩进不足应被检出。')
    p.paragraph_format.first_line_indent = Pt(12)  # 1字缩进，期望2字
    return _save(doc, tmp_path, 'wrong_indent.docx')


def _make_doc_correct_indent(tmp_path: Path) -> Path:
    """首行缩进 = 24pt（2字），符合要求"""
    doc = Document()
    p = doc.add_paragraph()
    p.add_run('这是一段首行缩进完全正确的正文段落内容。')
    p.paragraph_format.first_line_indent = Pt(24)
    return _save(doc, tmp_path, 'correct_indent.docx')


def _make_doc_heading_no_indent(tmp_path: Path) -> Path:
    """Heading 1 段落无缩进，不应被检测"""
    doc = Document()
    # 标题段落：无缩进，但应跳过
    doc.add_paragraph('第一章 绪论', style='Heading 1')
    # 也加一个正确缩进的正文，确保规则正常运行
    p = doc.add_paragraph()
    p.add_run('正文内容，缩进正确。')
    p.paragraph_format.first_line_indent = Pt(24)
    return _save(doc, tmp_path, 'heading_no_indent.docx')


def _make_doc_empty_paragraph(tmp_path: Path) -> Path:
    """空段落不应被检测"""
    doc = Document()
    # 空段落：无文本，无缩进
    doc.add_paragraph('')
    # 再加一个合规的正文
    p = doc.add_paragraph()
    p.add_run('合规正文。')
    p.paragraph_format.first_line_indent = Pt(24)
    return _save(doc, tmp_path, 'empty_para.docx')


def _make_doc_tolerance(tmp_path: Path) -> Path:
    """首行缩进 23.5pt，期望 24pt，误差 ≤1pt，应通过"""
    doc = Document()
    p = doc.add_paragraph()
    p.add_run('缩进 23.5pt，在容差范围内，不应被检出。')
    # python-docx Pt() 转 EMU 再转回可能有微小误差；直接赋值 EMU
    # Pt(23.5) = 23.5 * 12700 EMU
    p.paragraph_format.first_line_indent = Pt(23.5)
    return _save(doc, tmp_path, 'tolerance.docx')


def _make_doc_for_fix(tmp_path: Path) -> Path:
    """无缩进段落（含至少一个 run），用于 fix 测试"""
    doc = Document()
    p = doc.add_paragraph()
    p.add_run('这段文字需要被修复添加首行缩进。')
    return _save(doc, tmp_path, 'fix_target.docx')


# ===========================================================================
# 检测测试
# ===========================================================================

class TestDetect:
    def test_detect_finds_missing_indent(self, tmp_path):
        """无缩进段落 → 检出一个 issue"""
        path = _make_doc_no_indent(tmp_path)
        doc = Document(path)
        issues = ParagraphIndentRule.detect(doc, _DEFAULT_VALUE)
        assert len(issues) == 1
        issue = issues[0]
        assert issue.rule_id == 'paragraph.indent'
        assert issue.loc.run == 0
        assert issue.current['first_line_indent_pt'] is None
        assert issue.expected['first_line_indent_pt'] == _EXPECTED_PT
        assert issue.fix_available is True
        # snippet 和 context 不为空
        assert issue.snippet
        assert issue.context

    def test_detect_finds_wrong_indent(self, tmp_path):
        """缩进 12pt（1字）→ 检出，期望 24pt"""
        path = _make_doc_wrong_indent(tmp_path)
        doc = Document(path)
        issues = ParagraphIndentRule.detect(doc, _DEFAULT_VALUE)
        assert len(issues) == 1
        issue = issues[0]
        assert issue.current['first_line_indent_pt'] == pytest.approx(12.0, abs=0.1)
        assert issue.expected['first_line_indent_pt'] == _EXPECTED_PT

    def test_detect_returns_empty_for_correct_indent(self, tmp_path):
        """缩进 24pt → []"""
        path = _make_doc_correct_indent(tmp_path)
        doc = Document(path)
        issues = ParagraphIndentRule.detect(doc, _DEFAULT_VALUE)
        assert issues == []

    def test_detect_skips_heading(self, tmp_path):
        """Heading 1 段无缩进不被检出；正文缩进正确时结果为 []"""
        path = _make_doc_heading_no_indent(tmp_path)
        doc = Document(path)
        issues = ParagraphIndentRule.detect(doc, _DEFAULT_VALUE)
        # 标题不检测；正文已有正确缩进
        assert issues == []

    def test_detect_skips_empty_paragraph(self, tmp_path):
        """空段落不被检出；正文缩进正确时结果为 []"""
        path = _make_doc_empty_paragraph(tmp_path)
        doc = Document(path)
        issues = ParagraphIndentRule.detect(doc, _DEFAULT_VALUE)
        assert issues == []

    def test_detect_tolerance(self, tmp_path):
        """缩进 23.5pt，期望 24pt，误差 0.5pt ≤ 1pt → []"""
        path = _make_doc_tolerance(tmp_path)
        doc = Document(path)
        issues = ParagraphIndentRule.detect(doc, _DEFAULT_VALUE)
        assert issues == []

    def test_detect_skips_when_value_incomplete(self, tmp_path):
        """value 缺 first_line_chars → 静默返回 []"""
        path = _make_doc_no_indent(tmp_path)
        doc = Document(path)
        issues = ParagraphIndentRule.detect(doc, {'body_font_size_pt': 12})
        assert issues == []

    def test_detect_uses_default_body_size(self, tmp_path):
        """value 缺 body_font_size_pt → 默认 12pt，期望缩进仍为 24pt"""
        path = _make_doc_no_indent(tmp_path)
        doc = Document(path)
        issues = ParagraphIndentRule.detect(doc, {'first_line_chars': 2})
        assert len(issues) == 1
        assert issues[0].expected['first_line_indent_pt'] == _EXPECTED_PT

    def test_detect_issue_loc_run_is_zero(self, tmp_path):
        """一段一个 issue，loc.run 必须为 0"""
        path = _make_doc_no_indent(tmp_path)
        doc = Document(path)
        issues = ParagraphIndentRule.detect(doc, _DEFAULT_VALUE)
        assert issues[0].loc.run == 0


# ===========================================================================
# 修复测试
# ===========================================================================

class TestFix:
    def test_fix_and_reopen_no_issue(self, tmp_path):
        """fix 后保存重新打开，detect 应返回 []（reopen 硬测试）"""
        path = _make_doc_for_fix(tmp_path)
        doc = Document(path)

        # 先 detect 确认有问题
        issues = ParagraphIndentRule.detect(doc, _DEFAULT_VALUE)
        assert len(issues) == 1

        # fix
        result = ParagraphIndentRule.fix(doc, issues[0], _DEFAULT_VALUE)
        assert result.applied is True

        # 保存并重新打开
        fixed_path = tmp_path / 'fixed.docx'
        doc.save(fixed_path)
        doc2 = Document(fixed_path)
        issues2 = ParagraphIndentRule.detect(doc2, _DEFAULT_VALUE)
        assert issues2 == []

    def test_fix_marks_first_run_blue(self, tmp_path):
        """fix 后第一个 run 应被标记为蓝色 #0070C0"""
        path = _make_doc_for_fix(tmp_path)
        doc = Document(path)
        issues = ParagraphIndentRule.detect(doc, _DEFAULT_VALUE)
        assert len(issues) == 1

        ParagraphIndentRule.fix(doc, issues[0], _DEFAULT_VALUE)

        para = doc.paragraphs[issues[0].loc.para]
        assert para.runs, '段落应有 run'
        color = para.runs[0].font.color.rgb
        assert color == RGBColor(0x00, 0x70, 0xC0), f'期望蓝色 #0070C0，实际 {color}'

    def test_fix_sets_correct_indent(self, tmp_path):
        """fix 后 paragraph_format.first_line_indent 应等于 24pt"""
        path = _make_doc_for_fix(tmp_path)
        doc = Document(path)
        issues = ParagraphIndentRule.detect(doc, _DEFAULT_VALUE)
        ParagraphIndentRule.fix(doc, issues[0], _DEFAULT_VALUE)

        para = doc.paragraphs[issues[0].loc.para]
        actual_pt = para.paragraph_format.first_line_indent.pt
        assert abs(actual_pt - _EXPECTED_PT) <= 0.5, f'期望 {_EXPECTED_PT}pt，实际 {actual_pt}pt'

    def test_fix_diff_contains_expected_pt(self, tmp_path):
        """FixResult.diff 应包含期望磅值字样"""
        path = _make_doc_for_fix(tmp_path)
        doc = Document(path)
        issues = ParagraphIndentRule.detect(doc, _DEFAULT_VALUE)
        result = ParagraphIndentRule.fix(doc, issues[0], _DEFAULT_VALUE)
        assert '24' in result.diff
        assert result.xml_changed
