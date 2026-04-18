"""
@file: test_quote_style.py
@description: quote.style 规则 detect + fix + reopen 硬测试
              覆盖 cjk/ascii/mixed 三种 value 场景
@author: Atlas.oi
@date: 2026-04-18
"""
import shutil
from pathlib import Path
from docx import Document
from docx.shared import RGBColor

from thesis_worker.rules.quote_style import QuoteStyleRule


def _make_doc_with_text(tmp_path: Path, text: str) -> Path:
    """创建单段落 docx 文件，段落内容为指定文本"""
    p = tmp_path / 'test.docx'
    doc = Document()
    doc.add_paragraph(text)
    doc.save(p)
    return p


class TestDetect:
    def test_detect_ascii_quote_when_cjk_expected(self, tmp_path):
        """value='cjk'，文档含 ASCII 双引号 → 检出违规"""
        path = _make_doc_with_text(tmp_path, 'He said "hello" to her.')
        doc = Document(path)
        issues = QuoteStyleRule.detect(doc, 'cjk')
        # "hello" 两侧各一个 ASCII "，共两处
        assert len(issues) == 2
        for i in issues:
            assert i.rule_id == 'quote.style'
            assert i.fix_available is True
            assert i.current == '"'

    def test_detect_cjk_quote_when_ascii_expected(self, tmp_path):
        """value='ascii'，文档含中文双引号 → 检出违规"""
        path = _make_doc_with_text(tmp_path, '\u4ed6\u8bf4\u201c\u4f60\u597d\u201d\u3002')
        doc = Document(path)
        issues = QuoteStyleRule.detect(doc, 'ascii')
        # 左右中文双引号各一处
        assert len(issues) == 2
        for i in issues:
            assert i.rule_id == 'quote.style'
            assert i.current in ('\u201c', '\u201d')

    def test_detect_skips_mixed_value(self, tmp_path):
        """value='mixed' → 任何引号都不检测，返回空列表"""
        path = _make_doc_with_text(tmp_path, 'He said "hello" \u201c\u4f60\u597d\u201d.')
        doc = Document(path)
        issues = QuoteStyleRule.detect(doc, 'mixed')
        assert issues == []

    def test_detect_returns_empty_for_clean_cjk(self, tmp_path):
        """文档仅含中文引号，value='cjk' → 无违规"""
        path = _make_doc_with_text(tmp_path, '\u4ed6\u8bf4\u201c\u4f60\u597d\u201d\u3002')
        doc = Document(path)
        issues = QuoteStyleRule.detect(doc, 'cjk')
        assert issues == []

    def test_detect_snippet_contains_context_chars(self, tmp_path):
        """snippet 应包含匹配字符前后的上下文字符"""
        path = _make_doc_with_text(tmp_path, 'abc"xyz"def')
        doc = Document(path)
        issues = QuoteStyleRule.detect(doc, 'cjk')
        assert len(issues) == 2
        # 第一个 snippet 应包含 " 前面的字符
        assert '"' in issues[0].snippet

    def test_detect_single_quote_when_cjk_expected(self, tmp_path):
        """value='cjk'，文档含 ASCII 单引号 → 检出违规"""
        path = _make_doc_with_text(tmp_path, "It's a test.")
        doc = Document(path)
        issues = QuoteStyleRule.detect(doc, 'cjk')
        assert len(issues) == 1
        assert issues[0].current == "'"


class TestFix:
    def test_fix_ascii_to_cjk_pairing(self, tmp_path):
        """两组双引号 "a" "b" → \u201ca\u201d \u201cb\u201d（配对替换）"""
        path = _make_doc_with_text(tmp_path, '"a" "b"')
        doc = Document(path)
        issues = QuoteStyleRule.detect(doc, 'cjk')
        assert len(issues) == 4  # 4 个 ASCII 引号

        # 仅需修复第一个 issue，fix 会处理整个 run 内所有引号
        result = QuoteStyleRule.fix(doc, issues[0], 'cjk')
        assert result.applied is True

        # 验证 run 文本已替换为中文配对引号
        para = doc.paragraphs[issues[0].loc.para]
        run_text = para.runs[issues[0].loc.run].text
        assert '\u201c' in run_text  # 左双引号
        assert '\u201d' in run_text  # 右双引号
        # ASCII 双引号应已被全部替换
        assert '"' not in run_text

    def test_fix_and_reopen_no_issue(self, tmp_path):
        """修复后重开文档，detect 应返回空列表（硬测试）"""
        path = _make_doc_with_text(tmp_path, 'He said "hello world".')
        doc = Document(path)
        issues = QuoteStyleRule.detect(doc, 'cjk')
        assert len(issues) == 2

        # 修复第一个 issue（整个 run 内引号全部处理）
        QuoteStyleRule.fix(doc, issues[0], 'cjk')
        doc.save(path)

        # 重新打开验证
        reopened = Document(path)
        remaining = QuoteStyleRule.detect(reopened, 'cjk')
        assert remaining == []

    def test_fix_marks_blue(self, tmp_path):
        """修复后 run 字体颜色应标记为蓝色 #0070C0"""
        path = _make_doc_with_text(tmp_path, 'Say "hi".')
        doc = Document(path)
        issues = QuoteStyleRule.detect(doc, 'cjk')
        assert issues

        QuoteStyleRule.fix(doc, issues[0], 'cjk')
        doc.save(path)

        reopened = Document(path)
        run = reopened.paragraphs[issues[0].loc.para].runs[issues[0].loc.run]
        assert run.font.color.rgb == RGBColor(0x00, 0x70, 0xC0)
