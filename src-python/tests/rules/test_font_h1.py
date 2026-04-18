"""
@file: test_font_h1.py
@description: font.h1 规则的单元测试
              fixture 全部通过 tmp_path 在运行时生成（不依赖静态 .docx 文件）
@author: Atlas.oi
@date: 2026-04-18
"""
import pytest
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor

from thesis_worker.rules.font_h1 import FontH1Rule
from thesis_worker.models import Issue, Location


# ===========================================================================
# fixture 工厂函数（tmp_path 策略，CI 自包含）
# ===========================================================================

def _make_bad_h1_doc(tmp_path: Path) -> Path:
    """创建包含字体/字号不符一级标题的测试文档"""
    doc = Document()

    # Heading 1：字体 Calibri，字号 12pt，明显不符期望的 黑体 16pt
    h = doc.add_paragraph('一级标题测试内容', style='Heading 1')
    h.runs[0].font.name = 'Calibri'
    h.runs[0].font.size = Pt(12)

    # 正文段落：不应被 font.h1 检测到
    p = doc.add_paragraph()
    r = p.add_run('正文段落内容，字体随意。')
    r.font.name = 'Calibri'
    r.font.size = Pt(12)

    path = tmp_path / 'bad_h1.docx'
    doc.save(path)
    return path


def _make_clean_h1_doc(tmp_path: Path) -> Path:
    """创建完全合规的测试文档：黑体 16pt bold 一级标题"""
    doc = Document()

    h = doc.add_paragraph('合规一级标题', style='Heading 1')
    h.runs[0].font.name = '黑体'
    h.runs[0].font.size = Pt(16)
    h.runs[0].font.bold = True

    path = tmp_path / 'clean_h1.docx'
    doc.save(path)
    return path


# ===========================================================================
# detect 测试
# ===========================================================================

class TestDetect:
    def test_detect_finds_wrong_h1_font(self, tmp_path: Path):
        """Heading 1 用 Calibri 12pt → 检出 issue"""
        path = _make_bad_h1_doc(tmp_path)
        doc = Document(path)
        value = {'family': '黑体', 'size_pt': 16}
        issues = FontH1Rule.detect(doc, value)
        # 只有一个 Heading 1 段落有问题
        assert len(issues) == 1
        assert issues[0].rule_id == 'font.h1'

    def test_detect_skips_body_paragraph(self, tmp_path: Path):
        """正文段落（非 Heading 1）不应被 font.h1 检测到"""
        doc = Document()
        # 只有正文段落，没有 Heading 1
        p = doc.add_paragraph()
        r = p.add_run('正文内容，字体字号均不符合标题期望。')
        r.font.name = 'Calibri'
        r.font.size = Pt(12)
        path = tmp_path / 'body_only.docx'
        doc.save(path)

        doc2 = Document(path)
        issues = FontH1Rule.detect(doc2, {'family': '黑体', 'size_pt': 16})
        assert issues == []

    def test_detect_returns_empty_for_clean_h1(self, tmp_path: Path):
        """黑体 16pt bold 的 Heading 1 → 空列表"""
        path = _make_clean_h1_doc(tmp_path)
        doc = Document(path)
        issues = FontH1Rule.detect(doc, {'family': '黑体', 'size_pt': 16, 'bold': True})
        assert issues == []

    def test_detect_skips_when_value_incomplete(self, tmp_path: Path):
        """value 缺 family 或 size_pt → 静默返回 []（不完整 spec 不误报）"""
        path = _make_bad_h1_doc(tmp_path)
        doc = Document(path)
        # 缺 family
        assert FontH1Rule.detect(doc, {'size_pt': 16}) == []
        # 缺 size_pt
        assert FontH1Rule.detect(doc, {'family': '黑体'}) == []

    def test_detect_bold_mismatch(self, tmp_path: Path):
        """h1 字体字号正确但缺 bold → 检出包含 bold 的 issue"""
        doc = Document()
        h = doc.add_paragraph('标题加粗测试', style='Heading 1')
        h.runs[0].font.name = '黑体'
        h.runs[0].font.size = Pt(16)
        h.runs[0].font.bold = False  # 不符合 bold=True 要求
        path = tmp_path / 'bold_mismatch.docx'
        doc.save(path)

        doc2 = Document(path)
        issues = FontH1Rule.detect(doc2, {'family': '黑体', 'size_pt': 16, 'bold': True})
        assert len(issues) == 1
        assert 'bold' in issues[0].message


# ===========================================================================
# fix 测试
# ===========================================================================

class TestFix:
    def test_fix_and_reopen_no_issue(self, tmp_path: Path):
        """修复后保存并重新打开文档，detect 必须返回空列表（硬验证）"""
        path = _make_bad_h1_doc(tmp_path)
        doc = Document(path)
        value = {'family': '黑体', 'size_pt': 16}
        issues = FontH1Rule.detect(doc, value)
        assert len(issues) > 0

        # 逐条修复
        for issue in issues:
            FontH1Rule.fix(doc, issue, value)

        fixed_path = tmp_path / 'fixed_h1.docx'
        doc.save(fixed_path)

        # 重新打开验证
        doc2 = Document(fixed_path)
        remaining = FontH1Rule.detect(doc2, value)
        assert remaining == [], f'修复后仍有 {len(remaining)} 条 issue: {remaining}'

    def test_fix_marks_blue(self, tmp_path: Path):
        """修复后 run 的颜色必须变为蓝色 #0070C0"""
        doc = Document()
        h = doc.add_paragraph('需修复的一级标题', style='Heading 1')
        h.runs[0].font.name = 'Calibri'
        h.runs[0].font.size = Pt(12)
        path = tmp_path / 'for_fix_h1.docx'
        doc.save(path)

        doc2 = Document(path)
        value = {'family': '黑体', 'size_pt': 16}
        issues = FontH1Rule.detect(doc2, value)
        assert len(issues) == 1

        FontH1Rule.fix(doc2, issues[0], value)

        target_run = doc2.paragraphs[issues[0].loc.para].runs[issues[0].loc.run]
        color = target_run.font.color.rgb
        assert color == RGBColor(0x00, 0x70, 0xC0), f'颜色应为 #0070C0，实际为 {color}'

    def test_fix_sets_bold(self, tmp_path: Path):
        """value bold=True 时，fix 后 run.font.bold 必须为 True"""
        doc = Document()
        h = doc.add_paragraph('加粗修复测试', style='Heading 1')
        h.runs[0].font.name = 'Calibri'
        h.runs[0].font.size = Pt(12)
        h.runs[0].font.bold = False
        path = tmp_path / 'bold_fix_h1.docx'
        doc.save(path)

        doc2 = Document(path)
        value = {'family': '黑体', 'size_pt': 16, 'bold': True}
        issues = FontH1Rule.detect(doc2, value)
        assert len(issues) == 1

        result = FontH1Rule.fix(doc2, issues[0], value)
        assert result.applied is True

        target_run = doc2.paragraphs[issues[0].loc.para].runs[issues[0].loc.run]
        assert target_run.font.bold is True, 'fix 后 bold 应为 True'
