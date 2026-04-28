"""
@file: test_checkers.py
@description: 属性 checker 单元测试
              覆盖 CHECKER_MAP 完整性、字体、字号、对齐、缩进、间距、
              内容、页边距等核心检查器的命中/不命中两种路径
@author: Atlas.oi
@date: 2026-04-18
"""
import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from thesis_worker.engine_v2.checkers import (
    CHECKER_MAP,
    DOC_LEVEL_KEYS,
    check_citation_style,
    check_content_char_count_max,
    check_content_char_count_min,
    check_content_item_count_max,
    check_content_item_count_min,
    check_content_item_separator,
    check_content_specific_text,
    check_font_ascii,
    check_font_bold,
    check_font_cjk,
    check_font_italic,
    check_font_size_pt,
    check_layout_position,
    check_mixed_script_ascii_is_tnr,
    check_page_margin_bottom_cm,
    check_page_margin_left_cm,
    check_page_margin_right_cm,
    check_page_margin_top_cm,
    check_page_margin_gutter_cm,
    check_page_header_offset_cm,
    check_page_footer_offset_cm,
    check_page_print_mode,
    check_page_new_page_before,
    check_page_size,
    check_pagination_body_style,
    check_pagination_front_style,
    check_para_align,
    check_para_first_line_indent_chars,
    check_para_hanging_indent_chars,
    check_para_letter_spacing_chars,
    check_para_line_spacing,
    check_para_space_after_lines,
    check_para_space_before_lines,
    check_para_space_before_pt,
    check_para_space_after_pt,
    # T3.2: table.* namespace
    check_table_is_three_line,
    check_table_border_top_pt,
    check_table_border_bottom_pt,
    check_table_header_border_pt,
    # T3.3: numbering.* namespace
    check_numbering_figure_style,
    check_numbering_subfigure_style,
    check_numbering_formula_style,
)
from thesis_worker.engine_v2.field_defs import FIELD_DEFS


# ───────────────────────────────────────────────
# 辅助函数
# ───────────────────────────────────────────────

def _para_with_cjk_font(text: str, font_name: str):
    """创建带中文字体的段落"""
    doc = Document()
    p = doc.add_paragraph(text)
    run = p.runs[0]
    run.font.name = font_name
    # 必须显式设置 eastAsia 属性，python-docx 不会自动写入
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        from lxml import etree
        rfonts = etree.SubElement(rpr, qn('w:rFonts'))
    rfonts.set(qn('w:eastAsia'), font_name)
    return p


def _para_with_ascii_font(text: str, font_name: str):
    """创建带西文字体的段落"""
    doc = Document()
    p = doc.add_paragraph(text)
    p.runs[0].font.name = font_name
    return p


def _para_with_size(text: str, pt: float):
    """创建指定字号的段落"""
    doc = Document()
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(pt)
    return p


# ───────────────────────────────────────────────
# CHECKER_MAP 完整性
# ───────────────────────────────────────────────

class TestCheckerMapCompleteness:
    """确保 CHECKER_MAP 包含 field_defs 中所有 applicable_attributes 键"""

    # T4.2 已清空豁免：first_line_indent_pt / hanging_indent_pt / letter_spacing_pt
    # 三 checker 在 T4.2 实现完成后从此豁免清单移除。
    # 保留空 frozenset 与 invariant 引用，作为未来再次中间态时的扩展点。
    _PENDING_CHECKERS_T42 = frozenset()

    def test_all_applicable_attrs_have_checker(self):
        """field_defs 中每个属性 key 必须在 CHECKER_MAP 中有对应项

        例外：T4.2 待补的 _pt 系列 attr，T2.1 schema 先行，T4.2 加 checker 后清空豁免。
        """
        missing = []
        for field in FIELD_DEFS:
            for attr in field['applicable_attributes']:
                if attr not in CHECKER_MAP and attr not in self._PENDING_CHECKERS_T42:
                    missing.append(attr)
        assert missing == [], f'缺少 checker 的属性 key: {missing}'

    def test_checker_map_has_expected_keys(self):
        """抽查核心属性 key 存在"""
        required = [
            'font.cjk', 'font.ascii', 'font.size_pt', 'font.bold', 'font.italic',
            'para.align', 'para.first_line_indent_chars', 'para.line_spacing',
            'page.size', 'page.margin_top_cm', 'page.new_page_before',
            'content.char_count_max', 'content.specific_text',
            'mixed_script.ascii_is_tnr', 'citation.style', 'layout.position',
        ]
        for key in required:
            assert key in CHECKER_MAP, f'{key} 不在 CHECKER_MAP 中'

    def test_doc_level_keys_are_in_checker_map(self):
        """DOC_LEVEL_KEYS 中的 key 必须都在 CHECKER_MAP 中"""
        for key in DOC_LEVEL_KEYS:
            assert key in CHECKER_MAP, f'文档级 key {key} 不在 CHECKER_MAP 中'


# ───────────────────────────────────────────────
# 字体检查器
# ───────────────────────────────────────────────

class TestFontCjk:
    def test_match_returns_none(self):
        p = _para_with_cjk_font('测试文字', '宋体')
        assert check_font_cjk(p, '宋体') is None

    def test_mismatch_returns_issue(self):
        p = _para_with_cjk_font('测试文字', '宋体')
        issue = check_font_cjk(p, '黑体')
        assert issue is not None
        assert issue['attr'] == 'font.cjk'
        assert issue['actual'] == '宋体'
        assert issue['expected'] == '黑体'

    def test_empty_para_returns_issue(self):
        doc = Document()
        p = doc.add_paragraph('   ')  # 全空白
        # 无法读取字体，actual=None，expected='宋体' → issue
        issue = check_font_cjk(p, '宋体')
        assert issue is not None


class TestFontAscii:
    def test_match_returns_none(self):
        p = _para_with_ascii_font('Abstract', 'Times New Roman')
        assert check_font_ascii(p, 'Times New Roman') is None

    def test_mismatch_returns_issue(self):
        p = _para_with_ascii_font('Abstract', 'Times New Roman')
        issue = check_font_ascii(p, 'Arial')
        assert issue is not None
        assert issue['attr'] == 'font.ascii'


class TestFontSizePt:
    def test_exact_match(self):
        p = _para_with_size('标题', 12)
        assert check_font_size_pt(p, 12) is None

    def test_tolerance_ok(self):
        p = _para_with_size('标题', 12.05)
        # 容差 0.1，0.05 < 0.1，应通过
        assert check_font_size_pt(p, 12) is None

    def test_mismatch(self):
        p = _para_with_size('标题', 12)
        issue = check_font_size_pt(p, 14)
        assert issue is not None
        assert issue['attr'] == 'font.size_pt'
        assert abs(issue['actual'] - 12) < 0.01

    def test_no_size_set_returns_issue(self):
        doc = Document()
        p = doc.add_paragraph('内容')
        # 未设置字号 → None
        issue = check_font_size_pt(p, 12)
        assert issue is not None
        assert issue['actual'] is None


class TestFontBold:
    def test_bold_match(self):
        doc = Document()
        p = doc.add_paragraph('标题')
        p.runs[0].font.bold = True
        assert check_font_bold(p, True) is None

    def test_not_bold_mismatch(self):
        doc = Document()
        p = doc.add_paragraph('正文')
        p.runs[0].font.bold = False
        issue = check_font_bold(p, True)
        assert issue is not None
        assert issue['attr'] == 'font.bold'


class TestFontItalic:
    def test_italic_match(self):
        doc = Document()
        p = doc.add_paragraph('斜体')
        p.runs[0].font.italic = True
        assert check_font_italic(p, True) is None

    def test_none_treated_as_false(self):
        doc = Document()
        p = doc.add_paragraph('正文')
        # font.italic 默认为 None，视为 False
        assert check_font_italic(p, False) is None


# ───────────────────────────────────────────────
# 段落格式检查器
# ───────────────────────────────────────────────

class TestParaAlign:
    def test_center_match(self):
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document()
        p = doc.add_paragraph('居中标题')
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        assert check_para_align(p, 'center') is None

    def test_left_match(self):
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document()
        p = doc.add_paragraph('左对齐')
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        assert check_para_align(p, 'left') is None

    def test_mismatch(self):
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document()
        p = doc.add_paragraph('右对齐')
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        issue = check_para_align(p, 'center')
        assert issue is not None
        assert issue['attr'] == 'para.align'
        assert issue['actual'] == 'right'

    def test_none_alignment_defaults_to_left(self):
        doc = Document()
        p = doc.add_paragraph('默认')
        # alignment 未设置 → None → 视为 left
        assert check_para_align(p, 'left') is None


class TestParaFirstLineIndent:
    def test_two_chars_match(self):
        doc = Document()
        p = doc.add_paragraph('正文段落')
        p.paragraph_format.first_line_indent = Pt(24)  # 2 字 × 12pt
        assert check_para_first_line_indent_chars(p, 2, body_size_pt=12) is None

    def test_zero_indent_mismatch(self):
        doc = Document()
        p = doc.add_paragraph('无缩进')
        issue = check_para_first_line_indent_chars(p, 2)
        assert issue is not None
        assert issue['actual'] == 0


class TestParaHangingIndent:
    def test_hanging_match(self):
        doc = Document()
        p = doc.add_paragraph('参考文献条目')
        # first_line_indent < 0 表示悬挂缩进
        p.paragraph_format.first_line_indent = Pt(-24)  # 悬挂 2 字
        assert check_para_hanging_indent_chars(p, 2, body_size_pt=12) is None

    def test_no_hanging_returns_zero(self):
        doc = Document()
        p = doc.add_paragraph('正文')
        issue = check_para_hanging_indent_chars(p, 2)
        assert issue is not None
        assert issue['actual'] == 0


class TestParaLineSpacing:
    def test_match(self):
        doc = Document()
        p = doc.add_paragraph('行距测试')
        p.paragraph_format.line_spacing = 1.5
        assert check_para_line_spacing(p, 1.5) is None

    def test_none_treated_as_1(self):
        doc = Document()
        p = doc.add_paragraph('默认行距')
        # line_spacing 未设置 → None → 视为 1.0
        issue = check_para_line_spacing(p, 1.5)
        assert issue is not None
        assert abs(issue['actual'] - 1.0) < 0.01

    def test_tolerance(self):
        doc = Document()
        p = doc.add_paragraph('容差测试')
        p.paragraph_format.line_spacing = 1.52
        # 容差 0.05，1.52 vs 1.5 差 0.02 < 0.05，通过
        assert check_para_line_spacing(p, 1.5) is None


class TestParaSpaceBeforeAfter:
    def test_space_before_match(self):
        doc = Document()
        p = doc.add_paragraph('章节标题')
        # 1 行间距 = 12pt（body_size_pt=12）
        p.paragraph_format.space_before = Pt(12)
        assert check_para_space_before_lines(p, 1, body_size_pt=12) is None

    def test_space_after_mismatch(self):
        doc = Document()
        p = doc.add_paragraph('测试')
        p.paragraph_format.space_after = Pt(24)  # 24pt / 12pt body = 2 lines
        issue = check_para_space_after_lines(p, 0, body_size_pt=12)  # 期望 0 行
        assert issue is not None
        assert issue['actual'] == 2
        assert issue['expected'] == 0


# ───────────────────────────────────────────────
# 内容检查器
# ───────────────────────────────────────────────

class TestContentCharCountMax:
    def test_within_limit(self):
        doc = Document()
        p = doc.add_paragraph('短标题')
        assert check_content_char_count_max(p, 25) is None

    def test_exceeds_limit(self):
        doc = Document()
        p = doc.add_paragraph('非常非常非常非常非常非常非常非常长的标题超过了二十五个字符超过了')
        issue = check_content_char_count_max(p, 25)
        assert issue is not None
        assert issue['attr'] == 'content.char_count_max'
        assert issue['actual'] > 25


class TestContentCharCount:
    def test_min_match(self):
        doc = Document()
        p = doc.add_paragraph('摘要' * 100)
        assert check_content_char_count_min(p, 200) is None

    def test_min_fail(self):
        doc = Document()
        p = doc.add_paragraph('太短')
        issue = check_content_char_count_min(p, 200)
        assert issue is not None

    def test_max_match(self):
        doc = Document()
        p = doc.add_paragraph('刚好')
        assert check_content_char_count_max(p, 100) is None

    def test_max_fail(self):
        doc = Document()
        p = doc.add_paragraph('x' * 101)
        issue = check_content_char_count_max(p, 100)
        assert issue is not None


class TestContentItemCount:
    def test_item_count_min_match(self):
        doc = Document()
        p = doc.add_paragraph('关键词一；关键词二；关键词三')
        assert check_content_item_count_min(p, 3, separator='；') is None

    def test_item_count_min_fail(self):
        doc = Document()
        p = doc.add_paragraph('只有一个')
        issue = check_content_item_count_min(p, 3)
        assert issue is not None

    def test_item_count_max_match(self):
        doc = Document()
        p = doc.add_paragraph('a；b；c')
        assert check_content_item_count_max(p, 5, separator='；') is None

    def test_item_count_max_fail(self):
        doc = Document()
        p = doc.add_paragraph('a；b；c；d；e；f')
        issue = check_content_item_count_max(p, 5, separator='；')
        assert issue is not None


class TestContentItemSeparator:
    def test_separator_present(self):
        doc = Document()
        p = doc.add_paragraph('机器学习；深度学习；自然语言处理')
        assert check_content_item_separator(p, '；') is None

    def test_separator_missing(self):
        doc = Document()
        p = doc.add_paragraph('机器学习,深度学习')  # 用了英文逗号
        issue = check_content_item_separator(p, '；')
        assert issue is not None


class TestContentSpecificText:
    def test_contains(self):
        doc = Document()
        p = doc.add_paragraph('摘  要')
        assert check_content_specific_text(p, '摘') is None

    def test_not_contains(self):
        doc = Document()
        p = doc.add_paragraph('引言')
        issue = check_content_specific_text(p, '摘要')
        assert issue is not None
        assert issue['attr'] == 'content.specific_text'


# ───────────────────────────────────────────────
# 分页检查器
# ───────────────────────────────────────────────

class TestPageNewPageBefore:
    def test_page_break_true(self):
        doc = Document()
        p = doc.add_paragraph('第一章')
        p.paragraph_format.page_break_before = True
        assert check_page_new_page_before(p, True) is None

    def test_no_break_mismatch(self):
        doc = Document()
        p = doc.add_paragraph('正文')
        p.paragraph_format.page_break_before = False
        issue = check_page_new_page_before(p, True)
        assert issue is not None


# ───────────────────────────────────────────────
# 页面级检查器（文档级）
# ───────────────────────────────────────────────

class TestPageSize:
    def test_a4_match(self):
        doc = Document()
        from docx.shared import Inches
        doc.sections[0].page_width = Inches(8.27)
        doc.sections[0].page_height = Inches(11.69)
        assert check_page_size(doc, 'A4') is None

    def test_letter_detected(self):
        doc = Document()
        from docx.shared import Inches
        doc.sections[0].page_width = Inches(8.50)
        doc.sections[0].page_height = Inches(11.00)
        result = check_page_size(doc, 'A4')
        # Letter 尺寸 → 期望 A4 → issue
        assert result is not None
        assert result['actual'] == 'Letter'


class TestPageMargins:
    def test_top_margin_match(self):
        doc = Document()
        doc.sections[0].top_margin = Cm(2.54)
        assert check_page_margin_top_cm(doc, 2.54) is None

    def test_top_margin_mismatch(self):
        doc = Document()
        doc.sections[0].top_margin = Cm(3.0)
        issue = check_page_margin_top_cm(doc, 2.54)
        assert issue is not None
        assert issue['attr'] == 'page.margin_top_cm'

    def test_bottom_margin_match(self):
        doc = Document()
        doc.sections[0].bottom_margin = Cm(2.0)
        assert check_page_margin_bottom_cm(doc, 2.0) is None

    def test_left_margin_match(self):
        doc = Document()
        doc.sections[0].left_margin = Cm(3.0)
        assert check_page_margin_left_cm(doc, 3.0) is None

    def test_right_margin_mismatch(self):
        doc = Document()
        doc.sections[0].right_margin = Cm(2.0)
        issue = check_page_margin_right_cm(doc, 3.0)
        assert issue is not None


# ───────────────────────────────────────────────
# T4.1: check_layout_position 测试
# ───────────────────────────────────────────────

def _doc_with_drawing_then_caption() -> tuple:
    """构造"drawing 段落在前，caption 段落在后"的文档（caption 在 below 位置）。

    返回 (doc, caption_para)。
    """
    from lxml import etree
    doc = Document()
    # 第一段：注入 w:drawing，模拟图片段落
    fig_para = doc.add_paragraph()
    run = fig_para.add_run()
    etree.SubElement(run._element, qn('w:drawing'))
    # 第二段：caption 段落，紧接在图片之后
    caption_para = doc.add_paragraph('图1 测试图')
    return doc, caption_para


def _doc_with_caption_then_drawing() -> tuple:
    """构造"caption 段落在前，drawing 段落在后"的文档（caption 在 above 位置）。

    返回 (doc, caption_para)。
    """
    from lxml import etree
    doc = Document()
    # 第一段：caption，紧接在图片之前
    caption_para = doc.add_paragraph('图1 测试图')
    # 第二段：注入 w:drawing，模拟图片段落
    fig_para = doc.add_paragraph()
    run = fig_para.add_run()
    etree.SubElement(run._element, qn('w:drawing'))
    return doc, caption_para


class TestLayoutPosition:
    """check_layout_position 判别力测试（T4.1 实现）。

    判别力设计：
    - PASS below：前驱含 drawing，expected='below' → actual='below' → None。
    - FAIL below→above：前驱含 drawing，actual='below' 但 expected='above' → issue。
    - PASS above：后继含 drawing，expected='above' → actual='above' → None。
    - None path：前后都无图/表邻接 → 无法判定 → None（不论 expected 是什么）。
    """

    def test_caption_below_drawing_passes(self):
        # drawing 在前，caption 在后 → actual='below'，expected='below' → None
        _, caption = _doc_with_drawing_then_caption()
        assert check_layout_position(caption, 'below') is None

    def test_caption_below_drawing_fails_when_expected_above(self):
        # drawing 在前 → actual='below'，expected='above' → issue
        _, caption = _doc_with_drawing_then_caption()
        issue = check_layout_position(caption, 'above')
        assert issue is not None
        assert issue['attr'] == 'layout.position'
        assert issue['actual'] == 'below'
        assert issue['expected'] == 'above'

    def test_caption_above_drawing_passes(self):
        # caption 在前，drawing 在后 → actual='above'，expected='above' → None
        _, caption = _doc_with_caption_then_drawing()
        assert check_layout_position(caption, 'above') is None

    def test_caption_above_drawing_fails_when_expected_below(self):
        # caption 在前，drawing 在后 → actual='above'，expected='below' → issue
        _, caption = _doc_with_caption_then_drawing()
        issue = check_layout_position(caption, 'below')
        assert issue is not None
        assert issue['actual'] == 'above'
        assert issue['expected'] == 'below'

    def test_isolated_caption_returns_none(self):
        # 前后都没有图/表邻接 → 无法判定 → None（不论 expected）
        doc = Document()
        doc.add_paragraph('无关段落 A')
        caption = doc.add_paragraph('图1 孤立 caption')
        doc.add_paragraph('无关段落 B')
        assert check_layout_position(caption, 'below') is None
        assert check_layout_position(caption, 'above') is None

    def test_caption_after_independent_table_returns_below(self):
        """body 级独立 w:tbl 在前，caption 在后 → actual='below'。

        判别力：删除 _para_el_contains_figure_or_table 入口自身 tag 检查
        后此 case 必挂（body 级 tbl 漏检，误返回 None）。
        """
        doc = Document()
        # doc.add_table 直接在 body 写入独立 w:tbl 元素
        doc.add_table(2, 2)
        caption = doc.add_paragraph('表1 测试表格')
        # 前驱是独立 w:tbl → caption 在表格下方
        assert check_layout_position(caption, 'below') is None
        issue = check_layout_position(caption, 'above')
        assert issue is not None
        assert issue['attr'] == 'layout.position'
        assert issue['actual'] == 'below'
        assert issue['expected'] == 'above'

    def test_caption_before_independent_table_returns_above(self):
        """caption 在前，body 级独立 w:tbl 在后 → actual='above'。

        判别力：删除 _para_el_contains_figure_or_table 入口自身 tag 检查
        后此 case 必挂（body 级 tbl 漏检，误返回 None）。
        """
        doc = Document()
        caption = doc.add_paragraph('表1 测试表格')
        # 后继是独立 w:tbl → caption 在表格上方
        doc.add_table(2, 2)
        assert check_layout_position(caption, 'above') is None
        issue = check_layout_position(caption, 'below')
        assert issue is not None
        assert issue['attr'] == 'layout.position'
        assert issue['actual'] == 'above'
        assert issue['expected'] == 'below'


# ───────────────────────────────────────────────
# T4.1: check_citation_style 测试
# ───────────────────────────────────────────────

class TestCitationStyle:
    """check_citation_style 判别力测试（T4.1 实现）。

    判别力设计：
    - PASS gbt7714：[N] 开头，expected='gbt7714' → None。
    - FAIL gbt7714→apa：[N] 开头但 expected='apa' → issue（actual='gbt7714'）。
    - PASS apa：含 (YYYY)，expected='apa' → None。
    - FAIL apa→gbt7714：含 (YYYY) 但 expected='gbt7714' → issue（actual='apa'）。
    - None path（无法推断）：普通文本无 [N] 无 (YYYY) → actual=None → None。
    """

    def test_gbt7714_passes(self):
        # [N] 开头 → gbt7714，expected='gbt7714' → None
        doc = Document()
        p = doc.add_paragraph('[1] 张三, 李四. 机器学习综述. 计算机学报, 2020.')
        assert check_citation_style(p, 'gbt7714') is None

    def test_gbt7714_mismatch_apa(self):
        # [N] 开头 → actual='gbt7714'，expected='apa' → issue
        doc = Document()
        p = doc.add_paragraph('[1] 张三, 李四. 机器学习综述. 计算机学报, 2020.')
        issue = check_citation_style(p, 'apa')
        assert issue is not None
        assert issue['attr'] == 'citation.style'
        assert issue['actual'] == 'gbt7714'
        assert issue['expected'] == 'apa'

    def test_apa_passes(self):
        # 含 (2020) 括号年份 → apa，expected='apa' → None
        doc = Document()
        p = doc.add_paragraph('Zhang, S. (2020). Deep learning overview. Nature.')
        assert check_citation_style(p, 'apa') is None

    def test_apa_mismatch_gbt7714(self):
        # 含 (YYYY) → actual='apa'，expected='gbt7714' → issue
        doc = Document()
        p = doc.add_paragraph('Smith, J. (2019). Title. Journal, 12(3), 45-67.')
        issue = check_citation_style(p, 'gbt7714')
        assert issue is not None
        assert issue['actual'] == 'apa'
        assert issue['expected'] == 'gbt7714'

    def test_unrecognized_style_returns_none(self):
        # 普通段落文本，无 [N] 无 (YYYY) → actual=None → 不报 issue
        doc = Document()
        p = doc.add_paragraph('Smith A. Title of paper. Publisher, 2020.')
        # 无法推断 → 不论 expected 是什么都不报
        assert check_citation_style(p, 'gbt7714') is None
        assert check_citation_style(p, 'apa') is None
        assert check_citation_style(p, 'mla') is None


# ───────────────────────────────────────────────
# T4.1: check_pagination_front_style 测试
# ───────────────────────────────────────────────

def _inject_pgnumtype(doc: Document, section_idx: int, fmt_val: str) -> None:
    """向指定 section 的 sectPr 注入 w:pgNumType fmt 属性。

    业务场景：模拟 Word 在前置页（摘要目录页）设置罗马数字页码，
    或在正文页设置阿拉伯数字页码。
    """
    from lxml import etree
    section_pr = doc.sections[section_idx]._sectPr
    pg_num_type = etree.SubElement(section_pr, qn('w:pgNumType'))
    pg_num_type.set(qn('w:fmt'), fmt_val)


class TestPaginationFrontStyle:
    """check_pagination_front_style 判别力测试（T4.1 实现）。

    判别力设计：
    - PASS roman：注入 lowerRoman → actual='roman'，expected='roman' → None。
    - PASS arabic：不注入（默认）→ actual='arabic'，expected='arabic' → None。
    - FAIL roman→arabic：注入 lowerRoman → actual='roman'，expected='arabic' → issue。
    - FAIL arabic→roman：不注入（默认 arabic）→ expected='roman' → issue。
    """

    def test_lower_roman_passes(self):
        # 注入 lowerRoman → actual='roman'，expected='roman' → None
        doc = Document()
        _inject_pgnumtype(doc, 0, 'lowerRoman')
        assert check_pagination_front_style(doc, 'roman') is None

    def test_upper_roman_also_passes(self):
        # 注入 upperRoman 也映射为 'roman' → None
        doc = Document()
        _inject_pgnumtype(doc, 0, 'upperRoman')
        assert check_pagination_front_style(doc, 'roman') is None

    def test_no_pgnumtype_defaults_to_arabic(self):
        # 未注入 pgNumType → 默认 arabic，expected='arabic' → None
        doc = Document()
        assert check_pagination_front_style(doc, 'arabic') is None

    def test_roman_mismatch_arabic(self):
        # 注入 lowerRoman → actual='roman'，expected='arabic' → issue
        doc = Document()
        _inject_pgnumtype(doc, 0, 'lowerRoman')
        issue = check_pagination_front_style(doc, 'arabic')
        assert issue is not None
        assert issue['attr'] == 'pagination.front_style'
        assert issue['actual'] == 'roman'
        assert issue['expected'] == 'arabic'

    def test_default_arabic_mismatch_roman(self):
        # 未注入（默认 arabic），expected='roman' → issue
        doc = Document()
        issue = check_pagination_front_style(doc, 'roman')
        assert issue is not None
        assert issue['actual'] == 'arabic'
        assert issue['expected'] == 'roman'


# ───────────────────────────────────────────────
# T4.1: check_pagination_body_style 测试
# ───────────────────────────────────────────────

def _doc_with_two_sections() -> Document:
    """构造两个 section 的文档（模拟前置页 roman + 正文 arabic 的毕业论文）。

    python-docx 的 sections 解析规则：
    - 段落 pPr 中含 sectPr → 该段结束一个 section（section[0]）
    - body 末尾的 sectPr → 最后一个 section（section[-1]）

    所以要让 section[0]=roman、section[-1]=arabic（默认），
    需要把带 roman pgNumType 的 sectPr 插入到 body 末尾 sectPr **之前**：
    1. add_paragraph('前置页') 和 add_paragraph('正文')
    2. 在 body 末尾 sectPr 之前插入带 lowerRoman 的 p/pPr/sectPr
    """
    from lxml import etree
    doc = Document()
    doc.add_paragraph('前置页内容')
    doc.add_paragraph('正文内容')

    body = doc.element.body
    # body 末尾是 w:sectPr（文档结束 sectPr）
    last_sectpr = body[-1]

    # 在末尾 sectPr 之前插入 section break 段落（前置 section，roman 页码）
    p_el = etree.Element(qn('w:p'))
    p_pr = etree.SubElement(p_el, qn('w:pPr'))
    sect_pr_new = etree.SubElement(p_pr, qn('w:sectPr'))
    pg_num = etree.SubElement(sect_pr_new, qn('w:pgNumType'))
    pg_num.set(qn('w:fmt'), 'lowerRoman')
    # 插入到末尾 sectPr 之前，使其成为 section[0]
    body.insert(list(body).index(last_sectpr), p_el)

    return doc


class TestPaginationBodyStyle:
    """check_pagination_body_style 判别力测试（T4.1 实现）。

    判别力设计：
    - PASS arabic（单 section 默认）：不注入 → actual='arabic'，expected='arabic' → None。
    - FAIL arabic→roman：不注入（默认 arabic），expected='roman' → issue。
    - 多 section：最后 section 注入 decimal → actual='arabic'；前 section 是 roman，
      检验读的是最后一个而非第一个。
    """

    def test_no_pgnumtype_defaults_to_arabic(self):
        # 未注入 → 默认 arabic，expected='arabic' → None
        doc = Document()
        assert check_pagination_body_style(doc, 'arabic') is None

    def test_arabic_mismatch_roman(self):
        # 默认 arabic，expected='roman' → issue
        doc = Document()
        issue = check_pagination_body_style(doc, 'roman')
        assert issue is not None
        assert issue['attr'] == 'pagination.body_style'
        assert issue['actual'] == 'arabic'
        assert issue['expected'] == 'roman'

    def test_roman_body_fails_expected_arabic(self):
        # 单 section 注入 lowerRoman → actual='roman'，expected='arabic' → issue
        doc = Document()
        _inject_pgnumtype(doc, 0, 'lowerRoman')
        issue = check_pagination_body_style(doc, 'arabic')
        assert issue is not None
        assert issue['actual'] == 'roman'

    def test_multi_section_reads_last_section(self):
        # 两 section 文档：前 section=roman，正文（最后）section 无注入=arabic
        # check_pagination_body_style 应读最后 section → actual='arabic'
        doc = _doc_with_two_sections()
        # 最后一个 section（正文）无 pgNumType → arabic
        assert check_pagination_body_style(doc, 'arabic') is None


# ───────────────────────────────────────────────
# T3.1: 新增 checker 测试
# ───────────────────────────────────────────────

class TestParaSpaceBeforePt:
    """段前磅值检查器（容差 0.5pt）的判别力测试

    判别力设计：
    - actual=12.4, expected=12 → |12.4-12|=0.4 < 0.5 → 容差内 → None
    - actual=12.6, expected=12 → |12.6-12|=0.6 > 0.5 → 超出容差 → 返回 dict
    两个路径使用同一 expected=12 但不同 actual，确保测试真正检测了容差边界。
    """

    def test_within_tolerance_returns_none(self):
        doc = Document()
        p = doc.add_paragraph('章节标题')
        # 12.4pt：与 expected=12 差 0.4 < _TOL_SPACE_PT=0.5 → 符合
        p.paragraph_format.space_before = Pt(12.4)
        assert check_para_space_before_pt(p, 12.0) is None

    def test_exceeds_tolerance_returns_issue(self):
        doc = Document()
        p = doc.add_paragraph('章节标题')
        # 12.6pt：与 expected=12 差 0.6 > _TOL_SPACE_PT=0.5 → 违规
        p.paragraph_format.space_before = Pt(12.6)
        issue = check_para_space_before_pt(p, 12.0)
        assert issue is not None
        assert issue['attr'] == 'para.space_before_pt'
        assert abs(issue['actual'] - 12.6) < 0.05
        assert issue['expected'] == 12.0

    def test_none_space_before_treated_as_zero(self):
        doc = Document()
        p = doc.add_paragraph('无间距')
        # space_before 未设置 → None → 视为 0pt
        # expected=0 → 符合
        assert check_para_space_before_pt(p, 0.0) is None

    def test_none_space_before_mismatch(self):
        doc = Document()
        p = doc.add_paragraph('无间距')
        # space_before None（0pt）但 expected=12 → 违规
        issue = check_para_space_before_pt(p, 12.0)
        assert issue is not None
        assert issue['actual'] == 0.0


class TestParaSpaceAfterPt:
    """段后磅值检查器（容差 0.5pt）的判别力测试

    判别力设计：
    - actual=6.4, expected=6 → |6.4-6|=0.4 < 0.5 → 容差内 → None
    - actual=6.6, expected=6 → |6.6-6|=0.6 > 0.5 → 超出容差 → 返回 dict
    None 路径与 TestParaSpaceBeforePt 对称：None space_after 视为 0pt。
    """

    def test_within_tolerance_returns_none(self):
        doc = Document()
        p = doc.add_paragraph('段落')
        # 6.4pt vs expected=6 → 差 0.4 < 0.5 → 符合
        p.paragraph_format.space_after = Pt(6.4)
        assert check_para_space_after_pt(p, 6.0) is None

    def test_exceeds_tolerance_returns_issue(self):
        doc = Document()
        p = doc.add_paragraph('段落')
        # 6.6pt vs expected=6 → 差 0.6 > 0.5 → 违规
        p.paragraph_format.space_after = Pt(6.6)
        issue = check_para_space_after_pt(p, 6.0)
        assert issue is not None
        assert issue['attr'] == 'para.space_after_pt'
        assert issue['expected'] == 6.0

    def test_none_space_after_treated_as_zero(self):
        doc = Document()
        p = doc.add_paragraph('无间距')
        # space_after 未设置 → None → 视为 0pt
        # expected=0 → 符合
        assert check_para_space_after_pt(p, 0.0) is None

    def test_none_space_after_mismatch(self):
        doc = Document()
        p = doc.add_paragraph('无间距')
        # space_after None（0pt）但 expected=6 → 违规
        issue = check_para_space_after_pt(p, 6.0)
        assert issue is not None
        assert issue['actual'] == 0.0
        assert issue['expected'] == 6.0


class TestPageMarginGutterCm:
    """装订线宽度检查器（容差 0.05cm）"""

    def test_match_returns_none(self):
        doc = Document()
        doc.sections[0].gutter = Cm(1.0)
        assert check_page_margin_gutter_cm(doc, 1.0) is None

    def test_within_tolerance_returns_none(self):
        doc = Document()
        # 1.03cm vs expected=1.0 → 差 0.03 < 0.05 → 符合
        doc.sections[0].gutter = Cm(1.03)
        assert check_page_margin_gutter_cm(doc, 1.0) is None

    def test_exceeds_tolerance_returns_issue(self):
        doc = Document()
        # 1.8cm vs expected=1.0 → 差 0.8 > 0.05 → 违规
        doc.sections[0].gutter = Cm(1.8)
        issue = check_page_margin_gutter_cm(doc, 1.0)
        assert issue is not None
        assert issue['attr'] == 'page.margin_gutter_cm'

    def test_zero_gutter_valid(self):
        # 装订线为 0 也是有效值（无装订线）
        doc = Document()
        doc.sections[0].gutter = Cm(0)
        assert check_page_margin_gutter_cm(doc, 0.0) is None


class TestPageHeaderOffsetCm:
    """页眉距边界检查器（容差 0.05cm）"""

    def test_match_returns_none(self):
        doc = Document()
        doc.sections[0].header_distance = Cm(1.5)
        assert check_page_header_offset_cm(doc, 1.5) is None

    def test_within_tolerance_returns_none(self):
        doc = Document()
        # 1.53cm vs expected=1.5 → 差 0.03 < 0.05 → 符合
        doc.sections[0].header_distance = Cm(1.53)
        assert check_page_header_offset_cm(doc, 1.5) is None

    def test_mismatch_returns_issue(self):
        doc = Document()
        doc.sections[0].header_distance = Cm(2.0)
        # expected=1.5，actual=2.0，差 0.5 > 0.05 → 违规
        issue = check_page_header_offset_cm(doc, 1.5)
        assert issue is not None
        assert issue['attr'] == 'page.header_offset_cm'
        assert issue['expected'] == 1.5


class TestPageFooterOffsetCm:
    """页脚距边界检查器（容差 0.05cm）"""

    def test_match_returns_none(self):
        doc = Document()
        doc.sections[0].footer_distance = Cm(1.75)
        assert check_page_footer_offset_cm(doc, 1.75) is None

    def test_within_tolerance_returns_none(self):
        doc = Document()
        # 1.78cm vs expected=1.75 → 差 0.03 < 0.05 → 符合
        doc.sections[0].footer_distance = Cm(1.78)
        assert check_page_footer_offset_cm(doc, 1.75) is None

    def test_mismatch_returns_issue(self):
        doc = Document()
        doc.sections[0].footer_distance = Cm(1.0)
        # expected=1.75，actual=1.0 → 差 0.75 > 0.05 → 违规
        issue = check_page_footer_offset_cm(doc, 1.75)
        assert issue is not None
        assert issue['attr'] == 'page.footer_offset_cm'


class TestPagePrintMode:
    """打印模式检查器（严格相等，无容差）

    判别力设计：
    - 'single' vs 'double' 是互斥值，任何一个赋值都能独立区分两条路径。
    测试通过 lxml 注入 w:evenAndOddHeaders 来构造 double 模式文档，
    确保检测逻辑真实走通。
    """

    def test_single_mode_default(self):
        # 新建文档默认无 w:evenAndOddHeaders → single 模式
        doc = Document()
        assert check_page_print_mode(doc, 'single') is None

    def test_single_mode_mismatch(self):
        # 默认 single 但期望 double → 违规
        doc = Document()
        issue = check_page_print_mode(doc, 'double')
        assert issue is not None
        assert issue['attr'] == 'page.print_mode'
        assert issue['actual'] == 'single'
        assert issue['expected'] == 'double'

    def test_double_mode_with_even_odd_headers(self):
        # 注入 w:evenAndOddHeaders → double 模式
        from docx.oxml.ns import qn
        from lxml import etree
        doc = Document()
        settings_el = doc.settings.element
        etree.SubElement(settings_el, qn('w:evenAndOddHeaders'))
        assert check_page_print_mode(doc, 'double') is None

    def test_double_mode_mismatch(self):
        # 注入 w:evenAndOddHeaders → double，但 expected=single → 违规
        from docx.oxml.ns import qn
        from lxml import etree
        doc = Document()
        settings_el = doc.settings.element
        etree.SubElement(settings_el, qn('w:evenAndOddHeaders'))
        issue = check_page_print_mode(doc, 'single')
        assert issue is not None
        assert issue['actual'] == 'double'
        assert issue['expected'] == 'single'


# ───────────────────────────────────────────────
# T3.2: table.* namespace checker 测试
# ───────────────────────────────────────────────

def _doc_with_three_line_table(top_sz: int = 12, bottom_sz: int = 12, inside_h_sz: int = 4) -> Document:
    """构造含三线表 tblBorders 的文档辅助函数。

    参数均使用 eighth-points（OOXML 原生单位）：
    - top_sz=12 → 1.5pt（规范上线）
    - bottom_sz=12 → 1.5pt（规范下线）
    - inside_h_sz=4 → 0.5pt（规范表头下线）
    不注入 insideV → 视为 0，即无纵格线（三线表特征）。
    """
    from lxml import etree
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    tbl_pr = table._element.find(qn('w:tblPr'))
    borders = etree.SubElement(tbl_pr, qn('w:tblBorders'))
    for tag, sz in [('top', top_sz), ('bottom', bottom_sz), ('insideH', inside_h_sz)]:
        el = etree.SubElement(borders, qn(f'w:{tag}'))
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:val'), 'single')
    return doc


def _doc_with_full_border_table() -> Document:
    """构造含四边全线表格（非三线表）的文档：insideV 有值。"""
    from lxml import etree
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    tbl_pr = table._element.find(qn('w:tblPr'))
    borders = etree.SubElement(tbl_pr, qn('w:tblBorders'))
    for tag, sz in [('top', 12), ('bottom', 12), ('insideH', 4), ('insideV', 8)]:
        el = etree.SubElement(borders, qn(f'w:{tag}'))
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:val'), 'single')
    return doc


class TestTableIsThreeLine:
    """三线表判定检查器。

    判别力设计：
    - PASS case：仅有 top/bottom/insideH，无 insideV → 判定为三线表。
    - FAIL case：注入 insideV（sz=8）→ 存在纵格线，不是三线表。
    - None path：文档无表格时 actual=False，expected=False → None（无 issue）。
    删除 check_table_is_three_line 主体后两个 PASS/FAIL case 会互换结果，测试必挂。
    """

    def test_three_line_table_passes(self):
        # 三线表：无 insideV，top/bottom > 0 → actual=True，expected=True → None
        doc = _doc_with_three_line_table()
        assert check_table_is_three_line(doc, True) is None

    def test_full_border_table_fails(self):
        # 有 insideV（sz=8=1pt）→ 不是三线表，actual=False，expected=True → issue
        doc = _doc_with_full_border_table()
        issue = check_table_is_three_line(doc, True)
        assert issue is not None
        assert issue['attr'] == 'table.is_three_line'
        assert issue['actual'] is False
        assert issue['expected'] is True

    def test_full_border_table_returns_none_when_expected_false(self):
        # 有 insideV（纵格线）→ actual=False；expected=False → 符合规范，返回 None。
        # 判别力：删除 insideV==0 的逻辑后，actual 变为 True，与 expected=False 不符 → 挂。
        doc = _doc_with_full_border_table()
        assert check_table_is_three_line(doc, False) is None

    def test_no_table_returns_issue_when_expected_true(self):
        # 文档无表格 → actual=False；expected=True → 违规，返回 issue
        doc = Document()
        issue = check_table_is_three_line(doc, True)
        assert issue is not None
        assert issue['actual'] is False


class TestTableBorderTopPt:
    """表格上边框线宽检查器（容差 0.1pt）。

    判别力设计：
    - PASS：sz=12 eighth-pts = 1.5pt，expected=1.5 → 差 0 < 0.1 → None。
    - FAIL：sz=12 = 1.5pt，expected=3.0 → 差 1.5 > 0.1 → issue。
    - 容差边界：sz=12 = 1.5pt，expected=1.57（差 0.07 < 0.1）→ None。
    - None path：文档无表格，actual=0.0，expected=0.0 → None。
    """

    def test_pass_exact(self):
        # sz=12 eighth-pts = 1.5pt，expected=1.5 → 通过
        doc = _doc_with_three_line_table(top_sz=12)
        assert check_table_border_top_pt(doc, 1.5) is None

    def test_fail_mismatch(self):
        # sz=12 = 1.5pt，expected=3.0 → 违规
        doc = _doc_with_three_line_table(top_sz=12)
        issue = check_table_border_top_pt(doc, 3.0)
        assert issue is not None
        assert issue['attr'] == 'table.border_top_pt'
        # actual 应为 1.5（不是 3.0）
        assert abs(issue['actual'] - 1.5) < 0.01

    def test_tolerance_boundary_pass(self):
        # sz=12 = 1.5pt，expected=1.57（差 0.07 < 0.1）→ 容差内，通过
        doc = _doc_with_three_line_table(top_sz=12)
        assert check_table_border_top_pt(doc, 1.57) is None

    def test_tolerance_inner_boundary_pass(self):
        # sz=12 = 1.5pt，expected=1.45（差 0.05 < _TOL_FONT_PT=0.1）→ 容差内，通过。
        # 判别力：将 abs<容差 改为严格相等后，1.5 != 1.45 → 测试挂。
        doc = _doc_with_three_line_table(top_sz=12)
        assert check_table_border_top_pt(doc, 1.45) is None


class TestTableBorderBottomPt:
    """表格下边框线宽检查器（容差 0.1pt）。

    判别力设计：sz=12=1.5pt vs expected 两个互斥值覆盖两条路径。
    """

    def test_pass_exact(self):
        doc = _doc_with_three_line_table(bottom_sz=12)
        assert check_table_border_bottom_pt(doc, 1.5) is None

    def test_fail_mismatch(self):
        # sz=12 = 1.5pt，expected=0.5 → 差 1.0 > 0.1 → 违规
        doc = _doc_with_three_line_table(bottom_sz=12)
        issue = check_table_border_bottom_pt(doc, 0.5)
        assert issue is not None
        assert issue['attr'] == 'table.border_bottom_pt'
        assert abs(issue['actual'] - 1.5) < 0.01

    def test_tolerance_inner_boundary_pass(self):
        # sz=12 = 1.5pt，expected=1.45（差 0.05 < _TOL_FONT_PT=0.1）→ 容差内，通过。
        # 判别力：将 abs<容差 改为严格相等后，1.5 != 1.45 → 测试挂。
        doc = _doc_with_three_line_table(bottom_sz=12)
        assert check_table_border_bottom_pt(doc, 1.45) is None


class TestTableHeaderBorderPt:
    """表头下边框线宽检查器（insideH，容差 0.1pt）。

    判别力设计：
    - PASS：sz=4 eighth-pts = 0.5pt，expected=0.5 → 通过（三线表标准值）。
    - FAIL：sz=4 = 0.5pt，expected=1.5 → 差 1.0 > 0.1 → 违规（与上下线混淆可检出）。
    - 容差边界：sz=4 = 0.5pt，expected=0.55（差 0.05 < 0.1）→ 通过。
    - None path：文档无表格 → actual=0.0；expected=0.0 → None。
    """

    def test_pass_standard_half_pt(self):
        # sz=4 eighth-pts = 0.5pt，符合规范表头下线标准值
        doc = _doc_with_three_line_table(inside_h_sz=4)
        assert check_table_header_border_pt(doc, 0.5) is None

    def test_fail_wrong_expected(self):
        # sz=4 = 0.5pt，但期望 1.5pt → 违规，防止把 header 和 top/bottom 搞混
        doc = _doc_with_three_line_table(inside_h_sz=4)
        issue = check_table_header_border_pt(doc, 1.5)
        assert issue is not None
        assert issue['attr'] == 'table.header_border_pt'
        assert abs(issue['actual'] - 0.5) < 0.01

    def test_tolerance_boundary_pass(self):
        # sz=4 = 0.5pt，expected=0.55（差 0.05 < 0.1）→ 容差内通过
        doc = _doc_with_three_line_table(inside_h_sz=4)
        assert check_table_header_border_pt(doc, 0.55) is None

    def test_tolerance_inner_boundary_pass(self):
        # sz=4 = 0.5pt，expected=0.55（差 0.05 < _TOL_FONT_PT=0.1）→ 容差内，通过。
        # 判别力：将 abs<容差 改为严格相等后，0.5 != 0.55 → 测试挂。
        doc = _doc_with_three_line_table(inside_h_sz=4)
        assert check_table_header_border_pt(doc, 0.55) is None


# ───────────────────────────────────────────────
# T3.3: numbering.* namespace checker 测试
# ───────────────────────────────────────────────

def _doc_with_figure_captions(*captions: str) -> Document:
    """构造含多个图题段落的文档辅助函数。

    直接在段落文本中写入图题字符串（如"图1 系统架构"/"图1-1 数据流"），
    _read_numbering_styles 从文本正则推断风格，不依赖 Word 样式名。
    """
    doc = Document()
    for cap in captions:
        doc.add_paragraph(cap)
    return doc


def _doc_with_formula_numbers(*texts: str) -> Document:
    """构造含多个公式编号的文档辅助函数。

    段落文本末尾含括号编号（如"E = mc² (1)"/"F = ma (1-1)"）。
    """
    doc = Document()
    for text in texts:
        doc.add_paragraph(text)
    return doc


class TestNumberingFigureStyle:
    """图编号风格检查器测试。

    判别力设计：
    - PASS：5 个连续图题 + expected='continuous' → actual='continuous' → None。
    - FAIL：5 个连续图题 + expected='chapter_based' → actual='continuous' → issue。
    - None path（样本不足）：只有 1 个图题（< 2）→ actual=None → None。
    删除 check_numbering_figure_style 主体逻辑后 PASS/FAIL case 互换结果，测试必挂。
    """

    def test_continuous_passes(self):
        # 5 个"图N"格式图题 → actual=continuous，expected=continuous → None
        doc = _doc_with_figure_captions(
            '图1 系统架构图', '图2 数据流图', '图3 模块关系',
            '图4 时序图', '图5 部署图',
        )
        assert check_numbering_figure_style(doc, 'continuous') is None

    def test_continuous_mismatch_chapter_based(self):
        # 5 个连续图题，但 expected=chapter_based → actual 与 expected 互斥 → issue
        doc = _doc_with_figure_captions(
            '图1 系统架构图', '图2 数据流图', '图3 模块关系',
            '图4 时序图', '图5 部署图',
        )
        issue = check_numbering_figure_style(doc, 'chapter_based')
        assert issue is not None
        assert issue['attr'] == 'numbering.figure_style'
        assert issue['actual'] == 'continuous'
        assert issue['expected'] == 'chapter_based'

    def test_chapter_based_passes(self):
        # 5 个"图N-M"格式图题 → actual=chapter_based，expected=chapter_based → None
        doc = _doc_with_figure_captions(
            '图1-1 总体架构', '图1-2 子系统架构',
            '图2-1 数据模型', '图2-2 ER图', '图3-1 时序图',
        )
        assert check_numbering_figure_style(doc, 'chapter_based') is None

    def test_chapter_based_mismatch_continuous(self):
        # 5 个章节式图题，expected=continuous → actual 与 expected 互斥 → issue
        doc = _doc_with_figure_captions(
            '图1-1 总体架构', '图1-2 子系统架构',
            '图2-1 数据模型', '图2-2 ER图', '图3-1 时序图',
        )
        issue = check_numbering_figure_style(doc, 'continuous')
        assert issue is not None
        assert issue['actual'] == 'chapter_based'
        assert issue['expected'] == 'continuous'

    def test_insufficient_sample_returns_none(self):
        # 只有 1 个图题（样本不足），actual=None → 不报 issue（无法判定）
        doc = _doc_with_figure_captions('图1 唯一图题')
        assert check_numbering_figure_style(doc, 'continuous') is None
        assert check_numbering_figure_style(doc, 'chapter_based') is None

    def test_tie_vote_conservatively_selects_chapter_based(self):
        """等票场景：1 个连续式 + 1 个章节式 → 保守选 chapter_based。

        判别力设计（等票路径）：
        - 等票时 actual='chapter_based'，expected='chapter_based' → None（保守选通过）
        - 等票时 actual='chapter_based'，expected='continuous' → issue（actual=chapter_based）
        删除等票分支（else: result=chapter_based）后，两个 path 均会得到 actual=None
        （total=2 但无多数方），断言都会挂。
        """
        # 1 个连续式 + 1 个章节式 → fig_continuous=1, fig_chapter=1 → 等票
        doc = _doc_with_figure_captions('图1 连续图题', '图1-1 章节图题')
        # 等票 → actual=chapter_based → expected=chapter_based → None
        assert check_numbering_figure_style(doc, 'chapter_based') is None
        # 等票 → actual=chapter_based → expected=continuous → issue
        issue = check_numbering_figure_style(doc, 'continuous')
        assert issue is not None
        assert issue['actual'] == 'chapter_based'
        assert issue['expected'] == 'continuous'


class TestNumberingSubfigureStyle:
    """分图编号风格检查器测试。

    判别力设计：
    - PASS：含字母标记"图1(a)""图1(b)" + expected='a_b_c' → actual='a_b_c' → None。
    - FAIL：同样含字母标记，expected='1_2_3' → actual='a_b_c' → issue。
    - None path：图题无子图标记 → subfigure_style 不写入 → actual=None → None。
    """

    def test_letter_style_passes(self):
        # 含"图1(a)""图1(b)"分图标记，共有 >=2 个图题，expected=a_b_c → None
        doc = _doc_with_figure_captions(
            '图1(a) 方案A架构', '图1(b) 方案B架构',
            '图2 总体对比', '图3 详细流程',
        )
        assert check_numbering_subfigure_style(doc, 'a_b_c') is None

    def test_letter_style_mismatch(self):
        # 含字母分图标记，但 expected='1_2_3' → actual 与 expected 互斥 → issue
        doc = _doc_with_figure_captions(
            '图1(a) 方案A架构', '图1(b) 方案B架构',
            '图2 总体对比', '图3 详细流程',
        )
        issue = check_numbering_subfigure_style(doc, '1_2_3')
        assert issue is not None
        assert issue['attr'] == 'numbering.subfigure_style'
        assert issue['actual'] == 'a_b_c'
        assert issue['expected'] == '1_2_3'

    def test_no_subfigure_returns_none(self):
        # 无子图标记（普通图题），actual=None → 不报 issue（无法判定）
        doc = _doc_with_figure_captions(
            '图1 系统架构图', '图2 数据流图', '图3 模块关系',
        )
        assert check_numbering_subfigure_style(doc, 'a_b_c') is None
        assert check_numbering_subfigure_style(doc, '1_2_3') is None


class TestNumberingFormulaStyle:
    """公式编号风格检查器测试。

    判别力设计：
    - PASS：5 个"(N)"形式编号 + expected='continuous' → actual='continuous' → None。
    - FAIL：5 个连续编号，expected='chapter_based' → actual 与 expected 互斥 → issue。
    - None path：只有 1 个公式编号（< 2）→ actual=None → None。
    """

    def test_continuous_passes(self):
        # 5 个"(N)"连续编号，expected=continuous → None
        doc = _doc_with_formula_numbers(
            'E = mc² (1)', 'F = ma (2)', 'P = mv (3)',
            'W = Fs (4)', 'Q = mc△T (5)',
        )
        assert check_numbering_formula_style(doc, 'continuous') is None

    def test_continuous_mismatch_chapter_based(self):
        # 5 个连续编号，expected=chapter_based → actual 与 expected 互斥 → issue
        doc = _doc_with_formula_numbers(
            'E = mc² (1)', 'F = ma (2)', 'P = mv (3)',
            'W = Fs (4)', 'Q = mc△T (5)',
        )
        issue = check_numbering_formula_style(doc, 'chapter_based')
        assert issue is not None
        assert issue['attr'] == 'numbering.formula_style'
        assert issue['actual'] == 'continuous'
        assert issue['expected'] == 'chapter_based'

    def test_chapter_based_passes(self):
        # 5 个"(N-M)"章节式编号，expected=chapter_based → None
        doc = _doc_with_formula_numbers(
            'E = mc² (1-1)', 'F = ma (1-2)', 'P = mv (2-1)',
            'W = Fs (2-2)', 'Q = mc△T (3-1)',
        )
        assert check_numbering_formula_style(doc, 'chapter_based') is None

    def test_chapter_based_mismatch_continuous(self):
        # 5 个章节式编号，expected=continuous → actual 与 expected 互斥 → issue
        doc = _doc_with_formula_numbers(
            'E = mc² (1-1)', 'F = ma (1-2)', 'P = mv (2-1)',
            'W = Fs (2-2)', 'Q = mc△T (3-1)',
        )
        issue = check_numbering_formula_style(doc, 'continuous')
        assert issue is not None
        assert issue['actual'] == 'chapter_based'
        assert issue['expected'] == 'continuous'

    def test_insufficient_sample_returns_none(self):
        # 只有 1 个公式编号（样本不足），actual=None → 不报 issue
        doc = _doc_with_formula_numbers('唯一公式 E = mc² (1)')
        assert check_numbering_formula_style(doc, 'continuous') is None
        assert check_numbering_formula_style(doc, 'chapter_based') is None


# ───────────────────────────────────────────────
# W4 修复：三线表判定加 left/right == 0 的新测试
# ───────────────────────────────────────────────

def _doc_with_four_border_table() -> Document:
    """构造带 left/right 外侧竖线的表格（4 边框，非三线表）。

    构造 top/bottom/left/right 均有值，insideV=0，
    原版代码（只检查 insideV）会误判为三线表。
    W4 修复后必须判 False。
    """
    from lxml import etree
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    tbl_pr = table._element.find(qn('w:tblPr'))
    borders = etree.SubElement(tbl_pr, qn('w:tblBorders'))
    # top=12(1.5pt), bottom=12(1.5pt), left=8(1pt), right=8(1pt)，无 insideV
    for tag, sz in [('top', 12), ('bottom', 12), ('left', 8), ('right', 8)]:
        el = etree.SubElement(borders, qn(f'w:{tag}'))
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:val'), 'single')
    return doc


class TestTableIsThreeLineW4:
    """W4 修复：三线表判定加 left/right == 0 的补充测试。

    判别力设计：
    - FAIL case：表有 top/bottom/left/right 但无 insideV → 原版误判 True，W4 修复后 actual=False。
      删除 left_pt==0 or right_pt==0 的判定后 actual 变回 True，FAIL case 变 PASS，测试挂。
    """

    def test_table_with_left_right_borders_not_three_line(self):
        """带 left/right 外侧竖线、无 insideV 的表格不是三线表（W4 核心修复）。

        判别力：原版只检查 insideV，left/right 有值时 actual 仍为 True（误判）。
        W4 修复后 actual=False；expected=True → issue（不符合三线表规范）。
        """
        doc = _doc_with_four_border_table()
        issue = check_table_is_three_line(doc, True)
        assert issue is not None, (
            'W4 修复：带 left/right 外框的表不是三线表，应返回 issue'
        )
        assert issue['actual'] is False
        assert issue['expected'] is True

    def test_table_with_left_right_borders_expected_false_passes(self):
        """带 left/right 外框的表（actual=False），expected=False → None（符合期望）。

        判别力：actual=False 与 expected=False 相等 → None；
        若删除 left/right 判定 actual 变 True != False → 会变成 issue，测试挂。
        """
        doc = _doc_with_four_border_table()
        assert check_table_is_three_line(doc, False) is None

    def test_standard_three_line_still_passes_after_w4(self):
        """标准三线表（top/bottom/insideH，无 left/right/insideV）在 W4 修复后仍通过。

        防回归：W4 不应影响正常三线表判定。
        """
        doc = _doc_with_three_line_table()
        assert check_table_is_three_line(doc, True) is None


# ───────────────────────────────────────────────
# W5 修复：table.header_border_pt 优先读 tcBorders/bottom 的新测试
# ───────────────────────────────────────────────

def _doc_with_tc_bottom_border(tc_sz: int = 8) -> Document:
    """构造第一行 tc 含 tcBorders/bottom 的表格文档（不设 tblBorders/insideH）。

    W5 修复验证：checker 应优先从 tcBorders/bottom 读表头下线，
    而非 tblBorders/insideH（原有 fallback）。
    tc_sz=8 eighth-points → 1.0pt。
    """
    from lxml import etree
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    # 只设 tblBorders top/bottom，不设 insideH（确保 fallback 路径返回 0）
    tbl_pr = table._element.find(qn('w:tblPr'))
    tbl_borders = etree.SubElement(tbl_pr, qn('w:tblBorders'))
    for tag, sz in [('top', 12), ('bottom', 12)]:
        el = etree.SubElement(tbl_borders, qn(f'w:{tag}'))
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:val'), 'single')
    # 给第一行每个 tc 注入 tcBorders/bottom
    rows = table._element.findall(qn('w:tr'))
    first_row = rows[0]
    for tc in first_row.findall(qn('w:tc')):
        tc_pr = tc.find(qn('w:tcPr'))
        if tc_pr is None:
            tc_pr = etree.SubElement(tc, qn('w:tcPr'))
        tc_borders = etree.SubElement(tc_pr, qn('w:tcBorders'))
        bottom_el = etree.SubElement(tc_borders, qn('w:bottom'))
        bottom_el.set(qn('w:sz'), str(tc_sz))
        bottom_el.set(qn('w:val'), 'single')
    return doc


class TestTableHeaderBorderPtW5:
    """W5 修复：table.header_border_pt 优先读 tcBorders/bottom 的补充测试。

    判别力设计：
    - 优先路径：tc_sz=8(1.0pt)，无 insideH → 读 tcBorders，actual=1.0。
      删除 tcBorders 优先逻辑后 actual 变 0.0（insideH fallback 返回 0），
      断言 actual≈1.0 → 挂。
    - fallback 路径：有 insideH(sz=4=0.5pt)，无 tcBorders → fallback，actual=0.5。
      与现有 test_pass_standard_half_pt 测试共用。
    """

    def test_tc_bottom_border_takes_priority_over_inside_h(self):
        """tcBorders/bottom 存在时，优先读 tc 第一行 bottom（W5 核心）。

        构造：tcBorders/bottom sz=8(1.0pt)；tblBorders 无 insideH（fallback 返回 0）。
        期望 actual=1.0（来自 tcBorders），不是 0.0（来自 fallback）。
        """
        doc = _doc_with_tc_bottom_border(tc_sz=8)
        # 应从 tcBorders/bottom 读到 1.0pt（sz=8/8=1.0）
        assert check_table_header_border_pt(doc, 1.0) is None

    def test_tc_bottom_border_fail_when_expected_wrong(self):
        """tcBorders/bottom sz=8(1.0pt)，expected=0.5 → 差 0.5 > 0.1 → issue。

        判别力：如果 actual 仍是 0.0（即 W5 未生效），则 abs(0.0-0.5)=0.5>0.1 也会 issue，
        但 issue['actual'] 值不同（0.0 vs 1.0），可通过值断言区分。
        """
        doc = _doc_with_tc_bottom_border(tc_sz=8)
        issue = check_table_header_border_pt(doc, 0.5)
        assert issue is not None
        assert issue['attr'] == 'table.header_border_pt'
        # 关键：actual 必须是 1.0（来自 tcBorders），不是 0.0（fallback）
        assert abs(issue['actual'] - 1.0) < 0.01, (
            f"W5：actual 应为 1.0（tcBorders/bottom），实际为 {issue['actual']}"
        )

    def test_fallback_to_inside_h_when_no_tc_borders(self):
        """无 tcBorders/bottom 时，fallback 到 tblBorders/insideH（原有路径）。

        防回归：W5 不应破坏原 insideH fallback 逻辑。
        """
        doc = _doc_with_three_line_table(inside_h_sz=4)  # insideH=0.5pt，无 tcBorders
        # 原有路径：insideH sz=4 → 0.5pt
        assert check_table_header_border_pt(doc, 0.5) is None


# ───────────────────────────────────────────────
# T4.1: line_spacing checker 修 atLeast/exactly 漏判
# ───────────────────────────────────────────────


class TestCheckLineSpacingLineRule:
    """T4.1: line_spacing checker 修 atLeast/exactly 漏判

    背景：原 check_para_line_spacing 用 isinstance(ls, (int, float)) 判定，
    python-docx 在 atLeast/exactly 模式下返回 Length 对象，被静默忽略
    → 论文用"固定值 28 磅"时 detector 完全漏判。
    本组测试构造带 OOXML <w:spacing w:lineRule w:line/> 的段落，
    验证新增的 check_para_line_spacing_type / _pt 直接读 OOXML 区分 6 种状态。
    """

    def _make_para_with_line_rule(self, line_rule: str, line_value: int):
        """构造带 OOXML <w:spacing w:lineRule="X" w:line="Y"/> 的 paragraph

        参数说明：
          line_rule: 'auto' / 'atLeast' / 'exact'（OOXML 原始值）
          line_value: line 属性值，单位 twips（1/20 pt）
        """
        from docx import Document
        from docx.oxml.ns import qn
        from lxml import etree
        doc = Document()
        para = doc.add_paragraph('测试')
        pPr = para._element.get_or_add_pPr()
        spacing = etree.SubElement(pPr, qn('w:spacing'))
        spacing.set(qn('w:lineRule'), line_rule)
        spacing.set(qn('w:line'), str(line_value))
        return para

    def test_at_least_28pt_matches_at_least_28pt(self):
        """atLeast 28pt（line=560 twips）vs expected atLeast 28pt → None"""
        from thesis_worker.engine_v2.checkers import (
            check_para_line_spacing_type,
            check_para_line_spacing_pt,
        )
        para = self._make_para_with_line_rule('atLeast', 560)
        assert check_para_line_spacing_type(para, 'atLeast') is None
        assert check_para_line_spacing_pt(para, 28.0) is None

    def test_at_least_28pt_vs_exactly_expected_returns_issue(self):
        """atLeast 28pt vs expected exactly → issue（类型不同）"""
        from thesis_worker.engine_v2.checkers import check_para_line_spacing_type
        para = self._make_para_with_line_rule('atLeast', 560)
        out = check_para_line_spacing_type(para, 'exactly')
        assert out is not None
        assert out['actual'] == 'atLeast'
        assert out['expected'] == 'exactly'

    def test_auto_240twips_is_single(self):
        """auto + 240 twips → single 倍数（240/240 = 1.0）"""
        from thesis_worker.engine_v2.checkers import check_para_line_spacing_type
        para = self._make_para_with_line_rule('auto', 240)
        assert check_para_line_spacing_type(para, 'single') is None

    def test_auto_360twips_is_one_and_half(self):
        """auto + 360 twips → oneAndHalf（360/240 = 1.5）"""
        from thesis_worker.engine_v2.checkers import check_para_line_spacing_type
        para = self._make_para_with_line_rule('auto', 360)
        assert check_para_line_spacing_type(para, 'oneAndHalf') is None

    def test_exact_value_in_pt(self):
        """exact 30pt（line=600 twips）vs expected 30.0pt → None"""
        from thesis_worker.engine_v2.checkers import check_para_line_spacing_pt
        para = self._make_para_with_line_rule('exact', 600)
        assert check_para_line_spacing_pt(para, 30.0) is None

    def test_pt_tolerance_boundary(self):
        """容差 0.5pt 边界：actual 28.4 vs expected 28.0 → None；28.6 → issue

        判别力：568 twips = 28.4pt（差 0.4 < 0.5 容差），572 twips = 28.6pt（差 0.6 > 0.5）。
        若 _pt checker 错误用更宽松容差（如 1.0pt），第二条 case 会误通过。
        """
        from thesis_worker.engine_v2.checkers import check_para_line_spacing_pt
        para_a = self._make_para_with_line_rule('atLeast', 568)  # 28.4pt
        assert check_para_line_spacing_pt(para_a, 28.0) is None
        para_b = self._make_para_with_line_rule('atLeast', 572)  # 28.6pt
        out = check_para_line_spacing_pt(para_b, 28.0)
        assert out is not None

    def test_exact_lineRule_normalizes_to_exactly(self):
        """T-fix2: OOXML w:lineRule=exact 必须归一化为 type 字符串 'exactly'

        判别力：删除 checkers.py:419-420 'exact' → 'exactly' 归一化即挂。
        既有 test_exact_value_in_pt 仅间接覆盖；本 case 正反两侧直接断言。
        """
        from thesis_worker.engine_v2.checkers import check_para_line_spacing_type
        para = self._make_para_with_line_rule('exact', 600)  # 30pt exactly
        # 期望 'exactly' 应通过（无归一化则 actual='exact' ≠ expected → 误报 issue）
        assert check_para_line_spacing_type(para, 'exactly') is None
        # 反例：期望 'exact'（OOXML 原值）应 issue（对外统一暴露 'exactly'）
        out = check_para_line_spacing_type(para, 'exact')
        assert out is not None
        assert out['actual'] == 'exactly'
        assert out['expected'] == 'exact'


class TestCheckIndentPt:
    """T4.2: first_line/hanging_indent_pt checker"""

    def test_first_line_pt_match(self):
        from docx import Document
        from docx.shared import Pt
        from thesis_worker.engine_v2.checkers import check_para_first_line_indent_pt
        doc = Document()
        para = doc.add_paragraph('测试')
        para.paragraph_format.first_line_indent = Pt(24)
        assert check_para_first_line_indent_pt(para, 24.0) is None

    def test_first_line_pt_tolerance(self):
        from docx import Document
        from docx.shared import Pt
        from thesis_worker.engine_v2.checkers import check_para_first_line_indent_pt
        doc = Document()
        para = doc.add_paragraph('测试')
        para.paragraph_format.first_line_indent = Pt(24.4)
        assert check_para_first_line_indent_pt(para, 24.0) is None  # 0.4 < 0.5 容差
        para.paragraph_format.first_line_indent = Pt(24.6)
        out = check_para_first_line_indent_pt(para, 24.0)
        assert out is not None  # 0.6 超出

    def test_hanging_pt_match(self):
        from docx import Document
        from docx.shared import Pt
        from thesis_worker.engine_v2.checkers import check_para_hanging_indent_pt
        doc = Document()
        para = doc.add_paragraph('测试')
        # hanging 用 first_line_indent 负值表示
        para.paragraph_format.first_line_indent = Pt(-14)
        assert check_para_hanging_indent_pt(para, 14.0) is None


class TestCheckLetterSpacingPt:
    """T4.2: letter_spacing_pt checker"""

    def test_letter_spacing_pt_match(self):
        from docx import Document
        from docx.oxml.ns import qn
        from lxml import etree
        from thesis_worker.engine_v2.checkers import check_para_letter_spacing_pt
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run('测试')
        rPr = run._element.get_or_add_rPr()
        spacing_el = etree.SubElement(rPr, qn('w:spacing'))
        spacing_el.set(qn('w:val'), '20')  # 20 twips = 1pt
        assert check_para_letter_spacing_pt(para, 1.0) is None

    def test_letter_spacing_pt_mismatch(self):
        from docx import Document
        from docx.oxml.ns import qn
        from lxml import etree
        from thesis_worker.engine_v2.checkers import check_para_letter_spacing_pt
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run('测试')
        rPr = run._element.get_or_add_rPr()
        spacing_el = etree.SubElement(rPr, qn('w:spacing'))
        spacing_el.set(qn('w:val'), '40')  # 2pt
        out = check_para_letter_spacing_pt(para, 1.0)
        assert out is not None
        assert out['actual'] == 2.0

