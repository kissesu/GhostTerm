"""
@file: test_fixer.py
@description: v2 fixer 测试（含 reopen 验证）
              reopen 验证的意义：确保修改已写入磁盘，而非仅在内存中生效
@author: Atlas.oi
@date: 2026-04-18
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from thesis_worker.engine_v2.fixer import fix_v2


def make_bad_doc(tmp_path: Path) -> Path:
    """创建一个字体不合规的测试 docx（Calibri 14pt，正文应为 12pt）"""
    doc = Document()
    p = doc.add_paragraph('测试段')
    p.runs[0].font.name = 'Calibri'
    p.runs[0].font.size = Pt(14)
    path = tmp_path / 'bad.docx'
    doc.save(path)
    return path


class TestFixV2:
    def test_fix_font_size(self, tmp_path):
        """修复字号后 reopen 验证：磁盘上的 docx 已写入 12pt"""
        path = make_bad_doc(tmp_path)
        issue = {
            'rule_id': 'body_para',
            'attr': 'font.size_pt',
            'para_idx': 0,
        }
        value = {'font.size_pt': 12}
        result = fix_v2(str(path), issue, value)
        assert result['applied'] is True
        # reopen 验证：重新从磁盘读取，确认修改持久化
        doc2 = Document(path)
        assert doc2.paragraphs[0].runs[0].font.size.pt == 12.0

    def test_fix_marks_blue(self, tmp_path):
        """修复后 run 颜色应变为蓝色（0x0070C0）"""
        path = make_bad_doc(tmp_path)
        issue = {
            'rule_id': 'body_para',
            'attr': 'font.size_pt',
            'para_idx': 0,
        }
        value = {'font.size_pt': 12}
        fix_v2(str(path), issue, value)
        # reopen 验证：确认蓝色标记写入磁盘
        doc2 = Document(path)
        assert doc2.paragraphs[0].runs[0].font.color.rgb == RGBColor(0x00, 0x70, 0xC0)

    def test_fix_returns_diff(self, tmp_path):
        """返回的 diff 应包含前后值摘要"""
        path = make_bad_doc(tmp_path)
        issue = {
            'rule_id': 'body_para',
            'attr': 'font.size_pt',
            'para_idx': 0,
        }
        value = {'font.size_pt': 12}
        result = fix_v2(str(path), issue, value)
        assert 'font.size_pt' in result['diff']
        assert result['xml_changed'] == ['w:p[0]']

    def test_fix_unknown_attr_returns_not_applied(self, tmp_path):
        """未知属性应返回 applied=False，不修改文件"""
        path = make_bad_doc(tmp_path)
        issue = {
            'rule_id': 'body_para',
            'attr': 'nonexistent.attr',
            'para_idx': 0,
        }
        value = {'nonexistent.attr': 'something'}
        result = fix_v2(str(path), issue, value)
        assert result['applied'] is False

    def test_fix_missing_value_returns_not_applied(self, tmp_path):
        """value 中缺少对应 attr 时返回 applied=False"""
        path = make_bad_doc(tmp_path)
        issue = {
            'rule_id': 'body_para',
            'attr': 'font.size_pt',
            'para_idx': 0,
        }
        result = fix_v2(str(path), issue, {})
        assert result['applied'] is False

    def test_fix_invalid_para_idx_returns_not_applied(self, tmp_path):
        """para_idx 超出范围时返回 applied=False"""
        path = make_bad_doc(tmp_path)
        issue = {
            'rule_id': 'body_para',
            'attr': 'font.size_pt',
            'para_idx': 999,
        }
        value = {'font.size_pt': 12}
        result = fix_v2(str(path), issue, value)
        assert result['applied'] is False
