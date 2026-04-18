"""
@file: test_extractor.py
@description: 测试 extractor.extract_from_docx，覆盖 majority 提取、占位、空文档等场景
@author: Atlas.oi
@date: 2026-04-18
"""
import pytest
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from thesis_worker.extractor import extract_from_docx
from thesis_worker.rules import REGISTRY


def _make_docx(path: Path) -> str:
    """辅助：创建空 docx 并返回文件路径字符串"""
    doc = Document()
    doc.save(str(path))
    return str(path)


class TestExtractor:
    def test_extract_returns_full_rules_dict(self, tmp_path):
        """所有 REGISTRY 的 rule_id 都必须出现在 result['rules'] 中"""
        path = tmp_path / "empty.docx"
        doc = Document()
        doc.save(str(path))

        result = extract_from_docx(str(path))

        assert 'rules' in result
        assert 'evidence' in result
        # 11 条规则全部存在
        for rule_id in REGISTRY:
            assert rule_id in result['rules'], f"缺少规则: {rule_id}"
        # evidence 长度与 REGISTRY 一致
        assert len(result['evidence']) == len(REGISTRY)

    def test_extract_font_body_majority(self, tmp_path):
        """5 段宋体 12pt + 1 段黑体 14pt → majority 应提取宋体 12pt，confidence > 0.5"""
        path = tmp_path / "font_body.docx"
        doc = Document()

        # 正文段落：5 段宋体 12pt
        for _ in range(5):
            para = doc.add_paragraph()
            run = para.add_run("宋体正文段落")
            run.font.name = "宋体"
            run.font.size = Pt(12)

        # 少数：1 段黑体 14pt
        para = doc.add_paragraph()
        run = para.add_run("黑体段落")
        run.font.name = "黑体"
        run.font.size = Pt(14)

        doc.save(str(path))
        result = extract_from_docx(str(path))

        fb = result['rules']['font.body']
        assert fb['enabled'] is True
        assert fb['value']['family'] == '宋体'
        assert fb['value']['size_pt'] == 12

        # confidence 大于 0.5（宋体占多数）
        ev = next(e for e in result['evidence'] if e['rule_id'] == 'font.body')
        assert ev['confidence'] > 0.5

    def test_extract_font_h1_includes_bold(self, tmp_path):
        """Heading 1 段落：黑体 16pt bold → 提取 family=黑体 size_pt=16 bold=True"""
        path = tmp_path / "font_h1.docx"
        doc = Document()

        # 一级标题
        para = doc.add_heading("第一章 引言", level=1)
        for run in para.runs:
            run.font.name = "黑体"
            run.font.size = Pt(16)
            run.font.bold = True

        # 正文段落（不影响 h1 提取）
        body = doc.add_paragraph()
        body.add_run("正文内容")

        doc.save(str(path))
        result = extract_from_docx(str(path))

        fh = result['rules']['font.h1']
        assert fh['enabled'] is True
        assert fh['value']['family'] == '黑体'
        assert fh['value']['size_pt'] == 16
        assert fh['value']['bold'] is True

    def test_extract_paragraph_indent(self, tmp_path):
        """段落首行缩进 24pt（2 字 * 12pt）→ first_line_chars = 2"""
        path = tmp_path / "indent.docx"
        doc = Document()

        for _ in range(3):
            para = doc.add_paragraph()
            para.paragraph_format.first_line_indent = Pt(24)
            para.add_run("有缩进的正文段落")

        doc.save(str(path))
        result = extract_from_docx(str(path))

        pi = result['rules']['paragraph.indent']
        assert pi['enabled'] is True
        assert pi['value']['first_line_chars'] == 2

        ev = next(e for e in result['evidence'] if e['rule_id'] == 'paragraph.indent')
        assert ev['source_xml'] is not None
        assert 'w:ind' in ev['source_xml']

    def test_extract_cjk_ascii_space_unsupported(self, tmp_path):
        """cjk_ascii_space.extract 返回 None → enabled=False, confidence=0"""
        path = tmp_path / "cjk.docx"
        doc = Document()
        doc.add_paragraph("中文 English 混排")
        doc.save(str(path))

        result = extract_from_docx(str(path))

        cas = result['rules']['cjk_ascii_space']
        assert cas['enabled'] is False
        assert cas['value'] is None

        ev = next(e for e in result['evidence'] if e['rule_id'] == 'cjk_ascii_space')
        assert ev['confidence'] == 0.0

    def test_extract_other_rules_placeholder(self, tmp_path):
        """未实现 extract 的规则均为占位（enabled=False, confidence=0.0）"""
        path = tmp_path / "other.docx"
        doc = Document()
        doc.add_paragraph("普通段落")
        doc.save(str(path))

        result = extract_from_docx(str(path))

        # 这些规则在 P3 范围内不实现 extract，全为占位
        placeholder_rules = [
            'citation.format',
            'figure.caption_pos',
            'table.caption_pos',
            'chapter.new_page',
            'quote.style',
            'ai_pattern.check',
            'pagination',
        ]
        for rule_id in placeholder_rules:
            entry = result['rules'][rule_id]
            assert entry['enabled'] is False, f"{rule_id} 应为占位"
            assert entry['value'] is None, f"{rule_id} value 应为 None"

            ev = next(e for e in result['evidence'] if e['rule_id'] == rule_id)
            assert ev['confidence'] == 0.0, f"{rule_id} confidence 应为 0.0"

    def test_extract_empty_doc(self, tmp_path):
        """空文档（无段落内容）→ font.body 提取 None → enabled=False"""
        path = tmp_path / "empty2.docx"
        doc = Document()
        # 不添加任何段落内容
        doc.save(str(path))

        result = extract_from_docx(str(path))

        fb = result['rules']['font.body']
        # 空文档无 run 内容，font.body.extract 应返回 None
        assert fb['enabled'] is False
        assert fb['value'] is None
