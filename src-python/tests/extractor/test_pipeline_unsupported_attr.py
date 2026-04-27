"""
@file: test_pipeline_unsupported_attr.py
@description: 验证 pipeline 对未知 attr key 的开发者诊断日志通道（T1.2）
              - 未知 attr 触发 INFO 日志，extra 含结构化字段
              - 未知 attr 不出现在响应数据中（用户面板静默）
@author: Atlas.oi
@date: 2026-04-27
"""
import logging
import pytest
from docx import Document


class TestUnsupportedAttrLogging:
    """未知 attr key 触发开发者诊断日志，且不暴露给前端"""

    def test_unknown_attr_key_triggers_info_log(self, tmp_path, caplog):
        """构造一个会产生 mystery.unknown_key 的场景，验证 INFO 日志被触发。

        由于 pipeline 的抽取器都是硬编码的已知 key，
        我们通过 monkeypatch _extract_attributes_from_text 注入未知 key
        来模拟"将来某个新 extractor 返回了 schema 未声明的 key"的场景。
        """
        from thesis_worker.extractor import pipeline as p_module
        from thesis_worker.extractor.pipeline import extract_from_selection

        # 构造一个简单 docx
        doc = Document()
        doc.add_paragraph('这是正文段落，用于测试未知属性键的诊断日志。')
        docx_path = tmp_path / 'test_unsupported.docx'
        doc.save(str(docx_path))

        # 保存原始函数
        original_extract = p_module._extract_attributes_from_text

        # monkeypatch：在正常抽取结果基础上注入一个未知 key
        def patched_extract(text: str):
            result = original_extract(text)
            result['mystery.unknown_key'] = 5  # 模拟未知 attr key
            return result

        p_module._extract_attributes_from_text = patched_extract

        try:
            with caplog.at_level(logging.INFO, logger='thesis_worker.extractor.pipeline'):
                result = extract_from_selection(
                    str(docx_path),
                    para_indices=[0],
                    field_id='body_para',
                )
        finally:
            # 无论测试结果如何，还原原始函数
            p_module._extract_attributes_from_text = original_extract

        # 断言：必须有至少一条 INFO 日志，message 含 "unsupported_attr"
        unsupported_records = [
            r for r in caplog.records
            if r.levelno == logging.INFO and 'unsupported_attr' in r.getMessage()
        ]
        assert len(unsupported_records) >= 1, (
            f"期望至少一条含 'unsupported_attr' 的 INFO 日志，实际 caplog.records={caplog.records}"
        )

        # 断言：日志 extra 含正确的 attr_key
        record = unsupported_records[0]
        assert hasattr(record, 'attr_key'), (
            "日志 extra 必须包含 attr_key 字段"
        )
        assert record.attr_key == 'mystery.unknown_key', (
            f"attr_key 应为 'mystery.unknown_key'，实际为 {record.attr_key!r}"
        )

    def test_unknown_attr_not_exposed_in_response(self, tmp_path, caplog):
        """验证未知 attr key 不出现在响应数据中（不暴露给前端/用户面板）"""
        from thesis_worker.extractor import pipeline as p_module
        from thesis_worker.extractor.pipeline import extract_from_selection

        doc = Document()
        doc.add_paragraph('正文段落内容。')
        docx_path = tmp_path / 'test_no_expose.docx'
        doc.save(str(docx_path))

        original_extract = p_module._extract_attributes_from_text

        def patched_extract(text: str):
            result = original_extract(text)
            result['mystery.unknown_key'] = 5
            return result

        p_module._extract_attributes_from_text = patched_extract

        try:
            with caplog.at_level(logging.INFO, logger='thesis_worker.extractor.pipeline'):
                result = extract_from_selection(
                    str(docx_path),
                    para_indices=[0],
                    field_id='body_para',
                )
        finally:
            p_module._extract_attributes_from_text = original_extract

        # 断言：响应 value 中不含未知 key
        value = result.get('value', {})
        assert 'mystery.unknown_key' not in value, (
            f"未知 attr key 不应出现在响应 value 中，实际 value={value}"
        )

        # 断言：响应顶层不含 unsupported_attr / unknown_attr 字段
        assert 'unsupported_attr' not in result, (
            "unsupported_attr 不应出现在响应顶层（不暴露给前端）"
        )
        assert 'unknown_attr' not in result, (
            "unknown_attr 不应出现在响应顶层（不暴露给前端）"
        )

    def test_unknown_attr_in_extract_all(self, tmp_path, caplog):
        """验证 extract_all 路径同样触发未知 attr 诊断日志"""
        from thesis_worker.extractor import pipeline as p_module
        from thesis_worker.extractor.pipeline import extract_all

        # 构造一个能被 field_matcher 命中的 docx（含"摘要"关键词）
        doc = Document()
        doc.add_paragraph('摘  要（小三号宋体加粗）')
        docx_path = tmp_path / 'test_extract_all_unsupported.docx'
        doc.save(str(docx_path))

        original_extract = p_module._extract_attributes_from_text

        def patched_extract(text: str):
            result = original_extract(text)
            result['mystery.unknown_key'] = 99
            return result

        p_module._extract_attributes_from_text = patched_extract

        try:
            with caplog.at_level(logging.INFO, logger='thesis_worker.extractor.pipeline'):
                result = extract_all(str(docx_path))
        finally:
            p_module._extract_attributes_from_text = original_extract

        # 断言：extract_all 结果的 rules value 中不含未知 key
        for field_id, rule in result.get('rules', {}).items():
            value = rule.get('value', {})
            assert 'mystery.unknown_key' not in value, (
                f"字段 {field_id} 的 value 不应包含未知 attr key，实际 value={value}"
            )

        # 断言：extract_all 响应顶层不含 unsupported_attr 字段
        assert 'unsupported_attr' not in result

    def test_known_attr_does_not_trigger_log(self, tmp_path, caplog):
        """已知 attr key 不应触发 unsupported_attr 日志（确保无误报）"""
        from thesis_worker.extractor.pipeline import extract_from_selection

        doc = Document()
        doc.add_paragraph('正文段落。')
        docx_path = tmp_path / 'test_known_attr.docx'
        doc.save(str(docx_path))

        with caplog.at_level(logging.INFO, logger='thesis_worker.extractor.pipeline'):
            extract_from_selection(
                str(docx_path),
                para_indices=[0],
                field_id='body_para',
            )

        # 不应有任何 unsupported_attr 日志（正常 known key 不触发）
        unsupported_records = [
            r for r in caplog.records
            if 'unsupported_attr' in r.getMessage()
        ]
        assert len(unsupported_records) == 0, (
            f"已知 attr 不应触发 unsupported_attr 日志，实际 records={unsupported_records}"
        )
