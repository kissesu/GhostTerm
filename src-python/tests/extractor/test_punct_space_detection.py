"""
@file: test_punct_space_detection.py
@description: mixed_script.punct_space_after 全文启发式检测测试
@author: Atlas.oi
@date: 2026-04-18
"""
from docx import Document
from thesis_worker.extractor.pipeline import _detect_punct_space_after


def test_all_punct_followed_by_space_returns_true():
    """规范：所有英文标点后都空格 → True"""
    doc = Document()
    doc.add_paragraph('This is fine. Another sentence; semicolon: colon! exclamation? yes.')
    result = _detect_punct_space_after(doc)
    assert result is True


def test_all_punct_no_space_returns_false():
    """全部无空格 → False"""
    doc = Document()
    doc.add_paragraph('Bad.writing;like:this!everywhere?yes.no.end.')
    result = _detect_punct_space_after(doc)
    assert result is False


def test_mixed_majority_space_returns_true():
    """绝大多数有空格（>= 2:1） → True"""
    doc = Document()
    # 4 个有空格 + 1 个无空格 → 4 >= 2*1 → True
    doc.add_paragraph('Word. More. Lots. Also. Oops.no')
    result = _detect_punct_space_after(doc)
    assert result is True


def test_few_punct_returns_none():
    """样本不足 (< 3) → None

    'Hello world. Bye.' 中，第一个 '.' 后紧跟空格 → space_after=1；
    第二个句末 '.' 后是字符串结尾，两个 lookahead 都无法匹配（既非 \\s 也非 \\S）→ 不计数。
    total=1 < 3，返回 None。
    """
    doc = Document()
    doc.add_paragraph('Hello world. Bye.')
    result = _detect_punct_space_after(doc)
    assert result is None


def test_no_ascii_punct_returns_none():
    """纯中文无 ASCII 标点 → None"""
    doc = Document()
    doc.add_paragraph('这是一段纯中文内容，不含任何 ASCII 标点。')
    result = _detect_punct_space_after(doc)
    assert result is None


def test_boundary_2_to_1_returns_true():
    """边界：space_after=2, no_space=1（2:1 临界）→ True（2 >= 2*1）"""
    doc = Document()
    # 'A. B. C.end'：前两个 '.' 后空格（space_after=2）；第三个 '.' 后紧跟字母（no_space=1）
    # 末尾无额外标点
    doc.add_paragraph('A. B. C.end')
    result = _detect_punct_space_after(doc)
    assert result is True


def test_all_no_space_zero_vs_three_returns_false():
    """边界：space_after=0, no_space=3（全违规）→ False"""
    doc = Document()
    doc.add_paragraph('a.b,c;d')  # 三个标点后都是非空白字符
    result = _detect_punct_space_after(doc)
    assert result is False
