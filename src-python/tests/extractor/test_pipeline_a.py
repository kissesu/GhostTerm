"""
@file: test_pipeline_a.py
@description: extract_all 在 Template A (括号型) 上的集成测试；
              T3.1: 追加 Section 级属性抽取断言（space_before_pt / gutter / print_mode）
@author: Atlas.oi
@date: 2026-04-28
"""
from pathlib import Path
import pytest
from thesis_worker.extractor.pipeline import extract_all, extract_from_selection

FIXTURES = Path(__file__).parent.parent / 'fixtures'


class TestExtractAll:
    def test_returns_rules_dict(self):
        result = extract_all(str(FIXTURES / 'spec_template_a.docx'))
        assert 'rules' in result
        assert isinstance(result['rules'], dict)

    def test_returns_evidence_list(self):
        result = extract_all(str(FIXTURES / 'spec_template_a.docx'))
        assert 'evidence' in result
        assert isinstance(result['evidence'], list)

    def test_finds_title_zh(self):
        result = extract_all(str(FIXTURES / 'spec_template_a.docx'))
        # 模板 p0 "毕业论文（设计）题目（三号黑体，居中。理工农医类专业用）"
        assert 'title_zh' in result['rules']
        value = result['rules']['title_zh']['value']
        assert value.get('font.size_pt') == 16  # 三号 = 16pt
        assert value.get('para.align') == 'center'

    def test_finds_abstract_zh_title(self):
        result = extract_all(str(FIXTURES / 'spec_template_a.docx'))
        # 模板 p1 "摘  要（小三号宋体加粗，中间空2个字符，居中）"
        assert 'abstract_zh_title' in result['rules']
        value = result['rules']['abstract_zh_title']['value']
        assert value.get('font.cjk') == '宋体'
        assert value.get('font.size_pt') == 15  # 小三 = 15pt
        assert value.get('font.bold') is True
        assert value.get('para.align') == 'center'


class TestExtractFromSelection:
    def test_single_paragraph(self):
        result = extract_from_selection(
            str(FIXTURES / 'spec_template_a.docx'),
            para_indices=[1],
            field_id='abstract_zh_title',
        )
        assert result['field_id'] == 'abstract_zh_title'
        assert result['value'].get('font.cjk') == '宋体'
        assert result['value'].get('font.size_pt') == 15
        assert result['confidence'] > 0.8

    def test_empty_para_returns_low_confidence(self):
        # p6 是空段落，文本无内容，样式属性至多继承 1 个（行距），置信度 <= 0.5
        result = extract_from_selection(
            str(FIXTURES / 'spec_template_a.docx'),
            para_indices=[6],
            field_id='title_en',
        )
        assert result['confidence'] <= 0.5


class TestT31SectionAttrsExtraction:
    """T3.1: 验证 extract_all 能正确抽取 Section 级属性并写入 page_margin 字段

    构造临时 docx，设定 space_before/gutter/print_mode，
    验证 extract_all 返回的 rules['page_margin']['value'] 包含对应属性。
    """

    def test_space_before_pt_extracted(self, tmp_path):
        """para.space_before_pt 应从段落样式中抽出并正确路由到 chapter_title 字段

        路由验证：构造含 field_matcher 的 chapter_title 关键词"第一章"的段落，
        断言 extract_all 将 para.space_before_pt 写入 rules['chapter_title']['value']，
        而非仅检测某字段里"存在"该 key（后者无法区分路由是否正确）。
        """
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        # 文本含 chapter_title 关键词"第一章"→ field_matcher 应路由到 chapter_title
        p = doc.add_paragraph('第一章 引言（一级标题）')
        p.paragraph_format.space_before = Pt(12)
        docx_path = tmp_path / 'space_before_test.docx'
        doc.save(str(docx_path))

        result = extract_all(str(docx_path))
        # 定向断言：必须路由到 chapter_title 字段
        assert 'chapter_title' in result['rules'], (
            f'chapter_title 字段未被命中；实际 rules keys={list(result["rules"].keys())}'
        )
        value = result['rules']['chapter_title']['value']
        assert 'para.space_before_pt' in value, (
            f'para.space_before_pt 未写入 chapter_title.value；value={value}'
        )
        assert value['para.space_before_pt'] == pytest.approx(12.0, abs=0.5)

    def test_section_attrs_in_page_margin(self, tmp_path):
        """Section 级属性（gutter/header_offset/footer_offset/print_mode）应写入 page_margin"""
        from docx import Document
        from docx.shared import Cm
        doc = Document()
        doc.add_paragraph('测试内容')
        # 设定装订线 0.5cm、页眉距 1.5cm、页脚距 1.75cm
        doc.sections[0].gutter = Cm(0.5)
        doc.sections[0].header_distance = Cm(1.5)
        doc.sections[0].footer_distance = Cm(1.75)
        docx_path = tmp_path / 'section_attrs_test.docx'
        doc.save(str(docx_path))

        result = extract_all(str(docx_path))
        assert 'page_margin' in result['rules'], 'page_margin 字段应由 Section 级属性触发写入'
        value = result['rules']['page_margin']['value']
        assert 'page.margin_gutter_cm' in value
        assert value['page.margin_gutter_cm'] == pytest.approx(0.5, abs=0.05)
        assert 'page.header_offset_cm' in value
        assert value['page.header_offset_cm'] == pytest.approx(1.5, abs=0.05)
        assert 'page.footer_offset_cm' in value
        assert value['page.footer_offset_cm'] == pytest.approx(1.75, abs=0.05)
        # 默认文档无 w:evenAndOddHeaders → single
        assert value.get('page.print_mode') == 'single'

    def test_double_print_mode_extracted(self, tmp_path):
        """注入 w:evenAndOddHeaders 后，print_mode 应抽取为 'double'"""
        from docx import Document
        from docx.oxml.ns import qn
        from lxml import etree
        doc = Document()
        doc.add_paragraph('测试内容')
        # 注入 w:evenAndOddHeaders → 双面打印模式
        settings_el = doc.settings.element
        etree.SubElement(settings_el, qn('w:evenAndOddHeaders'))
        docx_path = tmp_path / 'double_print_test.docx'
        doc.save(str(docx_path))

        result = extract_all(str(docx_path))
        assert 'page_margin' in result['rules']
        value = result['rules']['page_margin']['value']
        assert value.get('page.print_mode') == 'double'


class TestExtractFromSelectionWithSelectedText:
    """Task 4：selected_text 按句选取路径的测试"""

    def test_selected_text_not_found_falls_back_to_full_paragraph(self):
        """selected_text 在段落中找不到时，回退到全段提取，不报错"""
        result = extract_from_selection(
            str(FIXTURES / 'spec_template_a.docx'),
            para_indices=[1],
            field_id='abstract_zh_title',
            selected_text='这段文字根本不存在于文档中',
        )
        # 回退全段提取：field_id 正确，结果非空，不抛异常
        assert result['field_id'] == 'abstract_zh_title'
        assert isinstance(result['value'], dict)
        assert isinstance(result['confidence'], float)

    def test_selected_text_restricts_to_matching_portion(self):
        """传入 selected_text 时，抽取结果来自该文本所在段落（端到端验证 selected_text 路径走通）

        验证策略：
        - 用段落中确实存在的文本片段调用；若能定位 run，结果不报错且有合理置信度
        - 无法用夹具文件精确控制 run 内容，因此只验证路径不崩溃、返回结构正确
        """
        # p1 文本包含 "摘  要"，取其前半段作为 selected_text
        result = extract_from_selection(
            str(FIXTURES / 'spec_template_a.docx'),
            para_indices=[1],
            field_id='abstract_zh_title',
            selected_text='摘',
        )
        assert result['field_id'] == 'abstract_zh_title'
        assert isinstance(result['value'], dict)
        # 段落级属性（行距/对齐等）应仍被抽取（来自 _extract_para_level_attrs）
        # p1 = 摘要标题行，居中对齐
        assert result['value'].get('para.align') == 'center'

class TestT32TableAttrsExtraction:
    """T3.2: 验证 extract_all 能正确抽取 table.* 属性并写入 table_header 字段"""

    def test_three_line_table_attrs_in_table_header(self, tmp_path):
        """构造含三线表 tblBorders 的 docx，断言 extract_all 将 4 个 table.* attr
        写入 rules['table_header']['value']。

        三线表规格（eighth-points）：
        - top=12 → 1.5pt（规范上线）
        - bottom=12 → 1.5pt（规范下线）
        - insideH=4 → 0.5pt（规范表头下线）
        无 insideV → 视为 0，三线表判定 True。
        """
        from docx import Document
        from docx.oxml.ns import qn
        from lxml import etree
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        tbl_pr = table._element.find(qn('w:tblPr'))
        borders = etree.SubElement(tbl_pr, qn('w:tblBorders'))
        for tag, sz in [('top', 12), ('bottom', 12), ('insideH', 4)]:
            el = etree.SubElement(borders, qn(f'w:{tag}'))
            el.set(qn('w:sz'), str(sz))
            el.set(qn('w:val'), 'single')
        docx_path = tmp_path / 'three_line_table.docx'
        doc.save(str(docx_path))

        result = extract_all(str(docx_path))
        assert 'table_header' in result['rules'], (
            f'table_header 字段应由表格属性触发写入；实际 rules keys={list(result["rules"].keys())}'
        )
        value = result['rules']['table_header']['value']

        # 全部 4 个 table.* attr 必须存在
        assert 'table.is_three_line' in value, f'缺 table.is_three_line；value={value}'
        assert 'table.border_top_pt' in value, f'缺 table.border_top_pt；value={value}'
        assert 'table.border_bottom_pt' in value, f'缺 table.border_bottom_pt；value={value}'
        assert 'table.header_border_pt' in value, f'缺 table.header_border_pt；value={value}'

        # 值正确性：eighth-points → pt 换算（sz/8）
        assert value['table.is_three_line'] is True
        assert abs(value['table.border_top_pt'] - 1.5) < 0.01
        assert abs(value['table.border_bottom_pt'] - 1.5) < 0.01
        assert abs(value['table.header_border_pt'] - 0.5) < 0.01

    def test_no_table_does_not_inject_table_attrs(self, tmp_path):
        """文档无表格时，extract_all 不应在 table_header 中注入 table.* attr。

        table_header 字段可能由段落路径被命中（如规范文档含"表头"关键词），
        但 value 中不应含 table.* 前缀的 attr。
        """
        from docx import Document
        doc = Document()
        doc.add_paragraph('无表格内容')
        docx_path = tmp_path / 'no_table.docx'
        doc.save(str(docx_path))

        result = extract_all(str(docx_path))
        # 若 table_header 字段存在（由段落路径命中），其 value 不应含 table.* attr
        if 'table_header' in result['rules']:
            value = result['rules']['table_header']['value']
            for key in value:
                assert not key.startswith('table.'), (
                    f'无表格文档的 table_header 不应含 table.* attr；实际 key={key}'
                )


    def test_extract_from_selection_with_selected_text_isolates_run(self, tmp_path):
        """构造 2-run 段落：run0 加粗，run1 不加粗；selected_text 只命中 run1。

        验证返回 value 中 font.bold 缺席（而非 True）——证明 run 隔离生效。
        若该测试通过，说明 _read_run_list_style_attrs 确实只读了匹配 run 的样式。

        注意：run 文本刻意用中性词（"标题："/"内容文本"），
        避免触发 _extract_attributes_from_text 的关键词识别污染断言。
        """
        from docx import Document
        doc = Document()
        para = doc.add_paragraph()
        # run 0：加粗"标题："
        run_bold = para.add_run('标题：')
        run_bold.bold = True
        # run 1：不加粗"内容文本"
        para.add_run('内容文本')
        # run_plain.bold 保持默认 None（未设）

        docx_path = tmp_path / 'two_run.docx'
        doc.save(str(docx_path))

        # selected_text 只匹配 run1 的文字
        result = extract_from_selection(
            str(docx_path),
            para_indices=[0],
            field_id='body_para',
            selected_text='内容文本',
        )

        # 关键断言：font.bold 不应出现（因 selected_text 只覆盖 run1，run1 无 bold）
        assert 'font.bold' not in result['value'], (
            f"run 隔离失败：font.bold 不应被报告，但实际 value={result['value']}"
        )


# ───────────────────────────────────────────────
# T3.3: numbering.* 集成测试（_read_numbering_styles via extract_all）
# ───────────────────────────────────────────────

class TestNumberingStylesIntegration:
    """extract_all 中 numbering 风格抽取的端到端集成测试。

    不依赖 spec_template_a.docx（该文档无图题，无法触发 numbering 路径），
    而是用 tmp_path 构造含图题/公式编号的合成文档进行验证。
    """

    def test_continuous_figure_style_in_extract_all(self, tmp_path):
        """5 个连续图题 → extract_all 后 figure_caption.value 含 figure_style='continuous'。

        判别力：构造 5 个"图N"格式图题，continuous 多数票必然命中。
        删除 pipeline 中 _read_numbering_styles 调用后，figure_caption.value 中
        不含 numbering.figure_style 键，断言 key in value 必挂。
        """
        from docx import Document as DocxDocument
        doc = DocxDocument()
        for i in range(1, 6):
            doc.add_paragraph(f'图{i} 这是图题文本')
        docx_path = tmp_path / 'test_fig_continuous.docx'
        doc.save(str(docx_path))

        result = extract_all(str(docx_path))
        rules = result['rules']

        # figure_caption 字段必须存在（由 numbering 全文扫描写入）
        assert 'figure_caption' in rules, f'figure_caption 未在 rules 中，rules keys={list(rules.keys())}'
        value = rules['figure_caption']['value']
        assert 'numbering.figure_style' in value, f'figure_style 未写入 value={value}'
        assert value['numbering.figure_style'] == 'continuous'

    def test_chapter_based_figure_style_in_extract_all(self, tmp_path):
        """5 个章节式图题 → extract_all 后 figure_caption.value 含 figure_style='chapter_based'。

        判别力：章节式（图N-M）与连续式（图N）各自命中路径互斥，删除判断分支后
        全部落入另一路径，断言 actual=='chapter_based' 必挂。
        """
        from docx import Document as DocxDocument
        doc = DocxDocument()
        for chap, seq in [(1, 1), (1, 2), (2, 1), (2, 2), (3, 1)]:
            doc.add_paragraph(f'图{chap}-{seq} 这是章节图题')
        docx_path = tmp_path / 'test_fig_chapter.docx'
        doc.save(str(docx_path))

        result = extract_all(str(docx_path))
        rules = result['rules']

        assert 'figure_caption' in rules
        value = rules['figure_caption']['value']
        assert value.get('numbering.figure_style') == 'chapter_based'

    def test_formula_style_in_extract_all(self, tmp_path):
        """5 个连续公式编号 → extract_all 后 formula_block.value 含 formula_style='continuous'。

        判别力：连续与章节式路径互斥，且仅在 >=2 个样本时写入。
        删除 numbering 合入逻辑后 formula_block.value 中无该 key，断言挂。
        """
        from docx import Document as DocxDocument
        doc = DocxDocument()
        for i in range(1, 6):
            doc.add_paragraph(f'这是公式 E = mc² ({i})')
        docx_path = tmp_path / 'test_formula_continuous.docx'
        doc.save(str(docx_path))

        result = extract_all(str(docx_path))
        rules = result['rules']

        assert 'formula_block' in rules, f'formula_block 未写入，rules keys={list(rules.keys())}'
        value = rules['formula_block']['value']
        assert 'numbering.formula_style' in value, f'formula_style 未写入 value={value}'
        assert value['numbering.formula_style'] == 'continuous'

    def test_fullwidth_space_in_figure_caption(self, tmp_path):
        """图题含全角空格 U+3000 时，regex [\s　] 分支应正确匹配并推断 figure_style。

        判别力：构造 2 个"图　1"/"图　2"格式图题（全角空格代替半角），
        验证 _read_numbering_styles 的 _RE_FIG_CONTINUOUS 能命中全角空格变体。
        删除 regex 中的 \\u3000 后，两个图题均漏匹配 → total_fig=0 < 2 → key 缺席，断言挂。
        """
        from docx import Document as DocxDocument
        doc = DocxDocument()
        # 全角空格 U+3000 位于图号与序号之间（"图　1 测试图"）
        doc.add_paragraph('图　1 测试图甲')
        doc.add_paragraph('图　2 测试图乙')
        docx_path = tmp_path / 'test_fig_fullwidth_space.docx'
        doc.save(str(docx_path))

        result = extract_all(str(docx_path))
        rules = result['rules']

        # figure_caption 字段必须存在（全角空格图题应被 regex 捕获）
        assert 'figure_caption' in rules, (
            f'全角空格图题未触发 figure_caption 写入；rules keys={list(rules.keys())}'
        )
        value = rules['figure_caption']['value']
        assert 'numbering.figure_style' in value, (
            f'全角空格路径漏抓 figure_style；value={value}'
        )
        assert value['numbering.figure_style'] == 'continuous'

    # ── W1 修复：APA 年份引用不应触发 formula 推断 ──────────────────────

    def test_apa_year_citations_not_counted_as_formula(self, tmp_path):
        """W1 修复：文档含 APA 年份引用 (2020)/(2021) 不应触发 formula_style 推断。

        判别力：构造含 5 个 APA 年份括号但无等号/数学符号的段落，
        旧版扫全文会将 (2020)(2021)(2022)(2023)(2024) 误计 5 票 continuous，
        W1 修复后这些段落被 _looks_like_formula_paragraph 过滤，票数 < 2，
        formula_block 字段不被写入（或不含 formula_style key）。

        判别力：删除 _looks_like_formula_paragraph 过滤后，5 个 APA 引用
        会触发 formula_block 写入且值为 'continuous'，断言"不存在"挂。
        """
        from docx import Document as DocxDocument
        doc = DocxDocument()
        # 5 个含 APA 年份引用的段落（无等号/数学符号/短编号）
        for year in [2020, 2021, 2022, 2023, 2024]:
            doc.add_paragraph(f'Smith et al. ({year}) 研究表明，该方法有效。')
        docx_path = tmp_path / 'test_apa_no_formula.docx'
        doc.save(str(docx_path))

        result = extract_all(str(docx_path))
        rules = result['rules']

        # formula_block 字段不应写入（票数不足），或即使写入也不含 formula_style
        if 'formula_block' in rules:
            value = rules['formula_block']['value']
            assert 'numbering.formula_style' not in value, (
                f'W1 修复失效：APA 年份引用不应触发 formula_style 推断；value={value}'
            )

    # ── W2 修复：正文图引用不应计票，只对图题开头计票 ──────────────────

    def test_inline_figure_ref_not_counted_as_caption(self, tmp_path):
        """W2 修复：正文"如图1所示"/"见图2"不应参与图编号计票。

        构造：10 个含"如图N所示"的正文段落（大量连续引用）
        + 2 个段落开头"图1 真实图题"/"图2 另一图题"
        期望 figure_style='continuous'（来自 2 个图题），不被 10 个正文引用污染计票。

        判别力：删除 _RE_FIG_CAPTION_START.match 过滤后，所有正文"如图N"也计票，
        连续票数 10 >> 图题 2，结果相同（continuous），但若图题是章节式而正文引用是连续式
        则会被误判——此 case 验证过滤逻辑确实限定了开头匹配。

        精确判别力 case：2 个章节式图题 + 10 个连续式正文引用：
        无过滤 → continuous（10 >> 2），有过滤 → chapter_based（2 章节 vs 0 连续）。
        """
        from docx import Document as DocxDocument
        doc = DocxDocument()
        # 2 个真实图题（章节式，段落开头）
        doc.add_paragraph('图1-1 系统架构图')
        doc.add_paragraph('图1-2 数据流程图')
        # 10 个正文引用（连续式，位于段落中间）
        for i in range(1, 11):
            doc.add_paragraph(f'如图{i}所示，系统的主要模块包括以下部分。')
        docx_path = tmp_path / 'test_caption_only_counted.docx'
        doc.save(str(docx_path))

        result = extract_all(str(docx_path))
        rules = result['rules']

        # figure_caption 应存在，且 figure_style 应为 'chapter_based'（来自 2 个图题）
        # 若正文引用也被计票，连续 10 > 章节 2，会误判 'continuous'
        assert 'figure_caption' in rules, (
            f'W2：figure_caption 未写入；rules keys={list(rules.keys())}'
        )
        value = rules['figure_caption']['value']
        assert value.get('numbering.figure_style') == 'chapter_based', (
            f"W2 修复失效：正文引用被计票，误判为连续式；实际 value={value}"
        )

    # ── W3 修复：数字子图 (1)(2)(3) 应推断为 1_2_3 ──────────────────

    def test_numeric_subfigure_style_detected(self, tmp_path):
        """W3 修复：5 个含数字子图标记的图题，应推断 subfigure_style='1_2_3'。

        判别力：构造"图1(1)"/"图1(2)"等数字子图，
        旧版只有 _RE_SUBFIG_LETTER（匹配字母），数字子图 actual=None，
        W3 修复后 actual='1_2_3'。

        删除 _RE_SUBFIG_NUMBER 正则后 number_count=0 < 2，
        不写 subfigure_style，断言 value 中不含该 key，测试挂。
        """
        from docx import Document as DocxDocument
        doc = DocxDocument()
        # 5 个含数字子图标记的图题
        for chap, sub in [(1, 1), (1, 2), (1, 3), (2, 1), (2, 2)]:
            doc.add_paragraph(f'图{chap}-{sub}({sub}) 数字子图示例')
        docx_path = tmp_path / 'test_numeric_subfig.docx'
        doc.save(str(docx_path))

        result = extract_all(str(docx_path))
        rules = result['rules']

        assert 'figure_caption' in rules, (
            f'W3：figure_caption 未写入；rules keys={list(rules.keys())}'
        )
        value = rules['figure_caption']['value']
        assert 'numbering.subfigure_style' in value, (
            f'W3 修复失效：数字子图未被识别；value={value}'
        )
        assert value['numbering.subfigure_style'] == '1_2_3', (
            f"W3：应为 '1_2_3' 不是 'a_b_c'；实际 value={value}"
        )


# ───────────────────────────────────────────────
# W6 修复：_log_and_filter_unsupported 字段级白名单过滤测试
# ───────────────────────────────────────────────

class TestW6FieldWhitelistFilter:
    """W6 修复：越界 attr（已知但不属于当前字段白名单）应被过滤。

    原有 _log_and_filter_unsupported 只按全局 _KNOWN_ATTR_KEYS 过滤，
    导致 section/table/numbering 全文注入可能把不属于该字段的 attr 写入。
    W6 修复后改为字段级白名单，越界 attr 被剔除。
    """

    def test_out_of_whitelist_attr_filtered_for_field(self, tmp_path):
        """注入 table.* attr 到 page_margin 字段，断言被字段白名单过滤掉。

        page_margin 的 applicable_attributes 不含 table.is_three_line，
        所以即使 attrs 里有这个 known attr，也应被过滤。

        判别力：删除字段级白名单逻辑（回退到全局 _KNOWN_ATTR_KEYS）后，
        table.is_three_line 是 known key，不会被过滤，断言 'not in value' 挂。
        """
        from thesis_worker.extractor import pipeline as p_module
        from thesis_worker.extractor.pipeline import _log_and_filter_unsupported

        # 构造越界 attrs：table.is_three_line 是已知 key，但不属于 page_margin 字段
        attrs_with_cross_field = {
            'page.margin_top_cm': 3.0,    # page_margin 白名单内
            'table.is_three_line': True,   # 越界：已知 key 但属于 table_header 字段
        }

        # 对 page_margin 字段调用过滤
        result = _log_and_filter_unsupported(
            attrs_with_cross_field,
            spec_file='test.docx',
            field_id='page_margin',
            context_snippet='[test context]',
        )

        # page.margin_top_cm 属于 page_margin 白名单，应保留
        assert 'page.margin_top_cm' in result, (
            f'page.margin_top_cm 不应被过滤（在白名单内）；result={result}'
        )
        # table.is_three_line 是越界 attr（known but not in page_margin），应被过滤
        assert 'table.is_three_line' not in result, (
            f'W6 修复失效：table.is_three_line 越界 attr 应被字段级白名单过滤；result={result}'
        )

    def test_known_attr_in_correct_field_preserved(self, tmp_path):
        """在正确字段内的 attr 不应被过滤（无误杀）。

        page_margin 字段注入 page.* attr，全部应保留。
        删除字段白名单也不影响本 case——测试只验证正路径不误杀。
        （判别力来自 test_out_of_whitelist_attr_filtered_for_field）
        """
        from thesis_worker.extractor.pipeline import _log_and_filter_unsupported

        attrs = {
            'page.margin_top_cm': 3.0,
            'page.margin_bottom_cm': 2.5,
            'page.margin_left_cm': 3.0,
            'page.margin_right_cm': 3.0,
        }

        result = _log_and_filter_unsupported(
            attrs,
            spec_file='test.docx',
            field_id='page_margin',
            context_snippet='[test context]',
        )

        # 全部 4 个 attr 属于 page_margin 白名单，应全部保留
        for key in attrs:
            assert key in result, f'{key!r} 不应被过滤（在 page_margin 白名单内）；result={result}'


class TestNullSectionAttrs:
    """修复回归：python-docx 未显式设置的 section 属性返回 None，真实 docx 必须不崩溃。

    T3.1 加入的 gutter/header_distance/footer_distance 假设非 None，
    textutil 从 .doc 转换的文档会触发 AttributeError: 'NoneType'.cm，
    本 class 验证修复后的 None 路径不抛异常。
    """

    def test_extract_all_handles_none_section_attrs(self, tmp_path):
        """构造未设任何 section attr 的 docx，extract_all 不应崩溃。

        判别力：临时把 _read_section_attrs 里的 None 守卫删掉，本 case 会抛
        AttributeError('NoneType' object has no attribute 'cm')。
        """
        from docx import Document
        doc = Document()
        doc.add_paragraph('测试内容：第一章 引言')
        # 不主动设置 gutter / header_distance / footer_distance，
        # 保持 python-docx 默认状态（新建文档这三项通常为 None）
        docx_path = tmp_path / 'no_section_attrs.docx'
        doc.save(str(docx_path))

        # 核心断言：不应抛 AttributeError
        result = extract_all(str(docx_path))
        assert isinstance(result, dict), '返回值应为 dict'
        assert 'rules' in result, '应包含 rules 键'

    def test_extract_all_handles_paragraph_with_no_spacing(self, tmp_path):
        """构造未设 space_before/space_after 的段落，extract_all 不应崩溃。

        判别力：若 _read_paragraph_style_attrs 在 None 守卫外写入 _pt，
        本 case 会抛 AttributeError('NoneType'.pt)。
        """
        from docx import Document
        doc = Document()
        # 添加纯文本段落，不设任何段间距
        para = doc.add_paragraph('摘要：本文研究了相关问题。')
        # 不调用 para.paragraph_format.space_before = ... 等
        docx_path = tmp_path / 'no_para_spacing.docx'
        doc.save(str(docx_path))

        result = extract_all(str(docx_path))
        assert isinstance(result, dict), '返回值应为 dict'
        assert 'rules' in result, '应包含 rules 键'


class TestT3PipelineIntegration:
    """T3.5: pipeline._extract_attributes_from_text 集成 5 个新抽取"""

    def test_para_spacing_pt_in_pipeline(self):
        from thesis_worker.extractor.pipeline import _extract_attributes_from_text
        out = _extract_attributes_from_text('段前 6 磅，段后 3 磅')
        assert out.get('para.space_before_pt') == 6.0
        assert out.get('para.space_after_pt') == 3.0

    def test_indent_chars_in_pipeline(self):
        from thesis_worker.extractor.pipeline import _extract_attributes_from_text
        out = _extract_attributes_from_text('首行缩进 2 字符')
        assert out.get('para.first_line_indent_chars') == 2.0

    def test_line_spacing_at_least_in_pipeline(self):
        from thesis_worker.extractor.pipeline import _extract_attributes_from_text
        out = _extract_attributes_from_text('行距：最小值 28 磅')
        assert out.get('para.line_spacing_type') == 'atLeast'
        assert out.get('para.line_spacing_pt') == 28.0

    def test_letter_spacing_pt_in_pipeline(self):
        from thesis_worker.extractor.pipeline import _extract_attributes_from_text
        out = _extract_attributes_from_text('字符间距 加宽 1 磅')
        assert out.get('para.letter_spacing_pt') == 1.0

    def test_table_three_line_in_pipeline(self):
        from thesis_worker.extractor.pipeline import _extract_attributes_from_text
        out = _extract_attributes_from_text('三线表，上下表线 1.5 磅，表头下线 0.5 磅')
        assert out.get('table.is_three_line') is True
        assert out.get('table.border_top_pt') == 1.5
        assert out.get('table.header_border_pt') == 0.5

    def test_combined_paragraph(self):
        from thesis_worker.extractor.pipeline import _extract_attributes_from_text
        # 单段含多种规范
        out = _extract_attributes_from_text('小三号黑体，居中，段前 6 磅，1.5 倍行距')
        assert out.get('font.size_pt') == 15.0   # 小三 = 15pt
        assert out.get('font.cjk') == '黑体'
        assert out.get('para.align') == 'center'
        assert out.get('para.space_before_pt') == 6.0
        assert out.get('para.line_spacing_type') == 'oneAndHalf'
        assert out.get('para.line_spacing') == 1.5


class TestT4PipelineParaStyle:
    """T4.3: _read_paragraph_style_attrs 同时写 _pt 兄弟 attr"""

    def test_first_line_indent_pt_in_attrs(self):
        from docx import Document
        from docx.shared import Pt
        from thesis_worker.extractor.pipeline import _read_paragraph_style_attrs
        doc = Document()
        para = doc.add_paragraph('测试')
        para.paragraph_format.first_line_indent = Pt(24)
        attrs = _read_paragraph_style_attrs(para)
        assert attrs.get('para.first_line_indent_pt') == 24.0
        # _chars 兄弟仍存在
        assert attrs.get('para.first_line_indent_chars') == 2  # 24pt / 12pt = 2 字

    def test_hanging_indent_pt_in_attrs(self):
        from docx import Document
        from docx.shared import Pt
        from thesis_worker.extractor.pipeline import _read_paragraph_style_attrs
        doc = Document()
        para = doc.add_paragraph('测试')
        para.paragraph_format.first_line_indent = Pt(-14)
        attrs = _read_paragraph_style_attrs(para)
        assert attrs.get('para.hanging_indent_pt') == 14.0

    def test_line_spacing_type_at_least(self):
        from docx import Document
        from docx.oxml.ns import qn
        from lxml import etree
        from thesis_worker.extractor.pipeline import _read_paragraph_style_attrs
        doc = Document()
        para = doc.add_paragraph('测试')
        pPr = para._element.get_or_add_pPr()
        spacing = etree.SubElement(pPr, qn('w:spacing'))
        spacing.set(qn('w:lineRule'), 'atLeast')
        spacing.set(qn('w:line'), '560')  # 28pt
        attrs = _read_paragraph_style_attrs(para)
        assert attrs.get('para.line_spacing_type') == 'atLeast'
        assert attrs.get('para.line_spacing_pt') == 28.0

    def test_line_spacing_type_single_auto(self):
        from docx import Document
        from docx.oxml.ns import qn
        from lxml import etree
        from thesis_worker.extractor.pipeline import _read_paragraph_style_attrs
        doc = Document()
        para = doc.add_paragraph('测试')
        pPr = para._element.get_or_add_pPr()
        spacing = etree.SubElement(pPr, qn('w:spacing'))
        spacing.set(qn('w:lineRule'), 'auto')
        spacing.set(qn('w:line'), '240')
        attrs = _read_paragraph_style_attrs(para)
        assert attrs.get('para.line_spacing_type') == 'single'
        assert attrs.get('para.line_spacing') == 1.0


class TestRealDocxRegression:
    """T5.1: 真实毕设规范模板抽取回归（依赖 /tmp/p4-real-test/spec.docx）"""

    SPEC = '/tmp/p4-real-test/spec.docx'

    def setup_method(self):
        """跳过 if 文件不存在（CI 环境无文件）"""
        import os
        import pytest
        if not os.path.exists(self.SPEC):
            pytest.skip(f'真实规范模板未就位：{self.SPEC}（用 textutil 转换毕设 .doc）')

    def test_extract_all_does_not_crash(self):
        """真实模板不应崩溃（防 None 防护回归）"""
        from thesis_worker.extractor.pipeline import extract_all
        result = extract_all(self.SPEC)
        assert isinstance(result, dict)
        assert 'rules' in result

    def test_real_docx_field_hit_rate_above_50pct(self):
        """真实文档字段命中率应 ≥ 50%（之前 54%；改造后期望保持或提升）"""
        from thesis_worker.extractor.pipeline import extract_all
        result = extract_all(self.SPEC)
        from thesis_worker.engine_v2.field_defs import FIELD_DEFS
        hit_rate = len(result['rules']) / len(FIELD_DEFS)
        assert hit_rate >= 0.50, f'命中率 {hit_rate:.1%} 低于 50%'

    def test_natural_language_extraction_hits(self):
        """规范文档"段前 N 磅"等关键词应触发 _pt attr"""
        from thesis_worker.extractor.pipeline import extract_all
        result = extract_all(self.SPEC)
        # 至少一个字段应含本批新加的 attr
        all_attrs = set()
        for fid, cfg in result['rules'].items():
            all_attrs.update(cfg.get('value', {}).keys())
        new_attrs = {
            'para.space_before_pt', 'para.space_after_pt',
            'para.first_line_indent_pt', 'para.hanging_indent_pt',
            'para.letter_spacing_pt',
            'para.line_spacing_type', 'para.line_spacing_pt',
        }
        # 至少 1 个新 attr 命中
        hit_new = all_attrs & new_attrs
        assert len(hit_new) >= 1, f'真实规范模板应命中至少 1 个新 attr；实际新 attr 命中数=0；总抽取 attr 数={len(all_attrs)}'
