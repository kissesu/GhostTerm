"""
@file: test_pipeline_attrs.py
@description: 测试 _read_paragraph_style_attrs 新增的行距/字间距/段前段后抽取
@author: Atlas.oi
@date: 2026-04-18
"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from thesis_worker.extractor.pipeline import _read_paragraph_style_attrs


def test_line_spacing_multiplier():
    """段落设置 1.5 倍行距 → para.line_spacing = 1.5"""
    doc = Document()
    para = doc.add_paragraph('测试段落')
    para.paragraph_format.line_spacing = 1.5
    attrs = _read_paragraph_style_attrs(para)
    assert attrs.get('para.line_spacing') == 1.5


def test_space_before_after_lines():
    """段前 12pt 段后 6pt → space_before_lines=1.0, space_after_lines=0.5"""
    doc = Document()
    para = doc.add_paragraph('测试段落')
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    attrs = _read_paragraph_style_attrs(para)
    assert attrs.get('para.space_before_lines') == 1.0
    assert attrs.get('para.space_after_lines') == 0.5


def test_letter_spacing_chars_from_xml():
    """run 的 w:spacing val=240 → letter_spacing_chars=1.0（1 字宽 @ 12pt）"""
    doc = Document()
    para = doc.add_paragraph('摘要')
    run = para.runs[0]
    # 注入 w:spacing XML 节点到 rPr
    rpr = run._element.get_or_add_rPr()
    spacing = rpr.makeelement(qn('w:spacing'), {qn('w:val'): '240'})
    rpr.append(spacing)
    attrs = _read_paragraph_style_attrs(para)
    assert attrs.get('para.letter_spacing_chars') == 1.0


def test_letter_spacing_chars_from_space_placeholder():
    """段落文本 "摘  要" 中间 2 空格 → letter_spacing_chars=2（fallback 路径）"""
    doc = Document()
    para = doc.add_paragraph('摘  要')
    attrs = _read_paragraph_style_attrs(para)
    assert attrs.get('para.letter_spacing_chars') == 2


def test_xml_spacing_wins_over_placeholder():
    """XML 字间距优先于空格占位 fallback（两路径结果必须不同才有判别力）"""
    doc = Document()
    para = doc.add_paragraph('摘  要')  # 文本有 2 空格占位（fallback 若命中会算出 2）
    run = para.runs[0]
    rpr = run._element.get_or_add_rPr()
    # 720 twips = 3 字宽 @ 12pt，与 fallback 的 2 明确错开，才能证明 XML 路径先命中
    spacing = rpr.makeelement(qn('w:spacing'), {qn('w:val'): '720'})
    rpr.append(spacing)
    attrs = _read_paragraph_style_attrs(para)
    # XML 路径先命中 → 3.0；若 fallback 抢先则会是 2
    assert attrs.get('para.letter_spacing_chars') == 3.0


def test_letter_spacing_chars_from_fullwidth_space_placeholder():
    """段落文本 "摘　要"（1 个全角空格 U+3000）→ letter_spacing_chars=1

    真实 Word 模板大量使用全角空格做字间距占位，fallback 正则必须同时接受
    半角 \\s 和全角 U+3000。
    """
    doc = Document()
    para = doc.add_paragraph('摘\u3000要')
    attrs = _read_paragraph_style_attrs(para)
    assert attrs.get('para.letter_spacing_chars') == 1


def test_no_attrs_when_not_set():
    """段落无任何样式设置 → 相关 key 缺席"""
    doc = Document()
    para = doc.add_paragraph('普通段落')
    attrs = _read_paragraph_style_attrs(para)
    assert 'para.line_spacing' not in attrs
    assert 'para.space_before_lines' not in attrs
    assert 'para.space_after_lines' not in attrs
    assert 'para.letter_spacing_chars' not in attrs
